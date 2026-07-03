from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

def normalizar_status_gerado(serie):
    """Converte a coluna 'Relatório Gerado' para bool de forma consistente."""
    valores_gerado = {"true", "verdadeiro", "sim", "s", "1", "yes", "y"}
    valores_pendente = {"false", "falso", "nao", "não", "n", "0", "no", ""}

    def parse(valor):
        if pd.isna(valor):
            return False
        if isinstance(valor, bool):
            return valor
        if isinstance(valor, (int, float)):
            return bool(valor) and valor != 0
        texto = str(valor).strip().lower()
        if texto in valores_gerado:
            return True
        if texto in valores_pendente:
            return False
        return False

    return serie.map(parse)


def adicionar_titulo_secao(doc, texto):
    """Adiciona um título de seção formatado."""
    secao = doc.add_paragraph()
    run = secao.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    secao.paragraph_format.space_before = Pt(12)
    secao.paragraph_format.space_after = Pt(6)


def ajustar_largura_colunas(caminho_planilha):
    """Ajusta a largura das colunas do Excel para caber o conteúdo."""
    wb = load_workbook(caminho_planilha)
    ws = wb.active

    for coluna in ws.columns:
        max_length = 0
        coluna_letra = coluna[0].column_letter

        for celula in coluna:
            try:
                if celula.value:
                    max_length = max(max_length, len(str(celula.value)))
            except:
                pass

        # Define largura da coluna com margem extra
        ajuste = max_length + 2
        ws.column_dimensions[coluna_letra].width = ajuste

    wb.save(caminho_planilha)


def arquivo_em_uso(caminho):
    """Verifica se arquivo está em uso/aberto por outro programa."""
    try:
        os.rename(caminho, caminho)
        return False
    except PermissionError:
        return True


def adicionar_texto_caixa_cinza(doc, texto, altura_cm=1.0):
    """Cria uma tabela 1x1 sem bordas com fundo cinza para simular uma caixa de texto."""
    from docx.shared import Cm
    from docx.enum.table import WD_ALIGN_VERTICAL
    table = doc.add_table(rows=1, cols=1)
    
    # Define a altura da linha
    table.rows[0].height = Cm(altura_cm)
    
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Aplica fundo cinza na célula
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9D9D9')
    tcPr.append(shd)
    
    # Adiciona o texto centralizado na célula
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    
    # Ajusta espaçamento interno do parágrafo para 0, já que a célula está centralizada verticalmente
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def formatar_data_extenso(data_val):
    """Converte datas (Timestamp, string, datetime) para o formato 'd de mês de ano' em português."""
    if pd.isna(data_val) or not data_val:
        return "data da assinatura eletrônica"
    
    try:
        from datetime import datetime
        if hasattr(data_val, "to_pydatetime"):
            dt = data_val.to_pydatetime()
        elif isinstance(data_val, datetime):
            dt = data_val
        else:
            data_str = str(data_val).strip()
            for fmt in ("%d/%m/%Y", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
                try:
                    dt = datetime.strptime(data_str, fmt)
                    break
                except ValueError:
                    continue
            else:
                return data_str
        
        meses = {
            1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
            5: "maio", 6: "junho", 7: "julho", 8: "agosto",
            9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
        }
        return f"{dt.day} de {meses[dt.month]} de {dt.year}"
    except Exception:
        return str(data_val)


def formatar_mes_ano(data_val):
    """Converte datas para o formato 'Mês, Ano' em português (ex: 'Junho, 2026')."""
    if pd.isna(data_val) or not data_val:
        return "Dezembro, 2025"
    
    try:
        from datetime import datetime
        if hasattr(data_val, "to_pydatetime"):
            dt = data_val.to_pydatetime()
        elif isinstance(data_val, datetime):
            dt = data_val
        else:
            data_str = str(data_val).strip()
            for fmt in ("%d/%m/%Y", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
                try:
                    dt = datetime.strptime(data_str, fmt)
                    break
                except ValueError:
                    continue
            else:
                return data_str
        
        meses = {
            1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril",
            5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto",
            9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"
        }
        return f"{meses[dt.month]}, {dt.year}"
    except Exception:
        return str(data_val)


def extrair_ano(data_val):
    """Extrai o ano de uma data."""
    if pd.isna(data_val) or not data_val:
        return "2026"
    try:
        from datetime import datetime
        if hasattr(data_val, "to_pydatetime"):
            return str(data_val.year)
        elif isinstance(data_val, datetime):
            return str(data_val.year)
        else:
            data_str = str(data_val).strip()
            for fmt in ("%d/%m/%Y", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
                try:
                    dt = datetime.strptime(data_str, fmt)
                    return str(dt.year)
                except ValueError:
                    continue
            if len(data_str) >= 4 and data_str[-4:].isdigit():
                return data_str[-4:]
            return "2026"
    except Exception:
        return "2026"
