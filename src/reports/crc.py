from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reports.base import BaseReport
import pandas as pd

class CrcReport(BaseReport):
    @property
    def key(self) -> str:
        return "CRC"

    @property
    def display_name(self) -> str:
        return "CRC"

    @property
    def default_contrato(self) -> str:
        return "CGPE – 001/2006"

    @property
    def capa_orgao_concedente(self) -> str:
        return "GOVERNO DO ESTADO DE PERNAMBUCO"

    @property
    def capa_secretaria(self) -> str:
        return "SECRETARIA DE PLANEJAMENTO, GESTÃO E DESENVOLVIMENTO REGIONAL - SEPLAG"

    @property
    def capa_titulo(self) -> str:
        return "RELATÓRIO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL"

    @property
    def capa_ctr_number_template(self) -> str:
        return "RELATÓRIO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL CTR Nº 05/{ano}"

    @property
    def capa_prestador_label(self) -> str:
        return "PRESTADOR DE SERVIÇO: CONCESSIONÁRIA ROTA DOS COQUEIROS (CRC)"

    @property
    def capa_contrato_label(self) -> str:
        return "CONTRATO DE CONCESSÃO PATROCINADA"

    @property
    def capa_regulador_label(self) -> str:
        return "AGÊNCIA DE REGULAÇÃO DOS SERVIÇOS PÚBLICOS DELEGADOS DO ESTADO DE PERNAMBUCO - ARPE"

    def get_capa_titulos(self, row, ano) -> list:
        return [
            "FISCALIZAÇÃO TÉCNICO-OPERACIONAL NO SISTEMA VIÁRIO DO PAIVA (PE – 024)",
            "PRESTADOR DE SERVIÇO: CONCESSIONÁRIA ROTA DOS COQUEIROS (CRC)"
        ]

    @property
    def sumario_before_abreviaturas(self) -> bool:
        return True

    def get_abbreviations(self) -> list:
        return [
            ("ABNT", "Associação Brasileira de Normas Técnicas"),
            ("ARPE", "Agência de Regulação de Pernambuco"),
            ("SEPPE", "Secretaria Executiva de Parcerias e Projetos Estratégicos"),
            ("NC", "Não Conformidade"),
            ("CRC", "Concessionária Rota dos Coqueiros S. A."),
            ("PER", "Programa de Exploração da Rodovia, anexo ao Contrato de Concessão Patrocinada CGPE-001/2006"),
            ("VI", "Verificador Independente, atualmente, o Consórcio formado pelas empresas Maciel Consultores S/S Ltda e Estratégica Serviços de Engenharia Consultiva Ltda")
        ]

    def get_sumario_linhas(self, row) -> list:
        return [
            "1.\tINTRODUÇÃO\t4",
            "2.\tOBJETIVO\t4",
            "3.\tINFORMAÇÕES GERAIS\t4",
            "4.\tMETODOLOGIA\t5",
            "5.\tFISCALIZAÇÃO\t7",
            "6.\tCONCLUSÕES\t11",
            "\tAPÊNDICE ÚNICO - MEMORIAL FOTOGRÁFICO - FISCALIZAÇÃO\t12"
        ]

    def get_intro_paragraphs(self, row, ano, data_extenso) -> list:
        return [
            [(f"A Coordenadoria de Transportes e Rodovias da ARPE, em cumprimento ao cronograma de fiscalização de {ano}, realizou fiscalização na Concessão Patrocinada da Ponte de Acesso e Sistema Viário da Praia do Paiva, sob responsabilidade da Concessionária Rota dos Coqueiros (CRC). A ação foi comunicada à concessionária por meio do Ofício ARPE/DTO nº 145/2026, de 26/05/2026 (Doc. SEI nº 86328395), e à Secretaria Executiva de Parcerias e Projetos Estratégicos (SEPPE) por meio do Ofício ARPE/DTO nº 146/2026, de 26/05/2026 (Doc. SEI nº 86328735).", False, False, None)],
            [("As ações de fiscalização registradas neste Relatório se referem à Rodovia Estadual PE-024, trecho Barra de Jangada – Itapuama, medindo 6,5 Km de extensão, compreendendo duas praças de pedágio, pela Ponte sobre o Rio Jaboatão e a via principal do Destino de Turismo e Lazer Praia do Paiva.", False, False, None)],
            [("Este relatório apresenta as observações realizadas pela equipe da Coordenadoria de Transportes e Rodovias da ARPE, sob a perspectiva técnico-operacional, com o objetivo de verificar as condições de operação, conservação, manutenção e segurança del trecho rodoviário concedido, bem como o cumprimento da legislação aplicável e a eficiência dos serviços prestados.", False, False, None)],
            [("As fiscalizações realizadas pela ARPE possuem, prioritariamente, caráter orientativo e corretivo, visando contribuir para o aperfeiçoamento dos procedimentos e para a correção de não conformidades identificadas na infraestrutura disponibilizada e nos serviços concedidos pelo Estado.", False, False, None)]
        ]

    def get_objective_paragraphs(self, row) -> list:
        return [
            [("A fiscalização direta e periódica do Sistema Viário do Paiva, concedido à CRC, tem por objetivo verificar as condições de operação, manutenção, segurança viária e níveis de serviço, bem como identificar não conformidades, subsidiar a adoção de medidas corretivas e assegurar o cumprimento das disposições contratuais, regulamentares e normativas aplicáveis, visando à preservação da segurança, da qualidade dos serviços e da adequada prestação do serviço público aos usuários.", False, False, None)],
            [("Nesse contexto, a ação fiscalizatória da ARPE avalia o grau de conformidade da infraestrutura e dos serviços prestados em relação ao Contrato de Concessão, à legislação vigente e às normas aplicáveis, podendo determinar ou recomendar medidas corretivas destinadas ao aprimoramento contínuo da qualidade dos serviços.", False, False, None)]
        ]

    def get_general_info_rows(self, row, responsaveis_formatted, periodo_val) -> list:
        return [
            ("3.1 DO PODER CONCEDENTE", "", True, False),
            ("Titular:", "Conselho do Programa de Parcerias Estratégicas de Pernambuco– CPPPE", False, False),
            ("Endereço:", "Rua da Moeda, nº 46, Recife-PE", False, False),
            ("Responsável:", "MARCELO BRUTO DA COSTA CORREIA", False, True),
            
            ("3.2 DO VERIFICADOR INDEPENDENTE", "", True, False),
            ("Verificador Independente:", "Consórcio Estratégica e Grupo Maciel", False, False),
            ("Endereço:", "Rua Irene Ramos Gomes de Mattos, Nº 176, Pina, Recife/PE CEP: 51011-530", False, False),
            ("Responsável:", "José Theodozio Netto", False, True),
            
            ("3.3 DO REGULADO", "", True, False),
            ("Regulado:", "Concessionária Rota dos Coqueiros (CRC)", False, False),
            ("Responsável:", "RAFAELA ELAINE DA COSTA LIMA ARAÚJO", False, True),
            ("Endereço:", "Rua Marmelo, s/nº - Praça de Pedágio – Barra de Jangada Jaboatão dos Guararapes – PE CEP: 54.495-760", False, False),
            ("Representantes por acompanhar:", "Larissa Almeida e Everton Albuquerque", False, False),
            
            ("3.4 DO FISCALIZADOR (ARPE)", "", True, False),
            ("Regulador:", "Agência de Regulação de Pernambuco (ARPE)", False, False),
            ("Diretor Presidente:", "CARLOS PORTO FILHO", False, True),
            ("Endereço:", "Avenida Conselheiro Rosa e Silva, 975, Aflitos, Recife/PE, CEP: 52.050-020. Estacionamento: Rua do Futuro, 150, Aflitos, Recife/PE.", False, False),
            
            ("Responsáveis pela fiscalização:", responsaveis_formatted, False, False),
            ("Período da Fiscalização:", periodo_val, False, False),
            ("Tipo de Fiscalização:", "Direta e periódica.", False, False)
        ]

    @property
    def general_info_col_widths(self) -> list:
        return [Inches(2.3), Inches(5.2)]

    @property
    def general_info_headers_indices(self) -> list:
        return [0, 4, 8, 13]

    def get_methodology_paragraphs(self, data_extenso) -> list:
        return [
            [(f"A Coordenadoria de Transporte e Rodovias da ARPE realizou fiscalização direta e periódica no Sistema Viário do Paiva em {data_extenso}, seguindo metodologia estruturada em três etapas: Preparação e Planejamento, Execução da Fiscalização e Monitoramento e Avaliação.", False, False, None)],
            [("Preparação e Planejamento", True, False, None),
             (" - compreende a organização e estruturação das atividades preliminares à execução da fiscalização, destacando-se a elaboração e o envio de avisos de fiscalização à Concessionária e demais atividades de suporte à fiscalização, bem como a análise de fiscalizações anteriores com a identificação de eventuais Não Conformidades pendentes.", False, False, None)],
            [("Execução da Fiscalização", True, False, None),
             (" - as possíveis Não Conformidades levantadas em campo são analisadas de acordo com os critérios elencados no PER (Contrato de Concessão). A execução da fiscalização é pautada por um arcabouço de normas e diretrizes, possibilitando que todas as etapas sejam desenvolvidas de maneira eficiente e em conformidade aos padrões estabelecidos, destacando-se:", False, False, None)]
        ]

    def get_references_bullets(self) -> list:
        return [
            ("Lei Estadual nº. 12.524, de 30 de dezembro de 2003, regulamentada pelo Decreto Estadual nº. 30.200, de 09 de fevereiro de 2007, que altera e consolida as disposições da Lei nº. 12.126, de 12 de dezembro de 2001, que cria a Agência de Regulação dos Serviços Públicos Delegados do Estado de Pernambuco - ARPE, e dá outras providências;", False, True),
            ("", False, False),
            ("Contrato de Concessão Patrocinada CGPE – 001/2006, de 28 de dezembro de 2006, e Termos Aditivos, para a exploração da ponte de acesso e sistema viário do destino de turismo e lazer Praia do Paiva, em conformidade com a Lei Federal nº 11.079/2004 e alterações; e a Lei Estadual nº 12.765/2005 e alterações.", False, True),
            ("", False, False),
            ("Resolução ARPE nº 083, de 30 de julho de 2013, que dispõe sobre os procedimentos de fiscalização, autuação e aplicação de penalidades aos prestadores de serviços públicos delegados no Estado de Pernambuco fiscalizados pela ARPE mediante delegação.", False, True),
            ("", False, False),
            ("Norma DNIT 005/2003 – que define os termos técnicos empregados em defeitos que ocorrem nos pavimentos flexíveis e semirrígidos e serve para padronizar a linguagem adotada na elaboração das normas, manuais, projetos e textos relativos aos pavimentos flexíveis e semirrígidos.", False, True)
        ]

    @property
    def references_left_indent_pt(self) -> float:
        return 53.4

    def get_post_methodology_paragraphs(self, total_ncs) -> list:
        return [
            [("Registra-se que as Não conformidades (NC) são codificadas de acordo com os seguintes níveis de informação, separados por “.” (ponto):", False, False, None)]
        ]

    def get_levels_data(self) -> dict:
        return {
            "niveis": [
                "Nível 1 - três dígitos, caracterizando a concessionária, nesse caso \"CRC\";",
                "Nível 2 - composto por cinco dígitos, os dois primeiros indicam a subdivisão utilizada na rodovia, neste caso, Segmento Homogêneo (SH) e os três últimos dígitos identificam a subdivisão utilizada no Contrato (nesse caso de 001 a 015);",
                "Nível 3 – composto por nove dígitos os quatro primeiros e os quatro últimos delimitam a localização das NC em escala de 0,01Km e o dígito intermediário informa se a NC é pontual \"+\" ou distribuída \"-\";",
                "Nível 4 – informa o ano da constatação da NC com quatro dígitos, antecedido de “/”; e",
                "Nível 5 - contém o sequencial numérico da NC no ano, expresso em três dígitos."
            ],
            "ex_str": "CRC.SH015.0646+0648/2026.001",
            "ex_desc": "Indica que esta Não Conformidade foi apontada para a CRC, no Segmento Homogêneo 15, concentrado na praça de pedágio do km 6,46 ao km 6,48, sendo a primeira NC registrada pela ARPE em 2026."
        }

    def get_post_metodologia_extra_paragraphs(self, row, data_extenso, total_achados) -> tuple:
        return (None, [])

    def render_quadros(self, doc, row, nc_df, criar_tabela_quadros_fn) -> None:
        from utils import formatar_data_extenso, adicionar_titulo_secao
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        data_extenso = formatar_data_extenso(row["Data"])
        
        # 1. Título da Seção
        adicionar_titulo_secao(doc, self.quadros_section_title)
        doc.add_paragraph() # Pula linha abaixo do título
        
        # 2. Parágrafos introdutórios
        intro_paragraphs = self.get_quadro_intro_paragraphs(row, data_extenso, "")
        for paragraph_runs in intro_paragraphs:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            for text, bold, italic, _ in paragraph_runs:
                run = p.add_run(text)
                run.font.name = 'Aptos'
                run.font.size = Pt(11)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
        
        doc.add_paragraph() # Parágrafo vazio
        
        # Obter dados de NC
        id_fisc = row["ID da Fiscalização"]
        current_ncs = nc_df[nc_df["ID da Fiscalização"] == id_fisc] if nc_df is not None and not nc_df.empty and "ID da Fiscalização" in nc_df.columns else pd.DataFrame()
        ncs_reais = pd.DataFrame()
        if not current_ncs.empty and "Não Conformidade" in current_ncs.columns:
            ncs_reais = current_ncs[current_ncs["Não Conformidade"].fillna("").astype(str).str.strip() != ""].copy()
            
        try:
            dt_obj = pd.to_datetime(row["Data"])
            data_abreviada = dt_obj.strftime("%d/%m/%Y")
        except Exception:
            data_abreviada = "27/05/2026"
            
        # 3. Quadro 1 title
        p7 = doc.add_paragraph()
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p7.paragraph_format.space_after = Pt(6)
        r7_1 = p7.add_run("QUADRO 1")
        r7_1.bold = True
        r7_1.font.name = 'Aptos'
        r7_1.font.size = Pt(11)
        r7_2 = p7.add_run(f" – NÃO CONFORMIDADES IDENTIFICADAS CRC - {data_abreviada}")
        r7_2.bold = True
        r7_2.font.name = 'Aptos'
        r7_2.font.size = Pt(11)
        
        # 4. Render Table
        criar_tabela_quadros_fn(doc, ncs_reais, is_pa=False, report_config=self)
        doc.add_paragraph() # Parágrafo vazio
        
        # 5. Nota abaixo da tabela
        p8 = doc.add_paragraph()
        p8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p8.paragraph_format.space_before = Pt(6)
        p8.paragraph_format.space_after = Pt(6)
        p8.paragraph_format.line_spacing = 1.15
        run8 = p8.add_run("É importante destacar que as Não Conformidades apontadas se referem à segurança dos pedestres na rodovia, visando evitar a ocorrência de acidentes.")
        run8.font.name = 'Aptos'
        run8.font.size = Pt(11)

    @property
    def quadros_section_title(self) -> str:
        return "5. FISCALIZAÇÃO"

    def get_quadro_intro_paragraphs(self, row, data_extenso, responsaveis_formatted) -> list:
        return [
            [(f"As NC identificadas, em {data_extenso}, pela equipe de fiscalização da ARPE estão detalhadas no Quadro 1 a seguir.", False, False, None)]
        ]

    @property
    def quadro_title_template(self) -> str:
        return "QUADRO 1 – NÃO CONFORMIDADES IDENTIFICADAS CRC - {data_abreviada}"

    @property
    def nc_table_headers(self) -> list:
        return [
            "IDENTIFICAÇÃO", 
            "DESCRIÇÃO DA EVIDÊNCIA", 
            "REGISTRO FOTOGRÁFICO", 
            "FUNDAMENTO DA NC (PER)", 
            "DETERMINAÇÃO"
        ]

    @property
    def nc_table_col_widths(self) -> list:
        return [Inches(1.34), Inches(1.68), Inches(0.81), Inches(2.46), Inches(0.98)]

    def format_nc_table_total_row(self, table, row_idx, total_ncs) -> None:
        r_total = table.rows[row_idx]
        r_total.cells[0].merge(r_total.cells[1]).merge(r_total.cells[2])
        r_total.cells[0].text = "TOTAL"
        r_total.cells[3].text = ""
        r_total.cells[4].text = str(total_ncs)
        
        col_widths = self.nc_table_col_widths
        r_total.cells[0].width = col_widths[0] + col_widths[1] + col_widths[2]
        r_total.cells[3].width = col_widths[3]
        r_total.cells[4].width = col_widths[4]

    @property
    def finalizacao_sections_config(self) -> dict:
        return {
            "6": "conclusões"
        }

    def get_determinations_paragraphs(self, total_ncs) -> list:
        return []

    def get_recommendations_paragraphs(self) -> list:
        return []

    def get_conclusions_paragraphs(self, total_ncs, local_val) -> list:
        from sections.finalizacao.finalizacao import numero_por_extenso
        f_extenso = {
            1: "uma", 2: "duas", 3: "três", 4: "quatro", 5: "cinco",
            6: "seis", 7: "sete", 8: "oito", 9: "nove", 10: "dez"
        }
        extenso_f = f_extenso.get(total_ncs, numero_por_extenso(total_ncs))
        extenso_com_parenteses = f"{total_ncs} ({extenso_f})" if total_ncs > 0 else "0 (zero)"
        return [
            [(f"Em decorrência das ações de fiscalização realizadas pela ARPE foram identificadas {extenso_com_parenteses} novas Não Conformidades na fiscalização do trecho rodoviário sob a responsabilidade da Concessionária Rota dos Coqueiros.", False, False, None)],
            [("Por fim, solicita-se o encaminhamento deste Relatório de Fiscalização à Concessionária Rota dos Coqueiros, para adoção das providências necessárias à regularização das não conformidades apontadas por esta Agência Reguladora, bem como dar conhecimento à SEPPE, gestora do Contrato de Concessão.", False, False, None)]
        ]

    def render_apendices(self, doc, row, ncs_reais, pas_reais, fotos_dir, data_fisc, ano, criar_grade_fotos_fn) -> None:
        from utils import adicionar_titulo_secao
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        p_ap_a = adicionar_titulo_secao(doc, f"APÊNDICE ÚNICO - MEMORIAL FOTOGRÁFICO - FISCALIZAÇÃO EM {data_fisc}")
        p_ap_a.paragraph_format.page_break_before = True
        
        total_ncs_val = len(ncs_reais)
        if total_ncs_val == 1:
            fotos_str = "foto 01"
        elif total_ncs_val == 2:
            fotos_str = "fotos 01 e 02"
        else:
            fotos_str = f"fotos 01 a {str(total_ncs_val).zfill(2)}"
            
        p_intro_fotos = doc.add_paragraph()
        p_intro_fotos.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_intro_fotos.paragraph_format.space_before = Pt(6)
        p_intro_fotos.paragraph_format.space_after = Pt(12)
        p_intro_fotos.paragraph_format.line_spacing = 1.15
        run_intro_fotos = p_intro_fotos.add_run(
            f"Estão evidenciadas a seguir as Não Conformidades apontadas neste Relatório de Fiscalização Técnico-Operacional "
            f"ARPE/CTR Nº 05/{ano} ({fotos_str})."
        )
        run_intro_fotos.font.name = 'Aptos'
        run_intro_fotos.font.size = Pt(11)
        
        if not ncs_reais.empty:
            criar_grade_fotos_fn(doc, ncs_reais, row.get("Local", ""), fotos_dir, data_fisc, self.key)
        else:
            p_empty = doc.add_paragraph()
            r_empty = p_empty.add_run("Nenhum registro fotográfico de não conformidade cadastrado.")
            r_empty.font.name = 'Aptos'
            r_empty.font.size = Pt(11)

    @property
    def analyst_title(self) -> str:
        return "Analista de Regulação"

    def get_process_sei_texts(self, ano) -> list:
        return [
            f"RELATÓRIO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL PROC ADM Nº 05/{ano} - CTR",
            f"SEI Nº xxxxxxxxxxxxxxxxxxxxxxx"
        ]
