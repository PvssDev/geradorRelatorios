# -*- coding: utf-8 -*-
from reports.crc import CrcReport
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
import pandas as pd


class CrcMonitoramentoReport(CrcReport):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.N_prev = "X"
        self.N_curr = "X"
        self.ctr_num = "XX/XXXX"

    @property
    def key(self) -> str:
        return "CRC_MONITORAMENTO"

    @property
    def display_name(self) -> str:
        return "CRC (Monitoramento)"

    @property
    def sumario_before_abreviaturas(self) -> bool:
        return False

    def get_abbreviations(self) -> list:
        return []

    # ------------------------------------------------------------------
    # Capa customizada exata do modelo de referência
    # ------------------------------------------------------------------
    def gerar_capa_monitoramento(self, doc, logo_path, row, documento_anterior):
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import os
        import re
        from docx import Document

        # 0. Ajuste de Estilo e Margens (Explicitas de 0.5 polegadas conforme a referência)
        style = doc.styles['Normal']
        style.font.name = 'Aptos'
        
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        # 1. Extração do número do monitoramento anterior, CTR e dados do Quadro 1 anterior
        self.N_prev = "X"
        self.ctr_num = "XX/XXXX"
        self.N_ncs_pendentes_prev = "X"
        self.data_vistoria_prev = "XX/XX/XXXX"
        self.processo_sei_prev = "XXXXXXXX"
        self.oficio_num_prev = "xxx/xxxx"
        self.oficio_data_prev = "xx/xx/xxxx"
        self.carta_num_prev = "xxxx/xxxx"
        self.carta_data_prev = "xx/xx/xxxx"
        self.carta_sei_prev = "xxxxxxxx"
        self.documento_anterior = documento_anterior
        
        if documento_anterior:
            try:
                # Se for um arquivo temporário/BytesIO do Streamlit, docx consegue ler
                prev_doc = Document(documento_anterior)
                
                # Buscar número do monitoramento anterior
                pattern = re.compile(r'(\d+)(?:\u00ba|\u00b0)?\s*monitoramento', re.IGNORECASE)
                for p in prev_doc.paragraphs:
                    m = pattern.search(p.text)
                    if m:
                        self.N_prev = int(m.group(1))
                        break
                
                # Buscar CTR
                pattern_ctr = re.compile(r'CTR\s*(?:N\u00ba|n\u00ba|N\u00ba)?\s*(\d+/\d+)', re.IGNORECASE)
                for p in prev_doc.paragraphs:
                    m = pattern_ctr.search(p.text)
                    if m:
                        self.ctr_num = m.group(1)
                        break

                # Extrair Processo SEI
                pattern_sei = re.compile(r'PROCESSO SEI\s*(?:N\u00ba|n\u00ba|N\u00ba)?\s*([\d\./-]+)', re.IGNORECASE)
                for p in prev_doc.paragraphs:
                    m = pattern_sei.search(p.text)
                    if m:
                        self.processo_sei_prev = m.group(1)
                        break
                
                # Extrair data da vistoria anterior a partir dos parágrafos (mais robusto)
                pattern_vist = re.compile(r'realizada na Rodovia Rota.*em\s*(\d{2}/\d{2}/\d{4})', re.IGNORECASE)
                for p in prev_doc.paragraphs:
                    m = pattern_vist.search(p.text)
                    if m:
                        self.data_vistoria_prev = m.group(1)
                        break
                        
                # Extrair Ofício anterior
                pattern_oficio = re.compile(r'Of(?:í|i)cio\s+Arpe\s+DTO\s+(?:n|N)(?:º|o)?\s*([\w/.-]+),\s*de\s*(\d{2}/\d{2}/\d{4})', re.IGNORECASE)
                for p in prev_doc.paragraphs:
                    m = pattern_oficio.search(p.text)
                    if m:
                        self.oficio_num_prev = m.group(1)
                        self.oficio_data_prev = m.group(2)
                        break
                        
                # Extrair Carta CRC anterior
                pattern_carta = re.compile(r'Carta\s+CRC/REG\s+(?:n|N)(?:º|o)?\s*([\w/.-]+),\s*de\s*(\d{2}/\d{2}/\d{4})\s*\(\s*Doc\.\s*SEI\s*(?:n|N)(?:º|o)?\s*([\w/.-]+)\)', re.IGNORECASE)
                for p in prev_doc.paragraphs:
                    m = pattern_carta.search(p.text)
                    if m:
                        self.carta_num_prev = m.group(1)
                        self.carta_data_prev = m.group(2)
                        self.carta_sei_prev = m.group(3)
                        break
                        
                # Extrair data da vistoria e quantidade de NCs pendentes do Quadro 1
                if len(prev_doc.tables) > 0:
                    quadro1_prev = prev_doc.tables[0]
                    # Se não encontrou a data de vistoria nos parágrafos, busca no cabeçalho do Quadro 1
                    if self.data_vistoria_prev == "XX/XX/XXXX":
                        try:
                            header_text = quadro1_prev.rows[0].cells[2].text
                            m_date = re.search(r'(\d{2}/\d{2}/\d{4})', header_text)
                            if m_date:
                                self.data_vistoria_prev = m_date.group(1)
                        except Exception:
                            pass
                    
                    # Contagem de Pendentes
                    try:
                        p_count = 0
                        for row_item in quadro1_prev.rows[1:]:
                            cells = row_item.cells
                            if len(cells) < 3:
                                continue
                            txt0 = cells[0].text.strip()
                            txt1 = cells[1].text.strip()
                            txt2 = cells[2].text.strip()
                            # Se for uma linha de título de pista mesclada, pula
                            if txt0 == txt1 == txt2:
                                continue
                            status_text = txt2.lower()
                            if "pendente" in status_text:
                                p_count += 1
                        self.N_ncs_pendentes_prev = p_count
                    except Exception:
                        pass
            except Exception as e:
                print(f"Erro ao extrair dados do documento anterior: {e}")

        if isinstance(self.N_prev, int):
            self.N_curr = self.N_prev + 1
        else:
            self.N_curr = "X"

        # Converte para strings para exibição
        n_curr_str = f"{self.N_curr}º" if self.N_curr != "X" else "Xº"

        def add_cover_p(text, bold=True, size_pt=12, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if text:
                run = p.add_run(text)
                run.font.name = "Aptos"
                run.font.size = Pt(size_pt)
                run.bold = bold
            return p

        # Cabeçalho Superior
        add_cover_p("", bold=False, size_pt=12, space_after=6)
        add_cover_p("", bold=False, size_pt=12, space_after=6)
        add_cover_p("COORDENADORIA DE TRANSPORTES E RODOVIAS", bold=True, size_pt=12, space_after=4)
        add_cover_p(f"RELATÓRIO DO {n_curr_str} MONITORAMENTO DO PROCESSO", bold=True, size_pt=12, space_after=4)
        add_cover_p(f"DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL CTR Nº {self.ctr_num}", bold=True, size_pt=12, space_after=12)

        # Imagem da Capa (logo_capa_crc.png)
        # Procuramos logo_capa_crc.png na pasta assets
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        logo_capa_crc = os.path.join(base_dir, "assets", "logo_capa_crc.png")
        if not os.path.exists(logo_capa_crc):
            logo_capa_crc = logo_path  # fallback
            
        if os.path.exists(logo_capa_crc):
            doc.add_picture(logo_capa_crc, width=Inches(5.90))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
            doc.paragraphs[-1].paragraph_format.space_before = Pt(12)

        # Títulos do Meio
        add_cover_p(
            f"VISTORIA TÉCNICA SOBRE AS NÃO CONFORMIDADES APONTADAS NO RELATÓRIO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL CTR {self.ctr_num}.",
            bold=True, size_pt=11, space_after=6
        )
        add_cover_p("", bold=False, size_pt=11, space_after=6)  # Parágrafo em branco P5
        add_cover_p("CONTRATO DE CONCESSÃO PATROCINADA – CGPE Nº 001/2006", bold=True, size_pt=11, space_after=4)
        add_cover_p("EXPLORAÇÃO DA PONTE DE ACESSO E SISTEMA VIÁRIO", bold=True, size_pt=11, space_after=4)
        add_cover_p("DO DESTINO DE TURISMO E LAZER PRAIA DO PAIVA", bold=True, size_pt=11, space_after=6)
        add_cover_p("", bold=False, size_pt=11, space_after=6)  # Parágrafo em branco P9
        add_cover_p("PROCESSO SEI Nº 0030200023.004570/2025-45", bold=True, size_pt=11, space_after=6)
        add_cover_p("", bold=False, size_pt=11, space_after=6)  # Parágrafo em branco P11
        add_cover_p("Recife, data da assinatura eletrônica", bold=False, size_pt=11, space_after=0)

    # ------------------------------------------------------------------
    # Linhas do Sumário dinâmicas
    # ------------------------------------------------------------------
    def get_sumario_linhas(self, row) -> list:
        # Pega a data da vistoria por extenso no formato Mes/Ano
        from utils import formatar_mes_ano
        try:
            mes_ano = formatar_mes_ano(row["Data"]).upper()
        except Exception:
            mes_ano = "MARÇO/2026"

        return [
            "1.\tINTRODUÇÃO\t4",
            "2.\tOBJETIVO\t4",
            "3.\tRESULTADOS DAS VISTORIAS DAS NÃO CONFORMIDADES\t5",
            f"4.\tRESUMO DA SITUAÇÃO DAS NÃO CONFORMIDADES MONITORADAS ({mes_ano})\t8",
            "5.\tCONCLUSÕES E RECOMENDAÇÕES\t9",
            f"\tAPÊNDÍCE ÚNCO - MEMORIAL FOTOGRÁFICO - {mes_ano}\t10",
        ]

    # ------------------------------------------------------------------
    # Método principal: renderiza as seções 1-4
    # ------------------------------------------------------------------
    def gerar_secoes_monitoramento(self, doc, row, nc_df, total_achados, documento_anterior=None):
        from utils import adicionar_titulo_secao, formatar_data_extenso, extrair_ano, formatar_mes_ano
        import re
        from docx import Document

        # Se N_prev ou ctr_num não foram calculados pela capa, calculamos aqui também
        if documento_anterior and (getattr(self, "N_prev", "X") == "X" or getattr(self, "ctr_num", "XX/XXXX") == "XX/XXXX"):
            try:
                prev_doc = Document(documento_anterior)
                pattern = re.compile(r'(\d+)(?:\u00ba|\u00b0)?\s*monitoramento', re.IGNORECASE)
                for p in prev_doc.paragraphs:
                    m = pattern.search(p.text)
                    if m:
                        self.N_prev = int(m.group(1))
                        self.N_curr = self.N_prev + 1
                        break
                
                pattern_ctr = re.compile(r'CTR\s*(?:N\u00ba|n\u00ba|N\u00ba)?\s*(\d+/\d+)', re.IGNORECASE)
                for p in prev_doc.paragraphs:
                    m = pattern_ctr.search(p.text)
                    if m:
                        self.ctr_num = m.group(1)
                        break
            except Exception as e:
                print(f"Erro ao extrair no gerar_secoes: {e}")

        N_prev = getattr(self, "N_prev", "X")
        N_curr = getattr(self, "N_curr", "X")
        ctr_num = getattr(self, "ctr_num", "XX/XXXX")

        # Converte para strings com ordinal
        n_prev_str = f"{N_prev}º" if N_prev != "X" else "Xº"
        n_curr_str = f"{N_curr}º" if N_curr != "X" else "Xº"

        data_extenso = formatar_data_extenso(row["Data"])
        try:
            from pandas import to_datetime
            data_vistoria = to_datetime(row["Data"]).strftime("%d/%m/%Y")
        except Exception:
            data_vistoria = "XX/XX/XXXX"
        id_fisc = str(row.get("ID da Fiscalização", "")).strip()
        mes_ano_cap = formatar_mes_ano(row["Data"]).upper()

        if nc_df is not None and not nc_df.empty and "ID da Fiscalização" in nc_df.columns:
            current_ncs = nc_df[nc_df["ID da Fiscalização"] == id_fisc].copy()
            if "Não Conformidade" in current_ncs.columns:
                ncs_reais = current_ncs[
                    current_ncs["Não Conformidade"].fillna("").astype(str).str.strip() != ""
                ].copy()
            else:
                ncs_reais = pd.DataFrame()
        else:
            ncs_reais = pd.DataFrame()

        def add_p(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(text)
            run.font.name = "Aptos"
            run.font.size = Pt(11)
            run.bold = bold
            return p

        def add_label_text(label, text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            run_l = p.add_run(label)
            run_l.bold = True
            run_l.font.name = "Aptos"
            run_l.font.size = Pt(11)
            run_t = p.add_run(text)
            run_t.font.name = "Aptos"
            run_t.font.size = Pt(11)
            return p

        # 1. INTRODUÇÃO
        adicionar_titulo_secao(doc, "1. INTRODUÇÃO")
        doc.add_paragraph()
        ncs_pend_prev = getattr(self, "N_ncs_pendentes_prev", "X")
        data_vist_prev = getattr(self, "data_vistoria_prev", "XX/XX/XXXX")
        
        processo_sei_prev = getattr(self, "processo_sei_prev", "XXXXXXXX")
        
        try:
            val_ncs = int(ncs_pend_prev)
        except Exception:
            val_ncs = 999
            
        if val_ncs == 1:
            nc_phrase = "1 Não Conformidade restava pendente"
        else:
            nc_phrase = f"{ncs_pend_prev} Não Conformidades restavam pendentes"
            
        oficio_num_prev = getattr(self, "oficio_num_prev", "xxx/xxxx")
        oficio_data_prev = getattr(self, "oficio_data_prev", "xx/xx/xxxx")
        carta_num_prev = getattr(self, "carta_num_prev", "xxxx/xxxx")
        carta_data_prev = getattr(self, "carta_data_prev", "xx/xx/xxxx")
        carta_sei_prev = getattr(self, "carta_sei_prev", "xxxxxxxx")
        
        add_p(
            f"Registra-se, preliminarmente, que na última vistoria técnica de monitoramento das Não Conformidades "
            f"apontadas no Relatório de Fiscalização Técnico-Operacional CTR {ctr_num}, realizada na Rodovia Rota dos "
            f"Coqueiros em {data_vist_prev}, verificou-se que {nc_phrase}. "
            f"Assim, o Relatório do {n_prev_str} Monitoramento (Doc. SEI nº {processo_sei_prev}) foi encaminhado à Concessionária "
            f"Rota dos Coqueiros (CRC) pelo Ofício Arpe DTO nº {oficio_num_prev}, de {oficio_data_prev}, reforçando a recomendação "
            f"para que a Concessionária mantenha a ARPE informada sobre as tratativas referentes às não conformidades pendentes."
        )
        add_p(
            f"Em sequência, a Concessionária, por meio da Carta CRC/REG nº {carta_num_prev}, de {carta_data_prev} "
            f"(Doc. SEI nº {carta_sei_prev}), enviou sua manifestação informando as ações tomadas referentes às Não Conformidades pendentes."
        )

        # 2. OBJETIVO
        adicionar_titulo_secao(doc, "2. OBJETIVO")
        doc.add_paragraph()
        add_p(
            f"Este {n_curr_str} Relatório de Monitoramento da ARPE tem como objetivo informar sobre o acompanhamento das ações da CRC "
            f"referentes às Não Conformidades pendentes do Relatório de Fiscalização Técnico-Operacional ARPE/CTR {ctr_num}, "
            f"de 27/05/2025, conforme as vistorias técnicas realizadas em {data_extenso}, considerando também o posicionamento da "
            f"Concessionária encaminhado à ARPE por meio da referida Carta CRC/REG nº {carta_num_prev}."
        )

        # 3. RESULTADOS DAS VISTORIAS
        adicionar_titulo_secao(doc, "3. RESULTADOS DAS VISTORIAS DAS NÃO CONFORMIDADES")
        doc.add_paragraph()
        add_p(
            f"Apresentam-se os resultados das vistorias técnicas realizadas em {data_extenso} na Rodovia Estadual PE-024 "
            f"referentes às Não Conformidades (NC) registradas no Relatório de Fiscalização Técnico-Operacional ARPE/CTR "
            f"nº {ctr_num}, organizadas em subitens pelo sentido da pista, associadas a cada Segmento Homogêneo (SH) "
            f"com indicação do ponto fiscalizado (KM)."
        )

        if ncs_reais.empty:
            add_p("Nenhuma Não Conformidade registrada para esta fiscalização.")
        else:
            if "Pista" in ncs_reais.columns:
                pistas_unicas = []
                for p_val in ncs_reais["Pista"].tolist():
                    p_str = str(p_val).strip() if not pd.isna(p_val) else ""
                    if p_str and p_str not in pistas_unicas:
                        pistas_unicas.append(p_str)
                if not pistas_unicas:
                    pistas_unicas = ["Unica"]
            else:
                pistas_unicas = ["Unica"]

            for sub_idx, pista_val in enumerate(pistas_unicas):
                sub_num = sub_idx + 1
                if pista_val == "Unica":
                    sub_title = "3.1 RODOVIA ESTADUAL PE – 024"
                    df_pista = ncs_reais
                else:
                    sub_title = f"3.{sub_num} RODOVIA ESTADUAL PE – 024 - PISTA {pista_val.upper()}"
                    mask = ncs_reais["Pista"].apply(
                        lambda x: str(x).strip() == pista_val if not pd.isna(x) else False
                    )
                    df_pista = ncs_reais[mask]

                adicionar_titulo_secao(doc, sub_title)
                doc.add_paragraph()

                for _, nc_row in df_pista.iterrows():
                    ident = str(nc_row.get("Identificação", "")).strip()
                    nc_desc = str(nc_row.get("Não Conformidade", "")).strip()
                    observacoes = str(nc_row.get("Observações", nc_row.get("Legenda da Foto", ""))).strip()
                    determinacao = str(nc_row.get("Determinação", "")).strip()
                    analise_arpe = str(nc_row.get("Análise ARPE", "")).strip()

                    if ident:
                        add_p(ident, bold=True)
                    add_label_text("NÃO CONFORMIDADE: ", nc_desc if nc_desc else "A ser preenchido.")
                    add_label_text("POSICIONAMENTO DA CRC: ", determinacao if determinacao else "A ser preenchido.")
                    add_label_text("CONSTATAÇÃO: ", observacoes if observacoes else "A ser preenchido.")
                    add_label_text("ANÁLISE ARPE: ", analise_arpe if analise_arpe else "A ser preenchido.")
                    doc.add_paragraph()

        # 4. RESUMO DA SITUAÇÃO
        adicionar_titulo_secao(doc, f"4. RESUMO DA SITUAÇÃO DAS NÃO CONFORMIDADES MONITORADAS ({mes_ano_cap})")
        doc.add_paragraph()
        add_p(
            f"Apresenta-se no Quadro 1, a seguir, um resumo da situação das Não Conformidades pendentes do Relatório "
            f"ARPE/CTR nº {ctr_num}, de acordo com as vistorias técnicas realizadas em {data_vistoria}, dando "
            f"continuidade ao monitoramento das respectivas soluções."
        )

    # ------------------------------------------------------------------
    # Seção de quadros
    # ------------------------------------------------------------------
    @property
    def quadros_section_title(self) -> str:
        return ""

    def render_quadros(self, doc, row, nc_df, criar_tabela_quadros_fn):
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from sections.quadros.quadros import set_cell_shading, set_cell_margins, set_table_borders
        import pandas as pd

        id_fisc = str(row.get("ID da Fiscalização", "")).strip()
        ctr_num = getattr(self, "ctr_num", "XX/XXXX")
        
        # Obter data da vistoria atual
        try:
            from pandas import to_datetime
            data_vistoria_atual = to_datetime(row["Data"]).strftime("%d/%m/%Y")
        except Exception:
            data_vistoria_atual = "XX/XX/XXXX"

        # Buscar Não Conformidades do preenchimento atual
        current_ncs = nc_df[nc_df["ID da Fiscalização"] == id_fisc] if nc_df is not None else pd.DataFrame()
        ncs_reais = pd.DataFrame()
        if not current_ncs.empty and "Não Conformidade" in current_ncs.columns:
            ncs_reais = current_ncs[
                current_ncs["Não Conformidade"].fillna("").astype(str).str.strip() != ""
            ].copy()

        # Adicionar o título do QUADRO 1
        p_q = doc.add_paragraph()
        p_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_q.paragraph_format.space_after = Pt(6)
        r_q1 = p_q.add_run("QUADRO 1")
        r_q1.bold = True
        r_q1.font.name = "Aptos"
        r_q1.font.size = Pt(11)
        r_q2 = p_q.add_run(f" – Situação das Não Conformidades Relatório ARPE/CTR Nº {ctr_num}")
        r_q2.bold = True
        r_q2.font.name = "Aptos"
        r_q2.font.size = Pt(11)

        # 1. Extrair linhas do Quadro 1 do documento anterior para manter id_nc e constatação fixos
        ncs_from_prev = []
        doc_anterior_path = getattr(self, "documento_anterior", None)
        if doc_anterior_path:
            try:
                prev_doc = Document(doc_anterior_path)
                if len(prev_doc.tables) > 0:
                    quadro1_prev = prev_doc.tables[0]
                    for row_item in quadro1_prev.rows[1:]:
                        cells = row_item.cells
                        if len(cells) < 3:
                            continue
                        txt0 = cells[0].text.strip()
                        txt1 = cells[1].text.strip()
                        txt2 = cells[2].text.strip()
                        if txt0 == txt1 == txt2:
                            ncs_from_prev.append({
                                "type": "section",
                                "text": txt0
                            })
                        elif txt0:
                            ncs_from_prev.append({
                                "type": "nc",
                                "id_nc": txt0,
                                "constatacao": txt1
                            })
            except Exception as e:
                print(f"Erro ao extrair linhas do Quadro 1 anterior: {e}")

        # Se não conseguiu extrair nada do anterior, usa fallbacks baseados nas NCs atuais
        if not ncs_from_prev:
            if ncs_reais.empty:
                p_empty = doc.add_paragraph()
                p_empty.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_empty.add_run("Nenhuma Não Conformidade registrada.")
                run.font.name = 'Aptos'
                run.font.size = Pt(10)
                return
            
            # Monta estrutura fallback baseada nas NCs do preenchimento atual
            pistas_ncs = {}
            if "Pista" in ncs_reais.columns:
                for _, nc_row in ncs_reais.iterrows():
                    p_val = str(nc_row.get("Pista", "")).strip() if not pd.isna(nc_row.get("Pista")) else ""
                    p_key = p_val if p_val else "Única"
                    if p_key not in pistas_ncs:
                        pistas_ncs[p_key] = []
                    pistas_ncs[p_key].append(nc_row)
            else:
                pistas_ncs["Única"] = [r for _, r in ncs_reais.iterrows()]

            for pista, items in pistas_ncs.items():
                if pista != "Única":
                    ncs_from_prev.append({
                        "type": "section",
                        "text": f"PE - 024, PISTA SENTIDO {pista.upper()}"
                    })
                for nc_row in items:
                    ncs_from_prev.append({
                        "type": "nc",
                        "id_nc": str(nc_row.get("Identificação", "")).strip(),
                        "constatacao": str(nc_row.get("Não Conformidade", "")).strip()
                    })

        # Criar tabela de 3 colunas
        table = doc.add_table(rows=1 + len(ncs_from_prev), cols=3)
        table.style = 'Table Grid'
        set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        table.allow_autofit = False

        # Configurar larguras das colunas
        col_widths = [Inches(1.15), Inches(3.60), Inches(1.06)]
        for r in table.rows:
            for c_idx, width in enumerate(col_widths):
                if c_idx < len(r.cells):
                    r.cells[c_idx].width = width

        # Cabeçalho Principal (Linha 0)
        headers = ["ID.NC(*)", "CONSTATAÇÃO (27/05/2025)", f"SITUAÇÃO {data_vistoria_atual}"]
        for c_idx, text in enumerate(headers):
            cell = table.cell(0, c_idx)
            set_cell_shading(cell, "BFBFBF")
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(text)
            run.font.name = "Aptos"
            run.font.size = Pt(10)
            run.bold = True

        # Preencher linhas
        for idx, item in enumerate(ncs_from_prev, 1):
            if item["type"] == "section":
                cell_start = table.cell(idx, 0)
                cell_end = table.cell(idx, 2)
                cell_start.merge(cell_end)
                
                set_cell_shading(cell_start, "FFFFCC")
                set_cell_margins(cell_start, top=100, bottom=100, left=150, right=150)
                cell_start.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell_start.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(item["text"])
                run.font.name = "Aptos"
                run.font.size = Pt(10)
                run.bold = True
            else:
                id_nc_val = item["id_nc"]
                constatacao_val = item["constatacao"]
                
                # Buscar a Situação correspondente do preenchimento do Streamlit
                situacao_val = "Pendente" # default se não houver preenchimento
                target_id = id_nc_val.strip().lower()
                for _, nc_row in ncs_reais.iterrows():
                    curr_id = str(nc_row.get("Identificação", "")).strip().lower()
                    if target_id == curr_id or target_id in curr_id or curr_id in target_id:
                        situacao_val = str(nc_row.get("Situação", "Pendente")).strip()
                        break
                
                # Coluna 0: ID.NC
                cell_id = table.cell(idx, 0)
                set_cell_margins(cell_id, top=100, bottom=100, left=150, right=150)
                cell_id.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p_id = cell_id.paragraphs[0]
                p_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_id.paragraph_format.space_before = Pt(0)
                p_id.paragraph_format.space_after = Pt(0)
                run_id = p_id.add_run(id_nc_val)
                run_id.font.name = "Aptos"
                run_id.font.size = Pt(10)
                run_id.bold = True

                # Coluna 1: Constatação
                cell_desc = table.cell(idx, 1)
                set_cell_margins(cell_desc, top=100, bottom=100, left=150, right=150)
                cell_desc.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p_desc = cell_desc.paragraphs[0]
                p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_desc.paragraph_format.space_before = Pt(0)
                p_desc.paragraph_format.space_after = Pt(0)
                p_desc.paragraph_format.line_spacing = 1.15
                run_desc = p_desc.add_run(constatacao_val)
                run_desc.font.name = "Aptos"
                run_desc.font.size = Pt(10)

                # Coluna 2: Situação
                cell_sit = table.cell(idx, 2)
                set_cell_margins(cell_sit, top=100, bottom=100, left=150, right=150)
                cell_sit.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p_sit = cell_sit.paragraphs[0]
                p_sit.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_sit.paragraph_format.space_before = Pt(0)
                p_sit.paragraph_format.space_after = Pt(0)
                run_sit = p_sit.add_run(situacao_val)
                run_sit.font.name = "Aptos"
                run_sit.font.size = Pt(10)
                run_sit.bold = True

        doc.add_paragraph()

        p_nota = doc.add_paragraph()
        p_nota.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_nota.paragraph_format.space_before = Pt(4)
        p_nota.paragraph_format.space_after = Pt(6)
        p_nota.paragraph_format.line_spacing = 1.15
        run_nota = p_nota.add_run(
            "(*) ID. NC: Identificação da Não Conformidade, composta com o Segmento Homogêneo (SH) "
            "e a quilometragem do ponto fiscalizado (KM)."
        )
        run_nota.font.name = "Aptos"
        run_nota.font.size = Pt(10)
        run_nota.italic = True

    # ------------------------------------------------------------------
    # Seção 5 — Conclusões e Recomendações
    # ------------------------------------------------------------------
    @property
    def finalizacao_sections_config(self) -> dict:
        return {"5": "conclusoes_monitoramento"}

    def get_conclusions_paragraphs(self, total_ncs, local_val) -> list:
        return []

    def get_conclusions_monitoramento_paragraphs(self, total_ncs, data_extenso) -> list:
        if total_ncs == 0:
            conc_nc = "todas as Não Conformidades foram sanadas"
            singular_plural_nc = "da Não Conformidade identificada"
        elif total_ncs == 1:
            conc_nc = "permanece uma Não Conformidade ainda sem solução definitiva"
            singular_plural_nc = "da Não Conformidade identificada"
        else:
            conc_nc = f"permanecem {total_ncs} Não Conformidades ainda sem solução definitiva"
            singular_plural_nc = "das Não Conformidades identificadas"

        N_curr = getattr(self, "N_curr", "X")
        n_curr_str = f"{N_curr}º" if N_curr != "X" else "Xº"
        ctr_num = getattr(self, "ctr_num", "XX/XXXX")

        return [
            [(
                f"As vistorias técnicas, realizadas em {data_extenso}, bem como das análises apresentadas neste "
                f"{n_curr_str} Relatório de Monitoramento do Relatório de Fiscalização Técnico-Operacional ARPE/CTR "
                f"nº {ctr_num}, indicam que {conc_nc}.",
                False, False, None
            )],
            [(
                "Recomenda-se que a Concessionária mantenha a ARPE informada sobre as interações com a "
                "Associação Geral da Reserva do Paiva para evitar definitivamente a presença de detritos no trecho concedido.",
                False, False, None
            )],
            [(
                "Diante desse resultado, considera-se necessário que a ARPE continue realizando vistorias até que a "
                f"Concessionária efetive a execução dos serviços para correção {singular_plural_nc}.",
                False, False, None
            )],
        ]

    # ------------------------------------------------------------------
    # Apêndice
    # ------------------------------------------------------------------
    def render_apendices(self, doc, row, ncs_reais, pas_reais, fotos_dir, data_fisc, ano, criar_grade_fotos_fn):
        from utils import adicionar_titulo_secao, formatar_mes_ano
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            mes_ano = formatar_mes_ano(row["Data"]).upper()
        except Exception:
            mes_ano = "MARÇO/2026"

        p_ap = adicionar_titulo_secao(doc, f"APÊNDÍCE ÚNCO - MEMORIAL FOTOGRÁFICO - {mes_ano}")
        p_ap.paragraph_format.page_break_before = True

        total_ncs_val = len(ncs_reais)
        if total_ncs_val == 0:
            fotos_str = ""
        elif total_ncs_val == 1:
            fotos_str = "foto 01"
        elif total_ncs_val == 2:
            fotos_str = "fotos 01 e 02"
        else:
            fotos_str = f"fotos 01 a {str(total_ncs_val).zfill(2)}"

        N_curr = getattr(self, "N_curr", "X")
        n_curr_str = f"{N_curr}º" if N_curr != "X" else "Xº"
        ctr_num = getattr(self, "ctr_num", "XX/XXXX")
        fotos_suffix = f" ({fotos_str})" if fotos_str else ""

        p_intro = doc.add_paragraph()
        p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_intro.paragraph_format.space_before = Pt(6)
        p_intro.paragraph_format.space_after = Pt(12)
        p_intro.paragraph_format.line_spacing = 1.15
        run_intro = p_intro.add_run(
            f"Apresenta-se, a seguir, as evidências da situação das Não Conformidades, de acordo com o sentido da pista "
            f"e as identificações utilizadas neste {n_curr_str} Relatório de Monitoramento do Processo de Fiscalização "
            f"Técnico-Operacional nº {ctr_num}{fotos_suffix}."
        )
        run_intro.font.name = "Aptos"
        run_intro.font.size = Pt(11)

        if not ncs_reais.empty:
            criar_grade_fotos_fn(doc, ncs_reais, row.get("Local", ""), fotos_dir, data_fisc, self.key)
        else:
            p_empty = doc.add_paragraph()
            r_empty = p_empty.add_run("Nenhuma foto de não conformidade anexada.")
            r_empty.font.name = "Aptos"
            r_empty.font.size = Pt(11)
