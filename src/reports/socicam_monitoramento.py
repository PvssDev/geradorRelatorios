# -*- coding: utf-8 -*-
from reports.base import BaseMonitoramentoMixin
from reports.socicam import SocicamReport
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document
import pandas as pd
import os
from monitoramento_utils import (
    extrair_metadados_anterior,
    extrair_linhas_quadro1_anterior,
    criar_tabela_quadro1_monitoramento,
)


class SocicamMonitoramentoReport(BaseMonitoramentoMixin, SocicamReport):
    @property
    def key(self) -> str:
        return "SOCICAM_MONITORAMENTO"

    @property
    def display_name(self) -> str:
        return "SOCICAM (Monitoramento)"

    @property
    def capa_titulo(self) -> str:
        return "RELATÓRIO DE MONITORAMENTO TÉCNICO-OPERACIONAL"

    @property
    def capa_ctr_number_template(self) -> str:
        return "RELATÓRIO DE MONITORAMENTO TÉCNICO-OPERACIONAL CTR Nº {mes_ano}"

    @property
    def sumario_before_abreviaturas(self) -> bool:
        return False

    @property
    def signatures_before_apendices(self) -> bool:
        return True

    def get_abbreviations(self) -> list:
        return []

    def get_process_sei_texts(self, row, ano=None) -> list:
        from utils import extrair_mes_ano_numerico, extrair_ano
        if isinstance(row, (str, int)) and ano is None:
            ano = str(row)
            mes_ano = f"01/{ano}"
        else:
            ano = ano or extrair_ano(row.get("Data", "") if hasattr(row, "get") else row["Data"])
            data_val = row.get("Data", "") if hasattr(row, "get") else (row["Data"] if isinstance(row, dict) or hasattr(row, "__getitem__") else "")
            mes_ano = extrair_mes_ano_numerico(data_val)
        return [
            f"RELATÓRIO DE MONITORAMENTO TÉCNICO-OPERACIONAL PROC ADM Nº {mes_ano} - CTR",
            f"SEI Nº 0030200023.002186/{ano}-99"
        ]

    def get_sumario_linhas(self, row, nc_df=None) -> list:
        from utils import formatar_mes_ano
        try:
            mes_ano = formatar_mes_ano(row["Data"]).upper()
        except Exception:
            mes_ano = "JULHO/2026"

        linhas = [
            "1.\tINTRODUÇÃO\t4",
            "2.\tOBJETIVO\t4",
            "3.\tRESULTADO DAS VISTORIAS DAS NÃO CONFORMIDADES\t5",
        ]

        id_fisc = str(row.get("ID da Fiscalização", "")).strip()
        ncs_reais = []
        if nc_df is not None and not nc_df.empty and "ID da Fiscalização" in nc_df.columns:
            current_ncs = nc_df[nc_df["ID da Fiscalização"] == id_fisc]
            if "Não Conformidade" in current_ncs.columns:
                ncs_reais = current_ncs[
                    current_ncs["Não Conformidade"].fillna("").astype(str).str.strip() != ""
                ].to_dict('records')

        if ncs_reais:
            for idx, nc_row in enumerate(ncs_reais, 1):
                term_val = str(nc_row.get("Terminal", nc_row.get("Local", "TIP"))).strip()
                if not term_val:
                    term_val = "PASSAGEIROS DO RECIFE (TIP)"
                
                t_upper = term_val.upper()
                if "GARANHUNS" in t_upper or "GAR" in t_upper:
                    term_name = "GARANHUNS (GAR)"
                elif "CARUARU" in t_upper or "CAR" in t_upper:
                    term_name = "CARUARU (CAR)"
                elif "RECIFE" in t_upper or "TIP" in t_upper:
                    term_name = "PASSAGEIROS DO RECIFE (TIP)"
                else:
                    term_name = f"{term_val.upper()}"

                sub_num = f"3.{idx}"
                sub_title = f"TERMINAL RODOVIÁRIO DE {term_name}"
                p_num = min(5 + (idx - 1) // 2, 7)
                linhas.append(f"{sub_num}\t{sub_title}\t{p_num}")
        else:
            linhas.append("3.1\tTERMINAL RODOVIÁRIO DE PASSAGEIROS DO RECIFE (TIP)\t5")

        linhas.extend([
            f"4.\tRESUMO DA SITUAÇÃO DAS NÃO CONFORMIDADES MONITORADAS ({mes_ano})\t8",
            "5.\tCONCLUSÕES E RECOMENDAÇÕES\t9",
            f"\tAPÊNDICE ÚNICO - MEMORIAL FOTOGRÁFICO - VISTORIAS REALIZADAS EM {mes_ano}\t10",
        ])

        return linhas

    # ------------------------------------------------------------------
    # Capa de Monitoramento SOCICAM
    # ------------------------------------------------------------------
    def gerar_capa_monitoramento(self, doc, logo_path, row, documento_anterior):
        style = doc.styles['Normal']
        style.font.name = 'Aptos'
        
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        self.documento_anterior = documento_anterior
        meta = extrair_metadados_anterior(documento_anterior)
        self.N_prev = meta["N_prev"]
        self.N_curr = meta["N_curr"]
        self.ctr_num = meta["ctr_num"] if meta["ctr_num"] != "XX/XXXX" else "04/2025"
        self.processo_sei_prev = meta["processo_sei_prev"]

        n_curr_str = f"{self.N_curr}º" if self.N_curr != "X" else "Xº"

        def add_cover_p(text, bold=True, size_pt=12, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if text:
                run = p.add_run(text)
                run.font.name = "Aptos"
                run.font.size = Pt(size_pt)
                run.bold = bold
            return p

        # Cabeçalho Superior
        add_cover_p("", bold=False, size_pt=12, space_after=6)
        add_cover_p("", bold=False, size_pt=12, space_after=6)
        add_cover_p("COORDENADORIA DE TRANSPORTES E RODOVIAS (CTR)", bold=True, size_pt=12, space_after=4)
        add_cover_p(f"RELATÓRIO DO {n_curr_str} MONITORAMENTO DAS NÃO CONFORMIDADES", bold=True, size_pt=12, space_after=4)
        add_cover_p(f"PROCESSO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL CTR Nº {self.ctr_num}", bold=True, size_pt=12, space_after=12)

        # Imagem da Capa (capa_monitoramento.png)
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        capa_mon = os.path.join(base_dir, "assets", "capa_monitoramento.png")
        if not os.path.exists(capa_mon):
            capa_mon = os.path.join(base_dir, "assets", "capa_1.png")
        if not os.path.exists(capa_mon):
            capa_mon = logo_path  # fallback
            
        if os.path.exists(capa_mon):
            doc.add_picture(capa_mon, width=Inches(5.90))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
            doc.paragraphs[-1].paragraph_format.space_before = Pt(12)

        # Títulos do Meio
        add_cover_p("PROCESSO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL DOS TERMINAIS RODOVIÁRIOS CONCEDIDOS À SOCICAM - ADMINISTRAÇÃO, PROJETOS E REPRESENTAÇÕES LTDA", bold=True, size_pt=11, space_after=6)
        add_cover_p("CONTRATO Nº 1.041.080/08", bold=True, size_pt=11, space_after=4)
        add_cover_p(f"PROCESSO SEI Nº {self.processo_sei_prev}", bold=True, size_pt=11, space_after=6)
        add_cover_p("", bold=False, size_pt=11, space_after=6)
        add_cover_p("Recife, data de assinatura eletrônica", bold=False, size_pt=11, space_after=0)

    # ------------------------------------------------------------------
    # Método principal: renderiza as seções 1-4
    # ------------------------------------------------------------------
    def gerar_secoes_monitoramento(self, doc, row, nc_df, total_achados, documento_anterior=None):
        from utils import adicionar_titulo_secao, formatar_data_extenso, formatar_mes_ano

        meta = extrair_metadados_anterior(documento_anterior)
        N_curr = meta["N_curr"] if meta["N_curr"] != "X" else getattr(self, "N_curr", "X")
        ctr_num = meta["ctr_num"] if meta["ctr_num"] != "XX/XXXX" else getattr(self, "ctr_num", "04/2025")

        n_curr_str = f"{N_curr}º" if N_curr != "X" else "Xº"

        data_extenso = formatar_data_extenso(row["Data"])
        try:
            from pandas import to_datetime
            data_vistoria = to_datetime(row["Data"]).strftime("%d/%m/%Y")
        except Exception:
            data_vistoria = "XX/XX/XXXX"

        id_fisc = str(row.get("ID da Fiscalização", "")).strip()
        mes_ano_cap = formatar_mes_ano(row["Data"]).upper()

        if nc_df is not None and not nc_df.empty and "ID da Fiscalização" in nc_df.columns:
            current_ncs = nc_df[nc_df["ID da Fiscalização"] == id_fisc].copy()
            if "Não Conformidade" in current_ncs.columns:
                ncs_reais = current_ncs[
                    current_ncs["Não Conformidade"].fillna("").astype(str).str.strip() != ""
                ].copy()
            else:
                ncs_reais = pd.DataFrame()
        else:
            ncs_reais = pd.DataFrame()

        def add_p(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(text)
            run.font.name = "Aptos"
            run.font.size = Pt(11)
            run.bold = bold
            return p

        def add_label_text(label, text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            run_l = p.add_run(label)
            run_l.bold = True
            run_l.font.name = "Aptos"
            run_l.font.size = Pt(11)
            run_t = p.add_run(text)
            run_t.font.name = "Aptos"
            run_t.font.size = Pt(11)
            return p

        # 1. INTRODUÇÃO
        adicionar_titulo_secao(doc, "1. INTRODUÇÃO")
        doc.add_paragraph()
        add_p(
            "O monitoramento das Não Conformidades apontadas nos relatórios de fiscalização é fundamental para acompanhar as ações "
            "desenvolvidas pelas concessionárias, garantindo o cumprimento adequado dos Contratos de Concessão."
        )
        add_p(
            f"Nesse contexto este {n_curr_str} Relatório de Monitoramento do Processo de Fiscalização Técnico-Operacional ARPE/CTR nº {ctr_num} "
            f"descreve a evolução das Não Conformidades listadas no Relatório de Fiscalização, referentes às vistorias realizadas nos Terminais "
            f"Rodoviários de Passageiros de Garanhuns, Caruaru e Recife (TIP) concedidos à Empresa SOCICAM - Administração, Projetos e Representações Ltda (SOCICAM)."
        )
        add_p(
            f"As informações disponibilizadas pela SOCICAM foram consideradas como referência para as vistorias realizadas em {data_extenso} "
            f"nos referidos Terminais Rodoviários."
        )

        # 2. OBJETIVO
        adicionar_titulo_secao(doc, "2. OBJETIVO")
        doc.add_paragraph()
        add_p(
            f"Este Relatório do {n_curr_str} Monitoramento objetiva apresentar os resultados das vistorias acerca das Não Conformidades "
            f"registradas no Relatório de Fiscalização Técnico-Operacional ARPE/CTR nº {ctr_num}, referentes aos Terminais Rodoviários "
            f"de Passageiros dos municípios de Garanhuns, Caruaru e Recife (TIP), conforme o cronograma encaminhado pela Concessionária SOCICAM."
        )

        # 3. RESULTADO DAS VISTORIAS
        adicionar_titulo_secao(doc, "3. RESULTADO DAS VISTORIAS DAS NÃO CONFORMIDADES")
        doc.add_paragraph()
        add_p(
            f"Estão registrados para cada Terminal Rodoviário os resultados da verificação pela ARPE das ações desenvolvidas pela SOCICAM, "
            f"visando solucionar as Não Conformidades apresentadas no Relatório de Fiscalização Técnico-Operacional ARPE/CTR nº {ctr_num}."
        )

        if ncs_reais.empty:
            add_p("Nenhuma Não Conformidade registrada para este monitoramento.")
        else:
            for sub_idx, (_, nc_row) in enumerate(ncs_reais.iterrows(), 1):
                term_val = str(nc_row.get("Terminal", nc_row.get("Local", "TIP"))).strip()
                if not term_val:
                    term_val = "PASSAGEIROS DO RECIFE (TIP)"
                
                t_upper = term_val.upper()
                if "GARANHUNS" in t_upper or "GAR" in t_upper:
                    term_name = "GARANHUNS (GAR)"
                elif "CARUARU" in t_upper or "CAR" in t_upper:
                    term_name = "CARUARU (CAR)"
                elif "RECIFE" in t_upper or "TIP" in t_upper:
                    term_name = "PASSAGEIROS DO RECIFE (TIP)"
                else:
                    term_name = f"{term_val.upper()}"

                sub_title = f"3.{sub_idx} – TERMINAL RODOVIÁRIO DE {term_name}"
                adicionar_titulo_secao(doc, sub_title)
                doc.add_paragraph()

                ident = str(nc_row.get("Identificação", "")).strip()
                nc_desc = str(nc_row.get("Não Conformidade", "")).strip()
                observacoes = str(nc_row.get("Observações", nc_row.get("Legenda da Foto", ""))).strip()
                determinacao = str(nc_row.get("Determinação", "")).strip()
                analise_arpe = str(nc_row.get("Análise ARPE", "")).strip()
                situacao = str(nc_row.get("Situação", "Pendente")).strip()

                title_text = f"Não Conformidade {ident} - {nc_desc}" if ident and not ident.lower().startswith("foto") else (nc_desc if nc_desc else "Não Conformidade")
                add_p(title_text, bold=True)

                add_label_text("Informação da SOCICAM: ", determinacao if determinacao else "A ser preenchido.")
                add_label_text("Constatação: ", observacoes if observacoes else "A ser preenchido.")

                if analise_arpe:
                    texto_analise = analise_arpe
                elif situacao:
                    texto_analise = f"Não Conformidade {situacao.upper()}."
                else:
                    texto_analise = "A ser preenchido."

                add_label_text("Análise da ARPE: ", texto_analise)
                doc.add_paragraph()

        # 4. RESUMO DA SITUAÇÃO
        adicionar_titulo_secao(doc, f"4. RESUMO DA SITUAÇÃO DAS NÃO CONFORMIDADES MONITORADAS ({mes_ano_cap})")
        doc.add_paragraph()
        add_p(
            f"Apresenta-se no Quadro 1, a seguir, um resumo da situação das Não Conformidades pendentes do Relatório ARPE/CTR nº {ctr_num}, "
            f"de acordo com as vistorias técnicas realizadas em {data_vistoria}, dando continuidade ao monitoramento das respectivas soluções."
        )

    # ------------------------------------------------------------------
    # Renderização dos Quadros (Quadro 1 de 3 colunas)
    # ------------------------------------------------------------------
    @property
    def quadros_section_title(self) -> str:
        return ""

    def render_quadros(self, doc, row, nc_df, criar_tabela_quadros_fn):
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        id_fisc = str(row.get("ID da Fiscalização", "")).strip()
        ctr_num = getattr(self, "ctr_num", "04/2025")

        current_ncs = nc_df[nc_df["ID da Fiscalização"] == id_fisc] if nc_df is not None and not nc_df.empty and "ID da Fiscalização" in nc_df.columns else pd.DataFrame()
        ncs_reais = pd.DataFrame()
        if not current_ncs.empty and "Não Conformidade" in current_ncs.columns:
            ncs_reais = current_ncs[
                current_ncs["Não Conformidade"].fillna("").astype(str).str.strip() != ""
            ].copy()

        # Título do Quadro 1
        p_q = doc.add_paragraph()
        p_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_q.paragraph_format.space_after = Pt(6)
        r_q1 = p_q.add_run("QUADRO 1")
        r_q1.bold = True
        r_q1.font.name = "Aptos"
        r_q1.font.size = Pt(11)
        r_q2 = p_q.add_run(f" – Resumo da Situação das Não Conformidades Relatório ARPE/CTR Nº {ctr_num}")
        r_q2.bold = True
        r_q2.font.name = "Aptos"
        r_q2.font.size = Pt(11)

        # Usar funções compartilhadas
        doc_anterior_path = getattr(self, "documento_anterior", None)
        ncs_from_prev = extrair_linhas_quadro1_anterior(doc_anterior_path)

        headers = ["NÃO CONFORMIDADE", "INFORMAÇÃO SOCICAM", "SITUAÇÃO"]
        col_widths = [Inches(3.20), Inches(3.20), Inches(1.17)]
        criar_tabela_quadro1_monitoramento(doc, headers, ncs_from_prev, ncs_reais, col_widths)

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Seção 5 — Conclusões e Recomendações
    # ------------------------------------------------------------------
    @property
    def finalizacao_sections_config(self) -> dict:
        return {"5": "conclusoes_monitoramento"}

    def get_conclusions_monitoramento_paragraphs(self, total_ncs, data_extenso) -> list:
        if total_ncs == 0:
            conc_nc = "todas as Não Conformidades foram sanadas"
            singular_plural_nc = "da Não Conformidade identificada"
        elif total_ncs == 1:
            conc_nc = "permanece uma Não Conformidade ainda sem solução definitiva"
            singular_plural_nc = "da Não Conformidade identificada"
        else:
            conc_nc = f"permanecem {total_ncs} Não Conformidades ainda sem solução definitiva"
            singular_plural_nc = "das Não Conformidades identificadas"

        N_curr = getattr(self, "N_curr", "X")
        n_curr_str = f"{N_curr}º" if N_curr != "X" else "Xº"
        ctr_num = getattr(self, "ctr_num", "04/2025")

        return [
            [(
                f"As vistorias técnicas, realizadas em {data_extenso}, bem como as análises apresentadas neste "
                f"{n_curr_str} Relatório de Monitoramento do Relatório de Fiscalização Técnico-Operacional ARPE/CTR "
                f"nº {ctr_num}, indicam que {conc_nc}.",
                False, False, None
            )],
            [(
                "Diante desse resultado, considera-se necessário que a ARPE continue realizando vistorias até que a "
                f"Concessionária efetive a execução dos serviços para correção {singular_plural_nc}.",
                False, False, None
            )],
        ]

    # ------------------------------------------------------------------
    # Apêndice
    # ------------------------------------------------------------------
    def render_apendices(self, doc, row, ncs_reais, pas_reais, fotos_dir, data_fisc, ano, criar_grade_fotos_fn):
        from utils import adicionar_titulo_secao, formatar_mes_ano
        from docx.shared import Pt

        try:
            mes_ano = formatar_mes_ano(row["Data"]).upper()
        except Exception:
            mes_ano = "JULHO/2026"

        p_ap = adicionar_titulo_secao(doc, f"APÊNDICE ÚNICO - MEMORIAL FOTOGRÁFICO - VISTORIAS REALIZADAS EM {mes_ano}")
        p_ap.paragraph_format.page_break_before = True

        if not ncs_reais.empty:
            criar_grade_fotos_fn(doc, ncs_reais, row.get("Local", ""), fotos_dir, data_fisc, self.key)
        else:
            p_empty = doc.add_paragraph()
            r_empty = p_empty.add_run("Nenhuma foto de não conformidade anexada.")
            r_empty.font.name = "Aptos"
            r_empty.font.size = Pt(11)
