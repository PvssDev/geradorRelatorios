from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reports.base import BaseReport
import pandas as pd

class SocicamReport(BaseReport):
    @property
    def key(self) -> str:
        return "SOCICAM"

    @property
    def display_name(self) -> str:
        return "SOCICAM"

    @property
    def default_contrato(self) -> str:
        return "CT. nº 1.041.080/08"

    @property
    def capa_orgao_concedente(self) -> str:
        return "GOVERNO DO ESTADO DE PERNAMBUCO"

    @property
    def capa_secretaria(self) -> str:
        return "SECRETARIA DE MOBILIDADE E INFRAESTRUTURA - SEMOBI"

    @property
    def capa_titulo(self) -> str:
        return "RELATÓRIO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL"

    @property
    def capa_ctr_number_template(self) -> str:
        return "RELATÓRIO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL CTR Nº 03/{ano}"

    @property
    def capa_prestador_label(self) -> str:
        return "PRESTADOR DE SERVIÇO: SOCICAM - ADMINISTRAÇÃO, PROJETOS E REPRESENTAÇÕES LTDA"

    @property
    def capa_contrato_label(self) -> str:
        return "CONTRATO DE CONCESSÃO DE SERVIÇO PÚBLICO"

    @property
    def capa_regulador_label(self) -> str:
        return "AGÊNCIA DE REGULAÇÃO DOS SERVIÇOS PÚBLICOS DELEGADOS DO ESTADO DE PERNAMBUCO - ARPE"

    def get_capa_titulos(self, row, ano) -> list:
        local_val = str(row.get("Local", "TIP")).upper()
        if "TERMINAL" not in local_val and "FISCALIZAÇÃO" not in local_val:
            local_val = f"TERMINAL RODOVIÁRIO DE PASSAGEIROS DO RECIFE ({local_val})" if "RECIFE" in local_val or "TIP" in local_val else f"TERMINAL RODOVIÁRIO DE PASSAGEIROS ({local_val})"
        return [
            f"FISCALIZAÇÃO NO {local_val}",
            "PRESTADOR DE SERVIÇO: SOCICAM - ADMINISTRAÇÃO, PROJETOS E REPRESENTAÇÕES LTDA"
        ]

    @property
    def sumario_before_abreviaturas(self) -> bool:
        return True

    def get_abbreviations(self) -> list:
        return [
            ("ABNT", "Associação Brasileira de Normas Técnicas"),
            ("ARPE", "Agência de Regulação de Pernambuco"),
            ("EPTI", "Empresa Pernambucana de Transporte Coletivo Intermunicipal"),
            ("NC", "Não Conformidade"),
            ("PCD", "Pessoa com Deficiência"),
            ("TIP", "Terminal Rodoviário de Passageiros do Recife")
        ]

    def get_sumario_linhas(self, row) -> list:
        local_val = str(row.get("Local", "TIP")).upper()
        local_sigla = "TIP"
        if "TIP" in local_val:
            local_sigla = "TIP"
        elif "(" in local_val:
            local_sigla = local_val.split("(")[0].strip()
        else:
            local_sigla = local_val
        return [
            "1.\tINTRODUÇÃO\t4",
            "2.\tOBJETIVO\t4",
            "3.\tMETODOLOGIA\t5",
            "4.\tFISCALIZAÇÃO\t5",
            "5.\tDETERMINAÇÕES GERAIS\t6",
            "6.\tRECOMENDAÇÕES\t6",
            "7.\tCONCLUSÕES\t7",
            f"\tAPÊNDICE A - REGISTROS FOTOGRÁFICOS DAS NÃO CONFORMIDADES APONTADAS PARA O {local_sigla}\t7"
        ]

    def get_intro_paragraphs(self, row, ano, data_extenso) -> list:
        local_val = str(row.get("Local", "Terminal Rodoviário de Passageiros do Recife (TIP)"))
        text = (
            f"A Coordenadoria de Transportes e Rodovias da Arpe realiza vistorias no {local_val} com o objetivo de "
            f"verificar as condições operacionais, de conservação, de manutenção e de segurança e da qualidade do "
            f"serviço prestado nos referidos terminais, conforme Contrato de Concessão de Serviço Público No 1.041.080/08, "
            f"firmado entre o Governo do Estado, atualmente representado pela Empresa Pernambucana de Transportes "
            f"Intermunicipal (EPTI) e a SOCICAM - Administração, Projetos e Representações Ltda (SOCICAM) visando a "
            f"operação, manutenção e administração de terminais rodoviários no Estado de Pernambuco, com execução de obras "
            f"de reforma e construção, incluindo, ainda, a cessão de uso de espaços para a exploração comercial através de "
            f"locação e publicidade."
        )
        if local_val in text:
            parts = text.split(local_val, 1)
            return [
                [
                    (parts[0], False, False, None),
                    (local_val, True, False, None),
                    (parts[1], False, False, None)
                ]
            ]
        return [[(text, False, False, None)]]

    def get_objective_paragraphs(self, row) -> list:
        local_val = str(row.get("Local", "Terminal Rodoviário de Passageiros do Recife (TIP)"))
        if "RECIFE" in local_val.upper() or "TIP" in local_val.upper():
            nome_negrito = "Terminal Rodoviário do Recife"
        else:
            nome_negrito = local_val.split("(")[0].strip()
        text = (
            f"A fiscalização direta e periódica dos Terminais Rodoviários de Passageiros concedidos à SOCICAM, tem por "
            f"objetivo verificar as condições de conservação, limpeza e higiene das áreas de embarque e desembarque, dos "
            f"sanitários, as condições do pavimento das vias de circulação interna, a infraestrutura oferecida, a segurança "
            f"e o atendimento ao usuário, bem como toda estrutura para funcionamento desses terminais. Dessa forma a ação de "
            f"fiscalização no {nome_negrito}, realizada pela ARPE verificou o grau de conformidade dessas instalações com o "
            f"Contrato de Concessão, bem como com a legislação e normas vigentes de modo a determinar e/ou recomendar "
            f"medidas corretivas, com foco na qualidade dos serviços prestados."
        )
        if nome_negrito in text:
            parts = text.split(nome_negrito, 1)
            return [
                [
                    (parts[0], False, False, None),
                    (nome_negrito, True, False, None),
                    (parts[1], False, False, None)
                ]
            ]
        return [[(text, False, False, None)]]

    def get_general_info_rows(self, row, responsaveis_formatted, periodo_val) -> list:
        return [
            ("3.1 DO TITULAR", "", True, False),
            ("Titular:", "Empresa Pernambucana de Transportes Intermunicipal (EPTI)", False, False),
            ("Endereço:", "Av. Caxangá, 2.200 Cordeiro Recife/PE CEP: 50.711-000", False, False),
            ("Responsável:", "THIAGO EDGLES SOBRAL DE SOUZA", False, True),
            
            ("3.2 DO REGULADO", "", True, False),
            ("Regulado:", "SOCICAM - Administração, Projetos e Representações Ltda", False, False),
            ("Responsável:", "THIAGO DUARTE PIMENTEL", False, True),
            ("Endereço:", "Avenida Prefeito Antônio Pereira, S/N Várzea Recife/PE CEP: 50.950-030", False, False),
            ("Representantes para acompanhar:", "Recife (TIP): Monalisa Pereira", False, False),
            
            ("3.3 DO REGULADOR", "", True, False),
            ("Regulador:", "Agência de Regulação de Pernambuco (Arpe)", False, False),
            ("Diretor Presidente:", "CARLOS PORTO FILHO", False, True),
            ("Endereço:", "Av. Cons. Rosa e Silva, 975, Aflitos, Recife/PE, CEP: 52.050-020", False, False),
            ("Estacionamento:", "Rua do Futuro, 150, Aflitos, Recife/PE", False, False),
            ("Responsáveis pela fiscalização:", responsaveis_formatted, False, False),
            ("Período da Fiscalização:", periodo_val, False, False),
            ("Tipo de Fiscalização:", "Direta e periódica.", False, False)
        ]

    @property
    def general_info_col_widths(self) -> list:
        return [Inches(3.0), Inches(4.57)]

    @property
    def general_info_headers_indices(self) -> list:
        return [0, 4, 9]

    def get_methodology_paragraphs(self, data_extenso) -> list:
        return [
            [("A fiscalização direta e periódica realizada pela Coordenadoria de Transportes e Rodovias da Arpe está submetida a uma metodologia organizada em três etapas: Preparação e Planejamento, Execução da Fiscalização e Monitoramento e Avaliação.", False, False, None)],
            [("Preparação e Planejamento", True, False, None),
             (" - compreende a organização e estruturação das atividades preliminares à execução da fiscalização, destacando-se a elaboração e o envio de avisos de fiscalização à Concessionária e demais atividades de suporte à fiscalização, bem como a análise de fiscalizações anteriores com a identificação de eventuais Não Conformidades pendentes.", False, False, None)],
            [("Execução da Fiscalização", True, False, None),
             (" - a execução da fiscalização é pautada por um arcabouço de normas e diretrizes, possibilitando que todas as etapas sejam desenvolvidas de maneira eficiente e em conformidade aos padrões estabelecidos, destacando-se:", False, False, None)]
        ]

    def get_references_bullets(self) -> list:
        return [
            ("Lei nº 13.254, de 21 de junho de 2007, alterada pela Lei nº 15.200, de 17 de dezembro de 2013, e regulamentada pelo Decreto nº 40.559, de 31 de março de 2014.", False, True),
            ("", False, False),
            ("Resoluções Arpe nº 46, de 07 de abril de 2008 (Antiga nº 06/2008), alterada pela Resolução ARPE nº 53, de 26 de janeiro de 2009 (Antiga 003/2009); e nº 083, de 30 de julho de 2013.", False, True),
            ("", False, False),
            ("Contrato de Concessão de Serviço Público Nº 1.041.080/08, de 19 de setembro de 2008 e aditivos, em especial, o Segundo Termo Aditivo ao Contrato de Concessão, de 29 de setembro de 2017.", False, True),
            ("", False, False),
            ("Normas Técnicas da ABNT.", False, True)
        ]

    @property
    def references_left_indent_pt(self) -> float:
        return 53.4

    def get_post_methodology_paragraphs(self, total_ncs) -> list:
        return [
            [("Monitoramento e Avaliação", True, False, None),
             (" - Esta etapa é fundamental para garantir a eficácia das ações corretivas a serem executadas pela Concessionária para a melhoria contínua dos serviços prestados. Os principais instrumentos do Monitoramento e Avaliação são: Termo de Notificação e respectivo Relatório de Fiscalização, Plano de Ação da Concessionária e Relatórios de Monitoramento e Avaliação Final.", False, False, None)]
        ]

    def get_levels_data(self) -> dict:
        return None

    def get_post_metodologia_extra_paragraphs(self, row, data_extenso, total_achados) -> tuple:
        return (None, [])

    def render_quadros(self, doc, row, nc_df, criar_tabela_quadros_fn) -> None:
        from utils import formatar_data_extenso, adicionar_titulo_secao
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        data_extenso = formatar_data_extenso(row["Data"])
        
        # 1. Título da Seção
        adicionar_titulo_secao(doc, self.quadros_section_title)
        doc.add_paragraph() # Pula linha abaixo do título
        
        # 2. Parágrafos introdutórios (todos em negrito para SOCICAM)
        intro_paragraphs = self.get_quadro_intro_paragraphs(row, data_extenso, "")
        for paragraph_runs in intro_paragraphs:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            for text, bold, italic, _ in paragraph_runs:
                run = p.add_run(text)
                run.font.name = 'Aptos'
                run.font.size = Pt(11)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                    
        doc.add_paragraph() # Parágrafo vazio
        
        # Obter dados de NC
        id_fisc = row["ID da Fiscalização"]
        current_ncs = nc_df[nc_df["ID da Fiscalização"] == id_fisc] if nc_df is not None else pd.DataFrame()
        ncs_reais = pd.DataFrame()
        if not current_ncs.empty and "Não Conformidade" in current_ncs.columns:
            ncs_reais = current_ncs[current_ncs["Não Conformidade"].fillna("").astype(str).str.strip() != ""].copy()
            
        local_val = str(row.get("Local", "Terminal Rodoviário de Passageiros do Recife (TIP)"))
        
        # 3. Quadro 1 title
        p7 = doc.add_paragraph()
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p7.paragraph_format.space_after = Pt(6)
        r7_1 = p7.add_run("Quadro 1")
        r7_1.bold = True
        r7_1.font.name = 'Aptos'
        r7_1.font.size = Pt(11)
        r7_2 = p7.add_run(f" – Não Conformidades do {local_val}")
        r7_2.font.name = 'Aptos'
        r7_2.font.size = Pt(11)
        
        # 4. Render Table
        criar_tabela_quadros_fn(doc, ncs_reais, is_pa=False, report_config=self)
        doc.add_paragraph()

    @property
    def quadros_section_title(self) -> str:
        return "4. FISCALIZAÇÃO"

    def get_quadro_intro_paragraphs(self, row, data_extenso, responsaveis_formatted) -> list:
        responsaveis_list = [r.strip() for r in str(row["Pessoal Responsável"]).split(",") if r.strip()]
        from database.manager import carregar_responsaveis
        db_resp = carregar_responsaveis()
        
        team_runs = []
        for idx, nome in enumerate(responsaveis_list):
            match = next((d for d in db_resp if d["nome"].strip().lower() == nome.lower()), None)
            matricula = match['matricula'] if match else "xxxxxxx/xx"
            nome_real = match['nome'] if match else nome
            
            if idx > 0:
                if idx == len(responsaveis_list) - 1:
                    team_runs.append((" e ", False, False, None))
                else:
                    team_runs.append((", ", False, False, None))
            team_runs.append((nome_real, True, False, None))
            team_runs.append((f" (matrícula nº {matricula})", False, False, None))
            
        local_val = str(row.get("Local", "Terminal Rodoviário de Passageiros do Recife (TIP)"))
        
        p1 = [
            ("As ações de fiscalização foram realizadas, em ", False, False, None),
            (f"{data_extenso}, ", False, False, None),
            ("pela equipe formada pelos Especialista em Regulação ", False, False, None)
        ] + team_runs + [
            (f", no {local_val}.", False, False, None)
        ]
        
        p2 = [
            ("As Não Conformidades constatadas estão relacionadas ao ", False, False, None),
            ("Programa de Manutenção dos Terminais Rodoviários", True, False, None),
            (", Anexo V do Contrato de Concessão, conforme descritas no ", False, False, None),
            ("Quadro 1", True, False, None),
            (", a seguir, com indicação dos respectivos registros fotográficos no ", False, False, None),
            ("Apêndice A", True, False, None),
            (".", False, False, None)
        ]
        return [p1, p2]

    @property
    def quadro_title_template(self) -> str:
        return "Quadro 1 – Não Conformidades do {local_val}"

    @property
    def nc_table_headers(self) -> list:
        return [
            "IDENTIFICAÇÃO", 
            "DESCRIÇÃO", 
            "REGISTRO\nFOTOGRÁFICO", 
            "FUNDAMENTO DA INFRAÇÃO\n(ANEXO V CONTRATO DE CONCESSÃO)", 
            "DETERMINAÇÃO"
        ]

    @property
    def nc_table_col_widths(self) -> list:
        return [Inches(1.24), Inches(1.62), Inches(1.26), Inches(1.69), Inches(2.25)]

    def format_nc_table_total_row(self, table, row_idx, total_ncs) -> None:
        r_total = table.rows[row_idx]
        r_total.cells[0].merge(r_total.cells[1]).merge(r_total.cells[2]).merge(r_total.cells[3])
        r_total.cells[0].text = "TOTAL"
        r_total.cells[4].text = str(total_ncs)
        
        col_widths = self.nc_table_col_widths
        r_total.cells[0].width = col_widths[0] + col_widths[1] + col_widths[2] + col_widths[3]
        r_total.cells[4].width = col_widths[4]

    @property
    def finalizacao_sections_config(self) -> dict:
        return {
            "5": "determinações",
            "6": "recomendações",
            "7": "conclusões"
        }

    def get_determinations_paragraphs(self, total_ncs) -> list:
        from sections.finalizacao.finalizacao import numero_por_extenso
        f_extenso = {
            1: "uma", 2: "duas", 3: "três", 4: "quatro", 5: "cinco",
            6: "seis", 7: "sete", 8: "oito", 9: "nove", 10: "dez"
        }
        extenso_ncs = f_extenso.get(total_ncs, numero_por_extenso(total_ncs))
        return [
            [("Considerando os dispositivos contratuais pertinentes e visando garantir a qualidade dos serviços prestados, determina-se que a SOCICAM tome as seguintes medidas através de plano de ação:", False, False, None)],
            [("Manutenção e Monitoramento", True, False, None),
             (": adotar medidas para assegurar a manutenção, o monitoramento contínuo e o cumprimento do Programa de Manutenção dos Terminais Rodoviários, constante da proposta da SOCICAM nos subitens 9.1.1 Manutenção Preventiva; 9.1.2 Manutenção Corretiva e 9.1.3 Tabela de Classificação de Níveis de Falha (tabela de tempos máximos para os níveis de atendimento).", False, False, None)],
            [("Medidas imediatas", True, False, None),
             (" para resolutividade das ", False, False, None),
             (f"{total_ncs} ({extenso_ncs}) novas Não Conformidades", True, False, None),
             (" constatadas, nos prazos estabelecidos, conforme disposto no Quadro 1, na coluna denominada Determinações.", False, False, None)]
        ]

    def get_recommendations_paragraphs(self) -> list:
        return [
            [("Considerando as disposições do Contrato de Concessão, em especial, o Anexo III – Regulamento Interno dos Terminais Rodoviários, aprovado pela Resolução Arpe nº 46, de 07 de abril de 2008 (Antiga nº 06/2008), bem como a legislação aplicável, devem ser observadas pela SOCICAM as seguintes recomendações:", False, False, None)],
            [("Garantir condições de segurança, higiene, acessibilidade e conforto aos usuários dos Terminais Rodoviários, sejam passageiros, público em geral, comerciantes neles estabelecidos, empresas de transportes e de seus empregados.", False, False, None)],
            [("Exigir a utilização de EPI adequados, inclusive por funcionários de empresas terceirizadas que prestem serviços nos Terminais Rodoviários.", False, False, None)],
            [("Providenciar a correta manutenção (evitar o vencimento) de extintores de incêndios nos Terminais Rodoviários.", False, False, None)],
            [("Instalar, sempre que necessário, aviso de sinalização de segurança, principalmente em pontos de risco de acidentes.", False, False, None)]
        ]

    def get_conclusions_paragraphs(self, total_ncs, local_val) -> list:
        from sections.finalizacao.finalizacao import numero_por_extenso
        f_extenso = {
            1: "uma", 2: "duas", 3: "três", 4: "quatro", 5: "cinco",
            6: "seis", 7: "sete", 8: "oito", 9: "nove", 10: "dez"
        }
        extenso_ncs = f_extenso.get(total_ncs, numero_por_extenso(total_ncs))
        return [
            [
                ("Tendo em vista as ações de fiscalização realizadas pela ARPE foram constatadas ", False, False, None),
                (f"{extenso_ncs} novas Não Conformidades no {local_val}", True, False, None),
                (", que devem ser solucionadas pela SOCICAM de acordo com as Determinações desta Agência de Regulação (v. Quadro 1).", False, False, None)
            ],
            [("Por fim, solicita-se o encaminhamento deste Processo de Fiscalização para conhecimento e acompanhamento da EPTI, na qualidade de Poder Concedente do Contrato de Concessão e gestora do Sistema de Transporte Coletivo Intermunicipal de Passageiros (STCIP-PE).", False, False, None)]
        ]

    def render_apendices(self, doc, row, ncs_reais, pas_reais, fotos_dir, data_fisc, ano, criar_grade_fotos_fn) -> None:
        from utils import adicionar_titulo_secao
        from docx.shared import Pt
        
        local_sigla = "TIP"
        local_val = str(row.get("Local", "TIP")).upper()
        if "TIP" in local_val:
            local_sigla = "TIP"
        elif "(" in local_val:
            local_sigla = local_val.split("(")[0].strip()
        else:
            local_sigla = local_val
            
        p_ap_a = adicionar_titulo_secao(doc, f"APÊNDICE A - REGISTROS FOTOGRÁFICOS DAS NÃO CONFORMIDADES APONTADAS PARA O {local_sigla}")
        p_ap_a.paragraph_format.page_break_before = True
        
        if not ncs_reais.empty:
            criar_grade_fotos_fn(doc, ncs_reais, row.get("Local", ""), fotos_dir, data_fisc, self.key)
        else:
            p_empty = doc.add_paragraph()
            r_empty = p_empty.add_run("Nenhum registro fotográfico de não conformidade cadastrado.")
            r_empty.font.name = 'Aptos'
            r_empty.font.size = Pt(11)

    @property
    def analyst_title(self) -> str:
        return "Especialista em Regulação"

    def get_process_sei_texts(self, ano) -> list:
        return [
            f"RELATÓRIO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL PROC ADM Nº 04/{ano} - CTR",
            f"SEI Nº 0030200023.002186/2026-99"
        ]
