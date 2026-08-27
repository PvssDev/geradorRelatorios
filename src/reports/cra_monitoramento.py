# -*- coding: utf-8 -*-
import os
import re
import pandas as pd
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx import Document

from reports.base import BaseMonitoramentoMixin
from reports.cra import CraReport
from monitoramento_utils import (
    extrair_metadados_anterior,
    extrair_linhas_quadro1_anterior,
    criar_tabela_quadro1_monitoramento,
)


class CraMonitoramentoReport(BaseMonitoramentoMixin, CraReport):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.N_prev = "X"
        self.N_curr = "X"
        self.ctr_num = "XX/XXXX"

    @property
    def key(self) -> str:
        return "CRA_MONITORAMENTO"

    @property
    def display_name(self) -> str:
        return "CRA (Monitoramento)"

    @property
    def capa_titulo(self) -> str:
        return "RELATÓRIO DO 1° MONITORAMENTO DO PROCESSO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL"

    @property
    def capa_ctr_number_template(self) -> str:
        return "COORDENADORIA DE TRANSPORTES E RODOVIAS\nRELATÓRIO DO 1° MONITORAMENTO DO PROCESSO\nDE FISCALIZAÇÃO TÉCNICO-OPERACIONAL CTR Nº 07/2025"

    @property
    def sumario_before_abreviaturas(self) -> bool:
        return False

    @property
    def signatures_before_apendices(self) -> bool:
        return True

    def get_abbreviations(self) -> list:
        return []

    def get_process_sei_texts(self, row, ano=None) -> list:
        from utils import extrair_mes_ano_numerico, extrair_ano
        if isinstance(row, (str, int)) and ano is None:
            ano = str(row)
            mes_ano = f"01/{ano}"
        else:
            ano = ano or extrair_ano(row.get("Data", "") if hasattr(row, "get") else row["Data"])
            data_val = row.get("Data", "") if hasattr(row, "get") else (row["Data"] if isinstance(row, dict) or hasattr(row, "__getitem__") else "")
            mes_ano = extrair_mes_ano_numerico(data_val)
        return [
            f"RELATÓRIO DO 1° MONITORAMENTO DO PROCESSO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL CTR Nº {mes_ano}",
            f"SEI Nº 0030200023.009194/{ano}-85"
        ]

    def get_objective_paragraphs(self, row) -> list:
        return [
            [(
                "Este Relatório de Monitoramento tem por objetivo apresentar os resultados das vistorias realizadas pela ARPE "
                "no Complexo Viário e Logístico de SUAPE referentes às Não Conformidades do relatório anterior.",
                False, False, None
            )]
        ]

    # ------------------------------------------------------------------
    # Capa de Monitoramento CRA
    # ------------------------------------------------------------------
    def gerar_capa_monitoramento(self, doc, logo_path, row, documento_anterior):
        style = doc.styles['Normal']
        style.font.name = 'Aptos'
        
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        self.documento_anterior = documento_anterior
        meta = extrair_metadados_anterior(documento_anterior)
        self.N_prev = meta["N_prev"]
        self.N_curr = meta["N_curr"]
        self.ctr_num = meta["ctr_num"] if meta["ctr_num"] != "XX/XXXX" else "07/2025"
        self.processo_sei_prev = meta["processo_sei_prev"] if meta["processo_sei_prev"] != "XXXXXXXX" else "0030200023.009194/2025-85"

        n_curr_str = f"{self.N_curr}º" if self.N_curr != "X" else "1º"

        def add_cover_p(text, bold=True, size_pt=12, space_after=6, align=WD_ALIGN_PARAGRAPH.CENTER):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(space_after)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if text:
                run = p.add_run(text)
                run.font.name = "Aptos"
                run.font.size = Pt(size_pt)
                run.bold = bold
            return p

        # Cabeçalho Superior em 3 linhas
        add_cover_p("", bold=False, size_pt=12, space_after=6)
        add_cover_p("", bold=False, size_pt=12, space_after=6)
        add_cover_p("COORDENADORIA DE TRANSPORTES E RODOVIAS", bold=True, size_pt=12, space_after=4)
        add_cover_p(f"RELATÓRIO DO {n_curr_str} MONITORAMENTO DO PROCESSO", bold=True, size_pt=12, space_after=4)
        add_cover_p(f"DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL CTR Nº {self.ctr_num}", bold=True, size_pt=12, space_after=12)

        # Imagem da Capa (capa_monitoramento.png)
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        capa_mon = os.path.join(base_dir, "assets", "capa_monitoramento.png")
        if not os.path.exists(capa_mon):
            capa_mon = os.path.join(base_dir, "assets", "capa_1.png")
        if not os.path.exists(capa_mon):
            capa_mon = logo_path  # fallback
            
        if os.path.exists(capa_mon):
            doc.add_picture(capa_mon, width=Inches(5.90))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
            doc.paragraphs[-1].paragraph_format.space_before = Pt(12)

        # Títulos Centrais
        add_cover_p("VISTORIA TÉCNICA DAS NÃO CONFORMIDADES REGISTRADAS NO", bold=True, size_pt=11, space_after=4)
        add_cover_p(f"RELATÓRIO DE FISCALIZAÇÃO TÉCNICO-OPERACIONAL CTR {self.ctr_num}", bold=True, size_pt=11, space_after=6)
        add_cover_p("", bold=False, size_pt=11, space_after=6)
        add_cover_p("CONTRATO DE CONCESSÃO PARA A EXPLORAÇÃO DO COMPLEXO", bold=True, size_pt=11, space_after=4)
        add_cover_p("VIÁRIO E LOGÍSTICO DE SUAPE – EXPRESS WAY - CT Nº 043/2011", bold=True, size_pt=11, space_after=6)
        add_cover_p("", bold=False, size_pt=11, space_after=6)
        add_cover_p(f"PROCESSO SEI Nº {self.processo_sei_prev}", bold=True, size_pt=11, space_after=6)
        add_cover_p("", bold=False, size_pt=11, space_after=6)
        add_cover_p("Recife, data da assinatura eletrônica", bold=False, size_pt=11, space_after=0)

    # ------------------------------------------------------------------
    # Sumário Dinâmico
    # ------------------------------------------------------------------
    def get_sumario_linhas(self, row, nc_df=None) -> list:
        from utils import formatar_mes_ano
        try:
            mes_ano = formatar_mes_ano(row["Data"]).upper()
        except Exception:
            mes_ano = "ABRIL/2026"

        linhas = [
            "1.\tINTRODUÇÃO\t4",
            "2.\tOBJETIVO\t4",
            f"3.\tRESULTADOS DAS VISTORIAS DAS NÃO CONFORMIDADES – {mes_ano}\t5",
        ]

        id_fisc = str(row.get("ID da Fiscalização", "")).strip()
        current_ncs = pd.DataFrame()
        if nc_df is not None and not nc_df.empty:
            if "ID da Fiscalização" in nc_df.columns:
                mask_id = nc_df["ID da Fiscalização"].astype(str).str.strip() == id_fisc
                current_ncs = nc_df[mask_id]
            if current_ncs.empty:
                current_ncs = nc_df.copy()

        pistas_unicas = []
        if not current_ncs.empty and "Pista" in current_ncs.columns:
            for p in current_ncs["Pista"].tolist():
                p_str = str(p).strip() if not pd.isna(p) else ""
                if not p_str:
                    p_str = "Única"
                if p_str not in pistas_unicas:
                    pistas_unicas.append(p_str)

        if not pistas_unicas:
            pistas_unicas = ["Sentido Sul", "Sentido Norte"]

        for idx, pista_name in enumerate(pistas_unicas, 1):
            p_up = pista_name.upper()
            if "SUL" in p_up and "SENTIDO" not in p_up:
                pista_display = "PISTA SENTIDO SUL"
            elif "NORTE" in p_up and "SENTIDO" not in p_up:
                pista_display = "PISTA SENTIDO NORTE"
            elif p_up in ["ÚNICA", "UNICA"]:
                pista_display = "PISTA ÚNICA"
            elif not p_up.startswith("PISTA"):
                pista_display = f"PISTA {p_up}"
            else:
                pista_display = p_up
            linhas.append(f"3.{idx}\tRODOVIA ESTADUAL PE - 009, {pista_display}\t5")

        linhas.extend([
            f"4.\tRESUMO DA SITUAÇÃO DAS NÃO CONFORMIDADES MONITORADAS ({mes_ano})\t8",
            "5.\tCONCLUSÕES E RECOMENDAÇÕES\t9",
            f"\tAPÊNDICE - MEMORIAL FOTOGRÁFICO - VISTORIAS REALIZADAS EM {mes_ano}\t10",
        ])

        return linhas

    # ------------------------------------------------------------------
    # Seções de Monitoramento 1 a 4
    # ------------------------------------------------------------------
    def gerar_secoes_monitoramento(self, doc, row, nc_df, total_achados, documento_anterior=None):
        from utils import adicionar_titulo_secao, formatar_data_extenso, formatar_mes_ano

        meta = extrair_metadados_anterior(documento_anterior)
        N_curr = meta["N_curr"] if meta["N_curr"] != "X" else getattr(self, "N_curr", "1")
        ctr_num = meta["ctr_num"] if meta["ctr_num"] != "XX/XXXX" else getattr(self, "ctr_num", "07/2025")

        n_curr_str = f"{N_curr}º" if N_curr != "X" else "1º"

        data_extenso = formatar_data_extenso(row["Data"])
        try:
            from pandas import to_datetime
            data_vistoria = to_datetime(row["Data"], dayfirst=True).strftime("%d/%m/%Y")
        except Exception:
            data_vistoria = "14 e 15/04/2026"

        id_fisc = str(row.get("ID da Fiscalização", "")).strip()
        mes_ano_cap = formatar_mes_ano(row["Data"]).upper()

        if nc_df is not None and not nc_df.empty:
            if "ID da Fiscalização" in nc_df.columns:
                mask_id = nc_df["ID da Fiscalização"].astype(str).str.strip() == id_fisc
                current_ncs = nc_df[mask_id].copy()
            else:
                current_ncs = nc_df.copy()
            
            if current_ncs.empty:
                current_ncs = nc_df.copy()

            cols_check = [c for c in ["Não Conformidade", "Identificação", "Observações", "Determinação"] if c in current_ncs.columns]
            if cols_check:
                mask_nc = current_ncs[cols_check].fillna("").astype(str).apply(lambda r_c: any(v.strip() != "" for v in r_c), axis=1)
                ncs_reais = current_ncs[mask_nc].copy()
            else:
                ncs_reais = current_ncs.copy()
        else:
            ncs_reais = pd.DataFrame()

        def add_p(text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(text)
            run.font.name = "Aptos"
            run.font.size = Pt(11)
            run.bold = bold
            return p

        def add_label_text(label, text):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            run_l = p.add_run(label)
            run_l.bold = True
            run_l.font.name = "Aptos"
            run_l.font.size = Pt(11)
            run_t = p.add_run(text)
            run_t.font.name = "Aptos"
            run_t.font.size = Pt(11)
            return p

        # 1. INTRODUÇÃO
        adicionar_titulo_secao(doc, "1. INTRODUÇÃO")
        doc.add_paragraph()
        add_p(
            f"Registra-se, preliminarmente, que o Relatório de Fiscalização Técnico-Operacional ARPE/CTR {ctr_num} "
            f"registrou as Não Conformidades identificadas e foi devidamente encaminhado à Concessionária Rota do "
            f"Atlântico (CRA) e aos órgãos concedentes para as providências cabíveis."
        )
        add_p(
            f"A Concessionária Rota do Atlântico (CRA) apresentou o cronograma de obras atualizado por trecho, com detalhamento "
            f"dos prazos de execução das ações de conservação especial do pavimento e sinalização viária."
        )
        add_p(
            f"Neste contexto, este {n_curr_str} Relatório de Monitoramento descreve a evolução das Não Conformidades apontadas, "
            f"com base nas vistorias técnicas realizadas pela equipe da ARPE em {data_extenso} no Sistema Viário Express Way."
        )

        # 2. OBJETIVO
        adicionar_titulo_secao(doc, "2. OBJETIVO")
        doc.add_paragraph()
        add_p(
            f"Este {n_curr_str} Relatório de Monitoramento tem como objetivo apresentar os resultados das vistorias realizadas pela ARPE, "
            f"em {data_vistoria}, no Sistema Viário Express Way referentes às Não Conformidades constantes do Relatório de Fiscalização "
            f"Técnico-Operacional ARPE/CTR {ctr_num}."
        )

        # 3. RESULTADOS DAS VISTORIAS
        adicionar_titulo_secao(doc, f"3. RESULTADOS DAS VISTORIAS DAS NÃO CONFORMIDADES – {mes_ano_cap}")
        doc.add_paragraph()
        add_p(
            f"Apresenta-se a seguir a análise das constatações resultantes das vistorias técnicas das Não Conformidades associadas "
            f"ao Relatório de Fiscalização Técnico-Operacional ARPE/CTR {ctr_num}, considerando o Cronograma das Obras por Trecho "
            f"e o posicionamento apresentado pela Concessionária Rota do Atlântico (CRA)."
        )

        if ncs_reais.empty:
            add_p("Nenhuma Não Conformidade registrada para este monitoramento.")
        else:
            pistas_unicas = []
            for p in ncs_reais["Pista"].tolist():
                p_str = str(p).strip() if not pd.isna(p) else ""
                if not p_str:
                    p_str = "Única"
                if p_str not in pistas_unicas:
                    pistas_unicas.append(p_str)

            for sub_idx, pista_val in enumerate(pistas_unicas, 1):
                p_upper = pista_val.upper()
                if "SUL" in p_upper and "SENTIDO" not in p_upper:
                    pista_display = "PISTA SENTIDO SUL"
                elif "NORTE" in p_upper and "SENTIDO" not in p_upper:
                    pista_display = "PISTA SENTIDO NORTE"
                elif p_upper in ["ÚNICA", "UNICA"]:
                    pista_display = "PISTA ÚNICA"
                elif not p_upper.startswith("PISTA"):
                    pista_display = f"PISTA {p_upper}"
                else:
                    pista_display = p_upper

                sub_title = f"3.{sub_idx} RODOVIA ESTADUAL PE - 009, {pista_display}"
                adicionar_titulo_secao(doc, sub_title)
                doc.add_paragraph()

                mask = ncs_reais["Pista"].apply(lambda x: (str(x).strip() if not pd.isna(x) else "") == (pista_val if pista_val != "Única" else ""))
                df_pista = ncs_reais[mask]

                for row_idx, (_, nc_row) in enumerate(df_pista.reset_index(drop=True).iterrows()):
                    ident = str(nc_row.get("Identificação", "")).strip()
                    
                    # Logic to fallback to ident from previous doc just like CRC
                    if not ident or ident.lower().startswith("foto"):
                        if row_idx < len(ncs_from_prev):
                            ident = ncs_from_prev[row_idx]["id_nc"]

                    nc_desc = str(nc_row.get("Não Conformidade", "")).strip()
                    observacoes = str(nc_row.get("Observações", nc_row.get("Legenda da Foto", ""))).strip()
                    determinacao = str(nc_row.get("Determinação", "")).strip()
                    analise_arpe = str(nc_row.get("Análise ARPE", "")).strip()
                    situacao = str(nc_row.get("Situação", "Pendente")).strip()

                    if ident and not ident.lower().startswith("foto"):
                        add_p(ident, bold=True)
                    
                    add_label_text("NÃO CONFORMIDADE: ", nc_desc if nc_desc else "A ser preenchido.")
                    add_label_text("POSICIONAMENTO CRA: ", determinacao if determinacao else "A ser preenchido.")
                    add_label_text("CONSTATAÇÃO: ", observacoes if observacoes else "A ser preenchido.")

                    if analise_arpe:
                        texto_analise = analise_arpe
                    elif situacao:
                        texto_analise = f"Não Conformidade {situacao.title()}."
                    else:
                        texto_analise = "A ser preenchido."

                    add_label_text("ANÁLISE ARPE: ", texto_analise)
                    doc.add_paragraph()

        # 4. RESUMO DA SITUAÇÃO
        adicionar_titulo_secao(doc, f"4. RESUMO DA SITUAÇÃO DAS NÃO CONFORMIDADES MONITORADAS ({mes_ano_cap})")
        doc.add_paragraph()
        add_p(
            f"Apresenta-se no Quadro 1, a seguir, um resumo da situação das Não Conformidades registradas no Relatório de "
            f"Fiscalização Técnico-Operacional ARPE/CTR {ctr_num}, de acordo com as vistorias realizadas em {data_vistoria}, "
            f"em continuidade às diligências necessárias ao monitoramento das respectivas soluções."
        )

    # ------------------------------------------------------------------
    # Renderização dos Quadros (Quadro 1 de 3 colunas)
    # ------------------------------------------------------------------
    @property
    def quadros_section_title(self) -> str:
        return ""

    def render_quadros(self, doc, row, nc_df, criar_tabela_quadros_fn):
        from utils import formatar_mes_ano
        id_fisc = str(row.get("ID da Fiscalização", "")).strip()
        ctr_num = getattr(self, "ctr_num", "07/2025")
        try:
            mes_ano = formatar_mes_ano(row["Data"]).title()
        except Exception:
            mes_ano = "Abril/2026"

        current_ncs = pd.DataFrame()
        if nc_df is not None and not nc_df.empty:
            if "ID da Fiscalização" in nc_df.columns:
                mask_id = nc_df["ID da Fiscalização"].astype(str).str.strip() == id_fisc
                current_ncs = nc_df[mask_id].copy()
            if current_ncs.empty:
                current_ncs = nc_df.copy()

        ncs_reais = current_ncs.copy() if not current_ncs.empty else pd.DataFrame()

        # Título do Quadro 1
        p_q = doc.add_paragraph()
        p_q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_q.paragraph_format.space_after = Pt(6)
        r_q1 = p_q.add_run("QUADRO 1")
        r_q1.bold = True
        r_q1.font.name = "Aptos"
        r_q1.font.size = Pt(11)
        r_q2 = p_q.add_run(f" – Resumo da Situação das Não Conformidades – {mes_ano}")
        r_q2.bold = True
        r_q2.font.name = "Aptos"
        r_q2.font.size = Pt(11)

        doc_anterior_path = getattr(self, "documento_anterior", None)
        ncs_from_prev = extrair_linhas_quadro1_anterior(doc_anterior_path)

        headers = ["REFERÊNCIA", f"CONSTATAÇÃO / NÃO CONFORMIDADE RELATÓRIO ARPE/CTR {ctr_num}", "SITUAÇÃO"]
        col_widths = [Inches(1.80), Inches(4.20), Inches(1.57)]
        criar_tabela_quadro1_monitoramento(doc, headers, ncs_from_prev, ncs_reais, col_widths)

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Seção 5 — Conclusões e Recomendações
    # ------------------------------------------------------------------
    @property
    def finalizacao_sections_config(self) -> dict:
        return {"5": "conclusoes_monitoramento"}

    def get_conclusions_monitoramento_paragraphs(self, total_ncs, data_extenso) -> list:
        N_curr = getattr(self, "N_curr", "1")
        n_curr_str = f"{N_curr}º" if N_curr != "X" else "1º"
        ctr_num = getattr(self, "ctr_num", "07/2025")

        return [
            [(
                f"O resultado das vistorias realizadas nos dias {data_extenso} e das respectivas análises apresentadas neste "
                f"{n_curr_str} Relatório de Monitoramento demonstram a evolução do tratamento das Não Conformidades identificadas no "
                f"Relatório de Fiscalização Técnico-Operacional ARPE/CTR {ctr_num}. Recomenda-se que a Concessionária mantenha atualizado "
                f"o cronograma das obras e soluções pendentes visando permitir à ARPE uma programação efetiva do monitoramento.",
                False, False, None
            )],
            [(
                f"Por fim, solicita-se dar conhecimento deste {n_curr_str} Relatório de Monitoramento do Relatório de Fiscalização "
                f"Técnico-Operacional ARPE/CTR {ctr_num} à Diretoria de Desenvolvimento e Gestão Portuária de SUAPE para as devidas "
                f"providências, em atendimento ao Convênio ARPE/Suape nº 003/2021.",
                False, False, None
            )],
        ]

    # ------------------------------------------------------------------
    # Apêndice Fotográfico de Monitoramento
    # ------------------------------------------------------------------
    def render_apendices(self, doc, row, ncs_reais, pas_reais, fotos_dir, data_fisc, ano, criar_grade_fotos_fn):
        from utils import adicionar_titulo_secao, formatar_mes_ano
        from docx.shared import Pt

        try:
            mes_ano = formatar_mes_ano(row["Data"]).upper()
        except Exception:
            mes_ano = "ABRIL/2026"

        p_ap = adicionar_titulo_secao(doc, f"APÊNDICE - MEMORIAL FOTOGRÁFICO - VISTORIAS REALIZADAS EM {data_fisc}")
        p_ap.paragraph_format.page_break_before = True

        p_desc = doc.add_paragraph()
        p_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_desc.paragraph_format.space_after = Pt(12)
        ctr_num = getattr(self, "ctr_num", "07/2025")
        r_desc = p_desc.add_run(
            f"Apresentam-se, a seguir, as evidências fotográficas coletadas no monitoramento das Não Conformidades apontadas "
            f"no Relatório de Fiscalização Técnico-Operacional ARPE/CTR {ctr_num}."
        )
        r_desc.font.name = "Aptos"
        r_desc.font.size = Pt(11)

        if not ncs_reais.empty:
            criar_grade_fotos_fn(doc, ncs_reais, row.get("Local", ""), fotos_dir, data_fisc, self.key)
        else:
            p_empty = doc.add_paragraph()
            r_empty = p_empty.add_run("Nenhuma foto de não conformidade anexada.")
            r_empty.font.name = "Aptos"
            r_empty.font.size = Pt(11)
