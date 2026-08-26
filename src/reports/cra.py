from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reports.base import BaseReport
import pandas as pd

class CraReport(BaseReport):
    @property
    def key(self) -> str:
        return "CRA"

    @property
    def display_name(self) -> str:
        return "CRA"

    @property
    def default_contrato(self) -> str:
        return "CT. nº 043/2011"

    @property
    def capa_orgao_concedente(self) -> str:
        return "GOVERNO DO ESTADO DE PERNAMBUCO"

    @property
    def capa_secretaria(self) -> str:
        return "SECRETARIA DE MOBILIDADE E INFRAESTRUTURA - SEMOBI"

    @property
    def capa_titulo(self) -> str:
        return "RELATÓRIO DE FISCALIZAÇÃO PROCESSO ADMINISTRATIVO"

    @property
    def capa_ctr_number_template(self) -> str:
        return "RELATÓRIO DE FISCALIZAÇÃO PROCESSO ADMINISTRATIVO CTR Nº {mes_ano}"

    @property
    def capa_prestador_label(self) -> str:
        return "PRESTADOR DE SERVIÇO: CONCESSIONÁRIA ROTA DO ATLÂNTICO (CRA)"

    @property
    def capa_contrato_label(self) -> str:
        return "CONTRATO DE CONCESSÃO"

    @property
    def capa_regulador_label(self) -> str:
        return "AGÊNCIA DE REGULAÇÃO DOS SERVIÇOS PÚBLICOS DELEGADOS DO ESTADO DE PERNAMBUCO - ARPE"

    def get_capa_titulos(self, row, ano) -> list:
        return [
            "FISCALIZAÇÃO DO COMPLEXO VIÁRIO E LOGÍSTICO DE SUAPE – EXPRESSWAY",
            "PRESTADOR DE SERVIÇO: CONCESSIONÁRIA ROTA DO ATLÂNTICO (CRA)",
            "CONTRATO DE CONCESSÃO CT. Nº 043/2011"
        ]

    @property
    def sumario_before_abreviaturas(self) -> bool:
        return False

    def get_abbreviations(self) -> list:
        return [
            ("CRA", "Concessionária Rota do Atlântico"),
            ("ECR", "ECR Engenharia Ltda"),
            ("FD", "Faixa Direita"),
            ("FE", "Faixa Esquerda"),
            ("IGG", "Índice de Gravidade Global"),
            ("IRI", "Índice Irregularidade Longitudinal"),
            ("NC", "Não Conformidade"),
            ("PDCL", "Programa de Desenvolvimento do Complexo Logístico, Anexo IV do Contrato de Concessão nº 043/2011"),
            ("SUAPE", "Poder Concedente e Regulador do Contrato de Concessão firmado com a CRA"),
            ("TPF", "TPF Engenharia Ltda"),
            ("TDR", "Tronco Distribuidor Rodoviário"),
            ("VI", "Verificador Independente contratado por SUAPE, atualmente o Consórcio formado pelas Empresas TPF e ECR")
        ]

    def get_sumario_linhas(self, row) -> list:
        return [
            "1.\tINTRODUÇÃO\t4",
            "2.\tOBJETIVO\t4",
            "3.\tINFORMAÇÕES GERAIS\t4",
            "4.\tMETODOLOGIA\t5",
            "5.\tFISCALIZAÇÃO\t7",
            "6.\tDETERMINAÇÕES GERAIS\t11",
            "7.\tRECOMENDAÇÕES\t11",
            "8.\tCONCLUSÕES\t11",
            "\tAPÊNDICE ÚNICO  – REGISTROS FOTOGRÁFICOS DAS NÃO CONFORMIDADES\t12"
        ]

    def get_intro_paragraphs(self, row, ano, data_extenso) -> list:
        return [
            [("A Coordenadoria de Transporte e Rodovias da Arpe implementou o cronograma de fiscalização para " + str(ano) + " do Complexo Viário e Logístico de SUAPE, sob a responsabilidade da Concessionária Rota do Atlântico (CRA), visualizando a necessidade de serem reservados dois dias consecutivos, para melhor estruturar suas ações de fiscalização técnico-operacionais da Rodovia.", False, False, None)],
            [("Ainda com a visão de implantar melhorias na contribuição da Agência sobre os trabalhos desenvolvidos na rodovia, foi introduzida uma visita técnica prévia, realizada uma semana antes do período de fiscalização, com o objetivo de elaborar um roteiro que é encaminhado para a CRA e SUAPE contendo os trechos mapeados com possíveis Não Conformidades, tornando os levantamentos fotográficos mais ágeis e a fiscalização efetiva.", False, False, None)],
            [("Posteriormente, as possíveis Não Conformidades levantadas in campo são analisadas de acordo com os critérios elencados no PDCL anexo ao Contrato de Concessão, em conjunto com o último Relatório Anual elaborado pelo Verificador Independente.", False, False, None)],
            [("Destaca-se preliminarmente que as ações de fiscalização registradas neste Relatório foram concentradas na Rodovia PE-009, do Entroncamento BR-101 ao Entroncamento PE-038, que abrange os subtrechos concedidos, especificamente, o Contorno do Cabo, TDR Norte, TDR Sul e a Ligação Rótula Curva do Boi a Nossa Senhora do Ó; e Rodovia Estadual VPE-034.", False, False, None)],
            [
                (f"É importante observar que foram realizadas ações de fiscalização, no dia {data_extenso}, conforme comunicado enviado à SUAPE por meio do Ofício Arpe/DTO nº ", False, False, None),
                ("302 (Doc. SEI 75605952)", False, False, (255, 0, 0)),
                (".", False, False, None)
            ],
            [("Destaca-se que as fiscalizações realizadas pela Arpe são tratadas com caráter educativo, preferencialmente, e contributivo para correção de procedimentos e solução de Não Conformidades evidenciados por defeitos e/ou problemas na infraestrutura disponibilizada e respectivos serviços concedidos pelo Estado.", False, False, None)]
        ]

    def get_objective_paragraphs(self, row) -> list:
        return [
            [("A fiscalização direta e periódica do complexo viário e logístico de Suape – Expressway concedido à CRA, "
            "tem por objetivo verificar as condições de operação, manutenção, segurança viária e níveis de serviço, bem "
            "como identificar não conformidades, subsidiar medidas corretivas e assegurar a observância das disposições "
            "contratuais, regulamentares e normativas aplicáveis, preservando a segurança, a qualidade do serviço, e a "
            "adequada prestação do serviço público aos usuários. Dessa forma a ação de fiscalização da Arpe verifica o grau "
            "de conformidade dessas instalações com o Contrato de Concessão, bem como com a legislação e normas vigentes "
            "de modo a determinar/ou recomendar medidas corretivas, com foco na qualidade dos serviços prestados.", False, False, None)]
        ]

    def get_general_info_rows(self, row, responsaveis_formatted, periodo_val) -> list:
        return [
            ("3.1 DO TITULAR E REGULADOR", "", True, False),
            ("Titular:", "SUAPE – Complexo Industrial Portuário Governador Eraldo Gueiros", False, False),
            ("Endereço:", "Engenho Massangana – Km 10 – Rodovia PE – 60 Ipojuca/PE CEP: 55.590-000", False, False),
            ("Responsável:", "JOSÉ CONSTANTINO DA SILVA FILHO", False, True),
            ("Representantes por acompanhar:", "Viviane Alves Walzertudes", False, False),
            
            ("3.2 DO VERIFICADOR INDEPENDENTE", "", True, False),
            ("Verificador Independente:", "Consórcio das Empresas TPF/ECR", False, False),
            ("Endereço:", "Rua Irene Ramos Gomes de Mattos, Nº 176, Pina, Recife/PE CEP: 51011-530", False, False),
            ("Responsável:", "RICARDO MEDEIROS PEREIRA DE CARVALHO", False, True),
            ("Representantes por acompanhar:", "Sónya Albuquerque; Ricardo Henrique Ferraz de Farias; Lauro Ricardo Torres Galindo e Maynara Milena Silva de Lima", False, False),
            
            ("3.3 DO REGULADO", "", True, False),
            ("Regulado:", "CRA - Concessionária Rota do Atlântico", False, False),
            ("Responsável:", "RAFAELA ELAINE DA COSTA LIMA ARAÚJO", False, True),
            ("Endereço:", "Rodovia PE-009, Km 38,5(TDR Norte, 2074) – Distrito Industrial Suape, Cabo de Santo Agostinho/PE – CEP: 54.590-000", False, False),
            ("Representantes por acompanhar:", "Vanessa Monteiro e OuvidoraXXXCRA", False, False),
            
            ("3.4 DO FISCALIZADOR (CONVÊNIO SUAPE/ARPE Nº 003/2021)", "", True, False),
            ("Regulador:", "Agência de Regulação de Pernambuco (Arpe)", False, False),
            ("Diretor Presidente:", "CARLOS PORTO FILHO", False, True),
            ("Endereço:", "Avenida Conselheiro Rosa e Silva, 975, Aflitos, Recife/PE, CEP: 52.050-020.\nEstacionamento: Rua do Futuro, 150, Aflitos, Recife/PE.", False, False),
            
            ("Responsáveis pela fiscalização:", responsaveis_formatted, False, False),
            ("Período da Fiscalização:", periodo_val, False, False),
            ("Tipo de Fiscalização:", "Direta e periódica.", False, False)
        ]

    @property
    def general_info_col_widths(self) -> list:
        return [Inches(2.3), Inches(5.2)]

    @property
    def general_info_headers_indices(self) -> list:
        return [0, 5, 10, 15]

    def get_methodology_paragraphs(self, data_extenso) -> list:
        ano = "2026"
        if data_extenso and len(data_extenso) > 4:
            ano = data_extenso[-4:]
        return [
            [(f"A Coordenadoria de Transporte e Rodovias da Arpe implementou o cronograma de fiscalização para {ano} do Complexo Viário e Logístico de SUAPE reservando dois dias consecutivos, para melhor estruturar suas ações de fiscalização técnico-operacionais no Complexo Rodoviário, após uma visita técnica prévia, realizada antes do período de fiscalização, com o objetivo de elaborar um roteiro que é encaminhado para a CRA e SUAPE contendo os trechos mapeados com possíveis Não Conformidades, tornando os levantamentos fotográficos mais ágeis e a fiscalização efetiva.", False, False, None)],
            [("Posteriormente, as possíveis Não Conformidades levantadas em campo são analisadas de acordo com os critérios elencados no PDCL (Anexo IV do Contrato de Concessão), em conjunto com o último Relatório Anual elaborado pelo Verificador Independente.", False, False, None)],
            [("Assim, a fiscalização direta e periódica realizada pela Coordenadoria de Transportes e Rodovias da Arpe está submetida a uma metodologia organizada em três etapas: Preparação e Planejamento, Execução da Fiscalização e Monitoramento e Avaliação.", False, False, None)],
            [("Preparação e Planejamento - compreende a organização e estruturação das atividades preliminares à execução da fiscalização, destacando-se a elaboração e o envio de avisos de fiscalização à Concessionária e demais atividades de suporte à fiscalização, bem como a análise de fiscalizações anteriores com a identificação de eventuais Não Conformidades pendentes.", False, False, None)],
            [("Execução da Fiscalização - a execução da fiscalização é pautada por um arcabouço de normas e diretrizes, possibilitando que todas as etapas sejam desenvolvidas de maneira eficiente e em conformidade aos padrões estabelecidos, destacando-se:", False, False, None)]
        ]

    def get_references_bullets(self) -> list:
        return [
            ("Lei Estadual nº. 12.524, de 30 de dezembro de 2003, regulamentada pelo Decreto Estadual nº. 30.200, de 09 de fevereiro de 2007, que altera e consolida as disposições da Lei nº. 12.126, de 12 de dezembro de 2001, que cria a Agência de Regulação dos Serviços Públicos Delegados do Estado de Pernambuco - ARPE, e dá outras providências;", False, True),
            ("Art. 3º Compete à ARPE a regulação de todos os serviços públicos delegados pelo Estado de Pernambuco, ou por ele diretamente prestados, embora sujeitos à delegação, quer de sua competência ou a ele delegados por outros entes federados, em decorrência de norma legal ou regulamentar, disposição convenial ou contratual.", True, False),
            ("§ 1º A atividade reguladora da ARPE deverá ser exercida, em especial, nas seguintes áreas:", True, False),
            ("[...]", True, False),
            ("III - rodovias;", True, False),
            ("[...]", True, False),
            ("Art. 4º Compete ainda à ARPE:", True, False),
            ("[...]", True, False),
            ("X - Fiscalizar diretamente ou mediante convênio com o Estado de Pernambuco, através de seus órgãos ou entidades vinculadas, com sua supervisão, os aspects técnico, econômico, contábil, financeiro, operacional e jurídico dos serviços públicos delegados, valendo-se inclusive, de indicadores e procedimentos amostrais.", True, False),
            ("", False, False),
            ("Contrato de Concessão CT. nº 043/2011, de 18 de julho de 2011, para a delegação da exploração do Complexo Viário e Logístico de SUAPE – EXPRESSWAY, conforme detalhado no Anexo IV do Edital (PDCL) e regidos pela Constituição Federal; pela Lei Federal Nº 8.987/95; Lei Federal Nº 9.074/95; Lei Federal 8.666/93 e Lei Estadual Nº 14.233/2010.", False, True),
            ("", False, False),
            ("Convênio de Cooperação Técnica n° 003/2021 de 22 de setembro 2021, firmado entre o Complexo Industrial Portuário Governador Eraldo Gueiros – SUAPE e a Agência de Regulação de Pernambuco – ARPE, e Renovação do Termo Aditivo que prorroga o prazo de vigência e execução contratual até 22 de setembro de 2026.", False, True),
            ("", False, False),
            ("Norma DNIT 005/2003 – que define os termos técnicos empregados in defeitos que ocorrem nos pavimentos flexíveis e semirrígidos e serve para padronizar a linguagem adotada na elaboração das normas, manuais, projetos e textos relativos aos pavimentos flexíveis e semirrígidos.", False, True),
            ("", False, False),
            ("Relatórios elaborados pelo Verificador Independente.", False, True)
        ]

    @property
    def references_left_indent_pt(self) -> float:
        return 36.0

    def get_post_methodology_paragraphs(self, total_ncs) -> list:
        return [
            [("Registra-se que as Não conformidades (NC) são codificadas de acordo com os seguintes níveis de informação, separados por “.” (ponto):", False, False, None)]
        ]

    def get_levels_data(self) -> dict:
        return {
            "niveis": [
                "Nível 1 - três dígitos, caracterizando a concessionária: CRA ou CRC.",
                "Nível 2 - composto por cinco dígitos: os dois primeiros representam as subdivisões utilizadas na rodovia concedida, por exemplo, Trecho (TR); Subtrecho (ST); Segmento Homogêneo (SH) e os três últimos dígitos identificam a subdivisão utilizada em cada contrato.",
                "Nível 3 – composto por nove dígitos os quatro primeiros e os quatro últimos delimitam a localização das NC em escala de 0,01Km e o dígito intermediário informa se a NC é pontual \"+\" ou distribuída.",
                "Nível 4 – informa o ano da constatação da NC com quatro dígitos, antecedido de “/”.",
                "Nível 5 - contendo o sequencial numérico da NC no ano, expresso em três dígitos."
            ],
            "ex_str": "CRA.ST601.2794-2834/2025.013",
            "ex_desc": "Indica que esta Não Conformidade foi apontada para a CRA, no subtrecho 6.01, distribuída do km 27,94 ao km 28,34, sendo a 13ª NC registrada pela Arpe em 2025"
        }

    def get_post_metodologia_extra_paragraphs(self, row, data_extenso, total_achados) -> tuple:
        # Formatar equipe de fiscalização
        responsaveis_list = [r.strip() for r in str(row["Pessoal Responsável"]).split(",") if r.strip()]
        from database.manager import carregar_responsaveis
        db_resp = carregar_responsaveis()
        
        team_parts = []
        for nome in responsaveis_list:
            match = next((d for d in db_resp if d["nome"].strip().lower() == nome.lower()), None)
            if match:
                team_parts.append({"nome": match["nome"], "matricula": match["matricula"]})
            else:
                team_parts.append({"nome": nome, "matricula": "xxxxxxx/xx"})
        
        # Primeiro parágrafo
        p1_runs = [
            ("As ações de fiscalização foram realizadas pela equipe formada pelos servidores: ", False, False, None)
        ]
        for i, member in enumerate(team_parts):
            p1_runs.append((member["nome"], True, False, None))
            p1_runs.append((f", matrícula nº {member['matricula']}", False, False, None))
            if i < len(team_parts) - 2:
                p1_runs.append((", ", False, False, None))
            elif i == len(team_parts) - 2:
                p1_runs.append((" e ", False, False, None))
                
        fisc_sufixo = f" nos dias {data_extenso}. Nessas ações foram identificados {total_achados} achados de fiscalização nas áreas sob concessão por meio de um levantamento de campo nas Rodovias com evidências de defeitos que poderiam se caracterizar Não Conformidades."
        p1_runs.append((fisc_sufixo, False, False, None))
        
        # Segundo parágrafo
        p2_runs = [
            ("Esses achados foram avaliados tomando por base as orientações do PDCL, em especial, o Subitem 4.2.2.1.3 - Parâmetros Mínimos Exigidos, visualizando que \"os pavimentos deverão ser analisados quanto às suas condições de superfície, conforto, deformabilidade, vida remanescente e segurança”, e ainda que os “parâmetros de aceitabilidade do pavimento para essas condições que deverão ser totalmente atendidas durante o período de CONCESSÃO”.", False, False, None)
        ]
        return ("5. FISCALIZAÇÃO", [p1_runs, p2_runs])

    def render_post_metodologia_extra_content(self, doc, row) -> None:
        """Renderiza o Quadro 1 – Demonstrativo da Elegibilidade do Trecho Fiscalizado para Indicação de NC."""
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        # Título do Quadro 1
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        
        r1 = p.add_run("Quadro 1")
        r1.bold = True
        r1.font.name = 'Aptos'
        r1.font.size = Pt(11)
        
        r2 = p.add_run(" – Demonstrativo da Elegibilidade do Trecho Fiscalizado para Indicação de NC")
        r2.font.name = 'Aptos'
        r2.font.size = Pt(11)

        # Dados da tabela
        dados_tabela = [
            ("PE - 009, PISTA SENTIDO SUL (CRESCENTE)", [
                ("29,08 - 29,14", "ST 6.02", "49,29\n(página 285)", "FD 4,16\n(página 125)"),
                ("33,88-33,90", "ST 6.06", "47,02\n(página 296)", "FD 3,58\n(página 125)"),
                ("34,00-34,02", "ST 6.07", "44,60\n(página 297)", "FD 3,21\n(página 125)"),
                ("34,28-34,40", "ST 6.07", "44,60\n(página 297)", "FD 3,21\n(página 125)"),
                ("34,52-34,56", "ST 6.07", "44,60\n(página 297)", "FD 3,21\n(página 125)"),
                ("34,66-34,68", "ST 6.07", "44,60\n(página 297)", "FD 3,21\n(página 125)"),
                ("34,72-34,74", "ST 6.07", "44,60\n(página 297)", "FD 3,21\n(página 125)"),
                ("36,02-36,13", "ST 4.01", "71,00\n(página 223)", "FE 3,92\n(página 125)"),
                ("36,18-36,32", "ST 4.01", "71,00\n(página 223)", "FE 3,92\n(página 125)"),
                ("36,94-36,98", "ST 4.02", "61,68\n(página 225)", "FE 4,30\n(página 125)"),
                ("37,00-37,10", "ST 4.02", "61,68\n(página 225)", "FE 4,30\n(página 125)"),
                ("37,10-37,12", "ST 4.02", "61,68\n(página 225)", "FE 4,30\n(página 125)"),
                ("37,38-37,38", "ST 4.02", "61,68\n(página 225)", "FE 4,30\n(página 125)"),
                ("37,54-37,58", "ST 4.02", "61,68\n(página 225)", "FE 4,30\n(página 125)"),
                ("37,62-37,66", "ST 4.03", "51,64\n(página 227)", "FD 3,22\n(página 125)"),
                ("38,32+38,32", "ST 4.03", "51,64\n(página 227)", "FD 3,22\n(página 125)"),
                ("41,16-41,22", "ST 3.01", "42,34\n(página 204)", "FD 3,00\n(página 125)"),
                ("41,84-41,88", "ST 3.02", "45,60\n(página 206)", "FD 2,88\n(página 125)"),
            ]),
            ("PE - 009, PISTA SENTIDO NORTE (DECRESCENTE)", [
                ("40,16-40,60", "ST 4.06", "70,36\n(página 233)", "FD 3,74\n(página 125)"),
                ("39,94-39,90", "ST 4.06", "70,36\n(página 233)", "FD 3,74\n(página 125)"),
            ]),
            ("VPE – 034, PISTA SENTIDO SUL (CRESCENTE)", [
                ("00,06-00,36", "ST 5.01", "58,46\n(página 250)", "FE 4,32\n(página 125)"),
                ("00,45-00,48\n(ROTATÓRIA)", "ST 5.01", "58,46\n(página 250)", "FE 4,32\n(página 125)"),
                ("00,52-00,56", "ST 5.01", "58,46\n(página 250)", "FE 4,32\n(página 125)"),
                ("00,58-00,62", "ST 5.01", "58,46\n(página 250)", "FE 4,32\n(página 125)"),
                ("00,68-00,86", "ST 5.01", "58,46\n(página 250)", "FE 4,32\n(página 125)"),
                ("01,56-01,58", "ST 5.02", "35,91\n(página 252)", "FE 4,86\n(página 125)"),
                ("01,76-01,78", "ST 5.02", "35,91\n(página 252)", "FE 4,86\n(página 125)"),
                ("02,04-02,12", "ST 5.03", "35,36\n(página 254)", "FD 4,23\n(página 125)"),
            ]),
            ("VPE - 034, PISTA SENTIDO NORTE (DECRESCENTE)", [
                ("02,78-02,72", "ST 5.11", "60,89\n(página 271)", "FE 4,23\n(página 125)"),
                ("02,16-02,08", "ST 5.11", "60,89\n(página 271)", "FE 4,23\n(página 125)"),
            ])
        ]

        total_rows = 1 + sum(1 + len(items) for _, items in dados_tabela)
        table = doc.add_table(rows=total_rows, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Cabeçalho Principal (Linha 0)
        headers = [
            "TRECHO ELEGÍVEL\n(KMinicial-KMfinal)",
            "SUBTRECHO",
            "IGG\n(Relatório Anual 01/2025)",
            "IRI\n(Relatório Anual 01/2025)"
        ]
        hdr_row = table.rows[0]
        for c_idx, text in enumerate(headers):
            cell = hdr_row.cells[c_idx]
            cell.text = text
            shd = parse_xml(r'<w:shd {} w:fill="DDDDDD"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shd)
            for p_cell in cell.paragraphs:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p_cell.runs:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(10)
                    run.bold = True

        # Preenchimento das seções e dados
        curr_row = 1
        for track_title, items in dados_tabela:
            # Linha de Sentido da Pista (mesclada nas 4 colunas)
            row_track = table.rows[curr_row]
            a = row_track.cells[0]
            b = row_track.cells[3]
            merged_cell = a.merge(b)
            merged_cell.text = track_title
            shd_track = parse_xml(r'<w:shd {} w:fill="FFFFCC"/>'.format(nsdecls('w')))
            merged_cell._tc.get_or_add_tcPr().append(shd_track)
            for p_cell in merged_cell.paragraphs:
                p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p_cell.runs:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(10)
                    run.bold = False
            curr_row += 1

            # Linhas de Dados
            for col0, col1, col2, col3 in items:
                row_data = table.rows[curr_row]
                row_data.cells[0].text = col0
                row_data.cells[1].text = col1
                row_data.cells[2].text = col2
                row_data.cells[3].text = col3
                for cell in row_data.cells:
                    for p_cell in cell.paragraphs:
                        p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p_cell.runs:
                            run.font.name = 'Aptos'
                            run.font.size = Pt(10)
                            run.bold = False
                curr_row += 1

        # Largura das colunas
        col_widths = [Inches(1.8), Inches(1.1), Inches(1.9), Inches(1.9)]
        for row_t in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row_t.cells):
                    row_t.cells[idx].width = width

        doc.add_paragraph()

    def render_quadros(self, doc, row, nc_df, criar_tabela_quadros_fn) -> None:
        from utils import formatar_mes_ano
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Obter dados de NC e PA
        id_fisc = row["ID da Fiscalização"]
        current_ncs = nc_df[nc_df["ID da Fiscalização"] == id_fisc] if nc_df is not None and not nc_df.empty and "ID da Fiscalização" in nc_df.columns else pd.DataFrame()
        
        ncs_reais = pd.DataFrame()
        pas_reais = pd.DataFrame()
        if not current_ncs.empty:
            if "Não Conformidade" in current_ncs.columns:
                ncs_reais = current_ncs[current_ncs["Não Conformidade"].fillna("").astype(str).str.strip() != ""].copy()
            if "Ponto de Atenção" in current_ncs.columns:
                pas_reais = current_ncs[current_ncs["Ponto de Atenção"].fillna("").astype(str).str.strip() != ""].copy()
                
        try:
            mes_ano = formatar_mes_ano(row["Data"]).replace(", ", "/").lower()
        except Exception:
            mes_ano = "junho/2026"

        # 1. Parágrafos introdutórios da seção de Fiscalização
        intro_paragraphs = self.get_quadro_intro_paragraphs(row, "", "")
        for paragraph_runs in intro_paragraphs:
            if not paragraph_runs:
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            for text, bold, italic, _ in paragraph_runs:
                run = p.add_run(text)
                run.font.name = 'Aptos'
                run.font.size = Pt(11)
                if bold:
                    run.bold = True
                if italic:
                    run.italic = True
                    
        # 2. Quadro 2 title
        p7 = doc.add_paragraph()
        p7.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p7.paragraph_format.space_after = Pt(6)
        r7_1 = p7.add_run("Quadro 2")
        r7_1.bold = True
        r7_1.font.name = 'Aptos'
        r7_1.font.size = Pt(11)
        r7_2 = p7.add_run(f" – Determinações para Não Conformidades Identificadas – {mes_ano}")
        r7_2.font.name = 'Aptos'
        r7_2.font.size = Pt(11)
        
        criar_tabela_quadros_fn(doc, ncs_reais, is_pa=False, report_config=self)
        doc.add_paragraph()
        
        # 3. Quadro 3 title
        p8 = doc.add_paragraph()
        p8.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p8.paragraph_format.space_after = Pt(6)
        r8_1 = p8.add_run("Quadro 3")
        r8_1.bold = True
        r8_1.font.name = 'Aptos'
        r8_1.font.size = Pt(11)
        r8_2 = p8.add_run(" – Pontos de Atenção por Rodovia/Sentido")
        r8_2.font.name = 'Aptos'
        r8_2.font.size = Pt(11)
        
        criar_tabela_quadros_fn(doc, pas_reais, is_pa=True, report_config=self)
        
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

    @property
    def quadros_section_title(self) -> str:
        return "5. FISCALIZAÇÃO"

    def get_quadro_intro_paragraphs(self, row, data_extenso, responsaveis_formatted) -> list:
        return [
            [("Os trechos com Não Conformidades registradas no Quadro 1, a seguir, associadas aos respectivos subtrechos, foram avaliadas pelos valores do Índice de Gravidade Global (IGG) que ultrapassaram o limite máximo previsto ≥ 30, como também os valores do Índice Irregularidade Longitudinal (IRI) que ultrapassaram o limite máximo previsto ≥ 2,7 m/km constantes do Relatório Anual 01 de novembro/2025 elaborado pelo Verificador Independente.", False, False, None)],
            [],
            [("A partir das vistorias de campo, e das avaliações das Não Conformidades registradas pelo Verificador Independente, foram consolidadas as informações no Quadro 2, a seguir, as Não Conformidades e Determinações/Recomendações expedidas pela Arpe.", False, False, None)]
        ]

    @property
    def quadro_title_template(self) -> str:
        return "Quadro 1 – Registro das Não Conformidades na Área de Concessão da CRA - {mes_ano}"

    @property
    def nc_table_headers(self) -> list:
        return [
            "TRECHO", 
            "IDENTIFICAÇÃO", 
            "DESCRIÇÃO / INFRAÇÃO", 
            "REGISTRO FOTOGRÁFICO", 
            "FUNDAMENTO DA INFRAÇÃO (CONTRATO DE CONCESSÃO)", 
            "DETERMINAÇÃO / RECOMENDAÇÃO"
        ]

    @property
    def nc_table_col_widths(self) -> list:
        return [Inches(1.40), Inches(1.40), Inches(1.22), Inches(1.69), Inches(1.56)]

    def format_nc_table_total_row(self, table, row_idx, total_ncs) -> None:
        r_total = table.rows[row_idx]
        r_total.cells[0].merge(r_total.cells[1]).merge(r_total.cells[2]).merge(r_total.cells[3])
        r_total.cells[0].text = "TOTAL"
        r_total.cells[4].text = str(total_ncs)
        
        col_widths = self.nc_table_col_widths
        r_total.cells[0].width = col_widths[0] + col_widths[1] + col_widths[2] + col_widths[3]
        r_total.cells[4].width = col_widths[4]

    @property
    def finalizacao_sections_config(self) -> dict:
        return {
            "6": "determinações",
            "7": "recomendações",
            "8": "conclusões"
        }

    def get_determinations_paragraphs(self, total_ncs) -> list:
        return [
            [("Considerando os dispositivos contratuais pertinentes e visando garantir a qualidade dos serviços prestados, determina-se que a CRA tome as seguintes medidas através de um plano de ação:", False, False, None)],
            [("Medidas de Manutenção / Conservação,", True, False, None),
             (" detalhando cronograma com trechos a executar de forma que permita à Arpe uma programação mais efetiva do monitoramento de suas soluções a execução de cada subtrecho, conforme o modelo encaminhado para {ano_anterior} (Cronograma de Conserva Especial do Pavimento CRA {ano_anterior}).", False, False, None)],
            [("Medidas imediatas", True, False, None),
             (" resolutividade das NC de Sinalização, nos prazos estabelecidos no subitem 4.1.3.3.2.4. Tachas e Tachões Refletivos do PDCL, conforme disposto no Quadro 1, na coluna denominada Determinações.", False, False, None)]
        ]

    def get_recommendations_paragraphs(self) -> list:
        return [
            [("Considerando as disposições do Contrato de Concessão, em especial, o Anexo IV - PDCL, Outras Sinalizações, dados do Relatório Anual 01 de novembro/2025 elaborado VI, e outras legislação aplicável, devem ser observadas pela CRA as seguintes recomendações:", False, False, None)],
            [("Levantar a necessidade de reposição da SINALIZAÇÃO por tachas e tachões em todo o complexo viário Express Way,  em especial a retirada das tachar anteriores danificadas que podem causar risco a segurança dos usuários. O vi verificou deficiências em todos os trechos.", False, False, None)]
        ]

    def get_conclusions_paragraphs(self, total_ncs, local_val) -> list:
        from sections.finalizacao.finalizacao import numero_por_extenso
        extenso_ncs = numero_por_extenso(total_ncs)
        return [
            [(f"Tendo em vista as ações de fiscalização realizadas pela Arpe foram constatadas {extenso_ncs} ({total_ncs}) pontos ou trechos fiscalizados que apresentaram Não Conformidades distribuídas majoritariamente na PE 009, estas foram analisadas e caracterizadas a partir dos indicadores do Grupo Condição de Superfície definidos no Contrato de Concessão CT. nº 043/2011.", False, False, None)],
            [(f"Assim considerando a Programação de Conserva Especial do Pavimento - CRA {{ano}}, solicita-se a inclusão destas não conformidades no cronograma detalhado de Conserva Especial do Pavimento que permita à Arpe uma realização mais efetiva dos monitoramentos ao longo de {{ano}}.", False, False, None)],
            [("Recomenda-se, por fim, o encaminhamento deste Relatório de Fiscalização para que Suape, na qualidade de Gestor do Contrato e Regulador desse Sistema Viário, realize as providências cabíveis junto à Concessionária Rota do Atlântico com o objetivo de garantir a regularização das Não Conformidades pendentes apontadas por esta Agência de Regulação.", False, False, None)]
        ]

    def render_apendices(self, doc, row, ncs_reais, pas_reais, fotos_dir, data_fisc, ano, criar_grade_fotos_fn) -> None:
        from utils import adicionar_titulo_secao
        from docx.shared import Pt
        
        p_ap_a = adicionar_titulo_secao(doc, "APÊNDICE A – REGISTROS FOTOGRÁFICOS DAS NÃO CONFORMIDADES")
        p_ap_a.paragraph_format.page_break_before = True
        
        if not ncs_reais.empty:
            criar_grade_fotos_fn(doc, ncs_reais, row.get("Local", ""), fotos_dir, data_fisc, self.key)
        else:
            p_empty = doc.add_paragraph()
            r_empty = p_empty.add_run("Nenhum registro fotográfico de não conformidade cadastrado.")
            r_empty.font.name = 'Aptos'
            r_empty.font.size = Pt(11)

        p_ap_b = adicionar_titulo_secao(doc, "APÊNDICE B – REGISTROS FOTOGRÁFICOS DAS PONTOS DE ATENÇÃO")
        p_ap_b.paragraph_format.page_break_before = True
        
        if not pas_reais.empty:
            criar_grade_fotos_fn(doc, pas_reais, row.get("Local", ""), fotos_dir, data_fisc, self.key)
        else:
            p_empty = doc.add_paragraph()
            r_empty = p_empty.add_run("Nenhum registro fotográfico de ponto de atenção cadastrado.")
            r_empty.font.name = 'Aptos'
            r_empty.font.size = Pt(11)

    @property
    def analyst_title(self) -> str:
        return "Analista de Regulação"

    def get_process_sei_texts(self, row, ano=None) -> list:
        from utils import extrair_mes_ano_numerico, extrair_ano
        if isinstance(row, (str, int)) and ano is None:
            ano = str(row)
            mes_ano = f"01/{ano}"
        else:
            ano = ano or extrair_ano(row.get("Data", "") if hasattr(row, "get") else row["Data"])
            data_val = row.get("Data", "") if hasattr(row, "get") else (row["Data"] if isinstance(row, dict) or hasattr(row, "__getitem__") else "")
            mes_ano = extrair_mes_ano_numerico(data_val)
        return [
            f"RELATÓRIO DE FISCALIZAÇÃO PROCESSO ADMINISTRATIVO Nº {mes_ano} - CTR",
            f"SEI Nº xxxxxxxxxxxx/{ano}-XX"
        ]
