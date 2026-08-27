# -*- coding: utf-8 -*-
import os
import re
import tempfile
import unicodedata
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

def strip_accents(s):
    if not s:
        return ""
    return ''.join(c for c in unicodedata.normalize('NFD', str(s)) if unicodedata.category(c) != 'Mn').upper()

def encontrar_quadro1_table(doc):
    """Localiza a tabela do Quadro 1 em um documento Word (.docx)."""
    if not doc or len(doc.tables) == 0:
        return None
    for t in doc.tables:
        header_text = ' '.join([c.text.upper() for r in t.rows[:min(2, len(t.rows))] for c in r.cells])
        if any(k in header_text for k in ['IDENTIFICAÇÃO', 'ID.NC', 'REGISTRO FOTOGRÁFICO', 'NÃO CONFORMIDADE', 'REFERÊNCIA', 'REFERENCIA']):
            return t
    return doc.tables[0]

def extrair_metadados_anterior(documento_anterior):
    """
    Extrai metadados do documento anterior (.docx):
    N_prev, N_curr, ctr_num, processo_sei_prev, data_vistoria_prev, oficio_num_prev, oficio_data_prev, carta_num_prev, etc.
    """
    meta = {
        "N_prev": "X",
        "N_curr": "X",
        "ctr_num": "XX/XXXX",
        "processo_sei_prev": "XXXXXXXX",
        "data_vistoria_prev": "XX/XX/XXXX",
        "oficio_num_prev": "xxx/xxxx",
        "oficio_data_prev": "xx/xx/xxxx",
        "carta_num_prev": "xxxx/xxxx",
        "carta_data_prev": "xx/xx/xxxx",
        "carta_sei_prev": "xxxxxxxx",
        "is_fiscalizacao": False
    }
    if not documento_anterior:
        return meta

    try:
        filename = getattr(documento_anterior, "name", str(documento_anterior)).lower()
        if "fiscaliza" in filename:
            meta["is_fiscalizacao"] = True
            meta["N_prev"] = 0

        doc = Document(documento_anterior)

        if not meta["is_fiscalizacao"]:
            pattern = re.compile(r'(\d+)(?:\u00ba|\u00b0)?\s*monitoramento', re.IGNORECASE)
            for p in doc.paragraphs:
                m = pattern.search(p.text)
                if m:
                    meta["N_prev"] = int(m.group(1))
                    break

        pattern_ctr = re.compile(r'CTR\s*(?:N\u00ba|n\u00ba|N\u00ba)?\s*(\d+/\d+)', re.IGNORECASE)
        for p in doc.paragraphs:
            m = pattern_ctr.search(p.text)
            if m:
                meta["ctr_num"] = m.group(1)
                break

        pattern_sei = re.compile(r'PROCESSO SEI\s*(?:N\u00ba|n\u00ba|N\u00ba)?\s*([\d\./-]+)', re.IGNORECASE)
        for p in doc.paragraphs:
            m = pattern_sei.search(p.text)
            if m:
                meta["processo_sei_prev"] = m.group(1)
                break

        pattern_vist = re.compile(r'(?:realizada|vistorias).*em\s*(\d{2}/\d{2}/\d{4})', re.IGNORECASE)
        for p in doc.paragraphs:
            m = pattern_vist.search(p.text)
            if m:
                meta["data_vistoria_prev"] = m.group(1)
                break

        pattern_oficio = re.compile(r'Of(?:í|i)cio\s+Arpe\s+DTO\s+(?:n|N)(?:º|o)?\s*([\w/.-]+),\s*de\s*(\d{2}/\d{2}/\d{4})', re.IGNORECASE)
        for p in doc.paragraphs:
            m = pattern_oficio.search(p.text)
            if m:
                meta["oficio_num_prev"] = m.group(1)
                meta["oficio_data_prev"] = m.group(2)
                break

        pattern_carta = re.compile(r'Carta\s+(?:CRC/REG|SAP/PER/ARPE|CRA/REG)\s+(?:n|N)(?:º|o)?\s*([\w/.-]+),\s*de\s*(\d{2}/\d{2}/\d{4})\s*\(\s*Doc\.\s*SEI\s*(?:n|N)(?:º|o)?\s*([\w/.-]+)\)', re.IGNORECASE)
        for p in doc.paragraphs:
            m = pattern_carta.search(p.text)
            if m:
                meta["carta_num_prev"] = m.group(1)
                meta["carta_data_prev"] = m.group(2)
                meta["carta_sei_prev"] = m.group(3)
                break
    except Exception as e:
        print(f"Erro ao extrair metadados do documento anterior: {e}")

    if isinstance(meta["N_prev"], int):
        meta["N_curr"] = meta["N_prev"] + 1

    return meta

def extrair_linhas_quadro1_anterior(documento_anterior):
    """
    Extrai as NCs e descrições do Quadro 1 do documento anterior.
    Retorna uma lista de dicionários [{"id_nc": ..., "info_ou_constatacao": ...}].
    """
    ncs = []
    if not documento_anterior:
        return ncs
    try:
        doc = Document(documento_anterior)
        quadro1 = encontrar_quadro1_table(doc)
        if quadro1:
            col_id = 0
            col_desc = 1
            if len(quadro1.columns) >= 2:
                for r_idx in range(min(2, len(quadro1.rows))):
                    for c_idx, cell in enumerate(quadro1.rows[r_idx].cells):
                        txt = cell.text.strip().upper()
                        if any(k in txt for k in ['REFERÊNCIA', 'REFERENCIA', 'IDENTIFICAÇÃO', 'IDENTIFICACAO', 'ID.NC', 'ID NC', 'ID_NC', 'ITEM']):
                            col_id = c_idx
                        elif any(k in txt for k in ['CONSTATAÇÃO', 'CONSTATACAO', 'DESCRIÇÃO', 'DESCRICAO', 'INFORMAÇÃO', 'INFORMACAO', 'NÃO CONFORMIDADE', 'NAO CONFORMIDADE']):
                            if c_idx != col_id:
                                col_desc = c_idx

            for row_item in quadro1.rows[1:]:
                cells = row_item.cells
                if len(cells) <= max(col_id, col_desc):
                    continue
                txt0 = cells[col_id].text.strip()
                txt1 = cells[col_desc].text.strip()
                if not txt0 or any(h in txt0.upper() for h in ['REFERÊNCIA', 'REFERENCIA', 'IDENTIFICAÇÃO', 'IDENTIFICACAO', 'ID.NC', 'TOTAL', 'QUADRO']):
                    continue
                if txt0 != txt1:
                    ncs.append({
                        "id_nc": txt0,
                        "info_ou_constatacao": txt1
                    })
                elif len(cells) > 1 and txt0:
                    ncs.append({
                        "id_nc": txt0,
                        "info_ou_constatacao": cells[1].text.strip()
                    })
    except Exception as e:
        print(f"Erro ao extrair linhas do Quadro 1 anterior: {e}")
    return ncs

def criar_tabela_quadro1_monitoramento(doc, headers, ncs_from_prev, ncs_reais, col_widths=None):
    """
    Cria a tabela comparativa do Quadro 1 (3 colunas) para relatórios de monitoramento.
    """
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from sections.quadros.quadros import set_cell_shading, set_cell_margins, set_table_borders

    if not ncs_from_prev:
        if ncs_reais is None or ncs_reais.empty:
            p_empty = doc.add_paragraph()
            p_empty.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_empty.add_run("Nenhuma Não Conformidade registrada.")
            run.font.name = 'Aptos'
            run.font.size = Pt(10)
            return None
        for _, nc_row in ncs_reais.iterrows():
            ident = str(nc_row.get("Identificação", "")).strip()
            desc = str(nc_row.get("Não Conformidade", "")).strip()
            ncs_from_prev.append({
                "id_nc": f"{ident} - {desc}" if ident else desc,
                "info_ou_constatacao": str(nc_row.get("Observações", nc_row.get("Legenda da Foto", ""))).strip()
            })

    table = doc.add_table(rows=1 + len(ncs_from_prev), cols=3)
    table.style = 'Table Grid'
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False

    if not col_widths:
        col_widths = [Inches(2.50), Inches(3.80), Inches(1.27)]

    for r in table.rows:
        for c_idx, width in enumerate(col_widths):
            if c_idx < len(r.cells):
                r.cells[c_idx].width = width

    # Cabeçalho
    for c_idx, text in enumerate(headers):
        cell = table.cell(0, c_idx)
        set_cell_shading(cell, "BFBFBF")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        run.bold = True

    # Preencher dados
    for idx, item in enumerate(ncs_from_prev, 1):
        id_nc_val = item["id_nc"]
        info_val = item["info_ou_constatacao"]

        situacao_val = "PENDENTE"
        target_id = id_nc_val.strip().lower()
        if ncs_reais is not None and not ncs_reais.empty:
            for _, nc_row in ncs_reais.iterrows():
                curr_id = str(nc_row.get("Identificação", "")).strip().lower()
                if curr_id and (target_id == curr_id or curr_id in target_id):
                    situacao_val = str(nc_row.get("Situação", "PENDENTE")).strip().upper()
                    break

        # Coluna 0
        c0 = table.cell(idx, 0)
        set_cell_margins(c0, top=100, bottom=100, left=150, right=150)
        c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p0.paragraph_format.space_before = Pt(0)
        p0.paragraph_format.space_after = Pt(0)
        p0.paragraph_format.line_spacing = 1.15
        r0 = p0.add_run(id_nc_val)
        r0.font.name = "Aptos"
        r0.font.size = Pt(10)

        # Coluna 1
        c1 = table.cell(idx, 1)
        set_cell_margins(c1, top=100, bottom=100, left=150, right=150)
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.15
        r1 = p1.add_run(info_val)
        r1.font.name = "Aptos"
        r1.font.size = Pt(10)

        # Coluna 2
        c2 = table.cell(idx, 2)
        set_cell_margins(c2, top=100, bottom=100, left=150, right=150)
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(situacao_val)
        r2.font.name = "Aptos"
        r2.font.size = Pt(10)
        r2.bold = True

    return table

def extrair_ncs_e_fotos_anterior(documento_anterior):
    """
    Lê o documento do monitoramento ou fiscalização anterior (.docx),
    extrai as NCs e suas respectivas fotos da tabela do apêndice.
    
    Retorna uma lista de dicionários com as fotos extraídas.
    """
    result = []
    if not documento_anterior:
        return result

    try:
        # Garante o reposicionamento do ponteiro do arquivo Streamlit (BytesIO)
        if hasattr(documento_anterior, "seek"):
            try:
                documento_anterior.seek(0)
            except Exception:
                pass

        doc = Document(documento_anterior)
        body = doc.element.body

        # Regex para IDs de NC (evitando capturar palavras como 'FOTO')
        nc_pattern = r'\b((?:CRC|CRA|SOCICAM|CAR|GAR|TIP|NC|PET|PETRO)(?:[\s_\.-]*[A-Z0-9\+\-/]+)+)\b'

        # 1. Localizar dinamicamente o Quadro 1 no documento
        quadro1 = None
        for t in doc.tables:
            header_text = ' '.join([c.text.upper() for r in t.rows[:min(2, len(t.rows))] for c in r.cells])
            if 'IDENTIFICAÇÃO' in header_text or 'ID.NC' in header_text or 'REGISTRO FOTOGRÁFICO' in header_text or 'NÃO CONFORMIDADE' in header_text:
                quadro1 = t
                break

        status_ncs = {}
        nc_details = {}
        foto_to_id_map = {}

        if quadro1:
            col_id = 0
            col_desc = 1
            col_foto = 2
            col_det = -1
            
            for r_idx in range(min(2, len(quadro1.rows))):
                for c_idx, cell in enumerate(quadro1.rows[r_idx].cells):
                    txt = cell.text.strip().upper()
                    if 'IDENTIFICAÇÃO' in txt or 'ID.NC' in txt:
                        col_id = c_idx
                    elif 'DESCRIÇÃO' in txt or 'CONSTATAÇÃO' in txt:
                        col_desc = c_idx
                    elif 'REGISTRO FOTOGRÁFICO' in txt:
                        col_foto = c_idx
                    elif 'DETERMINAÇÃO' in txt or 'SITUAÇÃO' in txt:
                        col_det = c_idx

            for r_idx, row in enumerate(quadro1.rows):
                cells = [c.text.strip() for c in row.cells]
                if not cells or len(cells) <= max(col_id, col_desc):
                    continue
                
                txt_id = cells[col_id]
                if not txt_id or any(h in txt_id.upper() for h in ['IDENTIFICAÇÃO', 'ID.NC', 'NÃO CONFORMIDADE', 'TOTAL']):
                    continue
                
                txt_desc = cells[col_desc]
                txt_foto = cells[col_foto] if col_foto < len(cells) else ""
                txt_det = cells[col_det] if col_det != -1 and col_det < len(cells) else "Pendente"
                
                status_ncs[txt_id] = txt_det
                nc_details[txt_id] = {
                    "constatacao": txt_desc,
                    "foto_ref": txt_foto
                }
                
                if txt_foto:
                    foto_to_id_map[txt_foto.upper().strip()] = txt_id
                    m = re.search(r'FOTO\s*0*(\d+)', txt_foto.upper())
                    if m:
                        num = int(m.group(1))
                        foto_to_id_map[f"FOTO {num}"] = txt_id
                        foto_to_id_map[f"FOTO {num:02d}"] = txt_id

        # 2. Varrer o corpo do documento para encontrar as fotos do apêndice
        temp_dir = tempfile.gettempdir()
        extracted_dir = os.path.join(temp_dir, "arpe_extracted_photos")
        os.makedirs(extracted_dir, exist_ok=True)

        current_trecho = "PE-024"
        current_pista = "SUL"
        current_nc_id = None
        current_nc_desc = ""
        
        in_appendix = False
        photo_count_per_nc = {}
        photo_counter_global = 0

        # Loop pelos elementos do corpo
        for child in body:
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                p_text = Paragraph(child, doc).text.strip()
                if not p_text:
                    continue
                
                norm_p = strip_accents(p_text)
                
                # Detecta início do Apêndice com insensibilidade a acentos
                if any(kw in norm_p for kw in ["APENDICE", "MEMORIAL", "FOTOGRAFICO", "REGISTRO FOTOGRAFICO", "EVIDENCIAS FOTOGRAFICAS", "FOTOGRAFIAS", "ANEXO FOTOGRAFICO"]):
                    in_appendix = True
                
                if in_appendix:
                    # Detecta pista ou trecho
                    if "PISTA" in norm_p or "SENTIDO" in norm_p:
                        pista_m = re.search(r'PISTA\s*(?:SENTIDO)?\s*([A-Za-z\s]+)', p_text, re.IGNORECASE)
                        if pista_m:
                            current_pista = pista_m.group(1).strip().upper()
                        trecho_m = re.search(r'(RODOVIA(?:\s+ESTADUAL)?\s*[A-Za-z0-9\s-]+)', p_text, re.IGNORECASE)
                        if trecho_m:
                            current_trecho = trecho_m.group(1).strip()
                    
                    # Detecta ID de NC ou Mapeamento por Foto XX
                    m_foto = re.search(r'(FOTO\s*\d+)', p_text, re.IGNORECASE)
                    if m_foto and m_foto.group(1).upper() in foto_to_id_map:
                        current_nc_id = foto_to_id_map[m_foto.group(1).upper()]
                    else:
                        m_nc = re.search(nc_pattern, p_text, re.IGNORECASE)
                        if m_nc and not m_nc.group(1).upper().startswith("FOTO"):
                            current_nc_id = m_nc.group(1).upper()
                            parts = p_text.split('–', 2) if '–' in p_text else p_text.split('-', 2)
                            if len(parts) >= 2:
                                current_nc_desc = parts[-1].strip()
                            else:
                                current_nc_desc = p_text

            elif tag == 'tbl':
                t = Table(child, doc)
                
                # Detecta início do Apêndice caso esteja na tabela banner
                if not in_appendix:
                    for r in t.rows:
                        for cell in r.cells:
                            c_txt = strip_accents(cell.text)
                            if any(kw in c_txt for kw in ["APENDICE", "MEMORIAL", "FOTOGRAFICO"]):
                                in_appendix = True
                                break
                        if in_appendix:
                            break
                
                if in_appendix:
                    cell_text_all = ' '.join([c.text for r in t.rows for c in r.cells])
                    m_foto_tbl = re.search(r'(FOTO\s*\d+)', cell_text_all, re.IGNORECASE)
                    if m_foto_tbl and m_foto_tbl.group(1).upper() in foto_to_id_map:
                        active_id = foto_to_id_map[m_foto_tbl.group(1).upper()]
                    elif current_nc_id:
                        active_id = current_nc_id
                    else:
                        photo_counter_global += 1
                        active_id = f"NC_{photo_counter_global:02d}"
                    
                    status = status_ncs.get(active_id, "Pendente")
                    if status.strip().lower() == "sanada":
                        continue

                    # Seleção inteligente de célula da imagem e legenda
                    cell_img = None
                    cell_leg = None
                    
                    if len(t.rows) >= 2 and len(t.columns) == 2:
                        embeds_right = t.cell(0, 1)._element.xpath('.//@r:embed')
                        if embeds_right:
                            cell_img = t.cell(0, 1)
                            cell_leg = t.cell(1, 1) if len(t.rows) > 1 else t.cell(0, 1)
                        else:
                            cell_img = t.cell(0, 0)
                            cell_leg = t.cell(1, 0) if len(t.rows) > 1 else t.cell(0, 0)
                    elif len(t.rows) >= 1 and len(t.columns) >= 1:
                        cell_img = t.cell(0, 0)
                        cell_leg = t.cell(1, 0) if len(t.rows) > 1 else t.cell(0, 0)

                    if cell_img:
                        embeds = cell_img._element.xpath('.//@r:embed')
                        if embeds:
                            for rId in embeds:
                                if rId in doc.part.related_parts:
                                    img_part = doc.part.related_parts[rId]
                                    img_bytes = img_part.image.blob
                                    
                                    ext = ".jpg"
                                    if "png" in img_part.content_type:
                                        ext = ".png"
                                    
                                    idx = photo_count_per_nc.get(active_id, 0) + 1
                                    photo_count_per_nc[active_id] = idx
                                    
                                    safe_id = re.sub(r'[/\\:*?"<>|]', '_', active_id)
                                    photo_name = f"{safe_id}_{idx}{ext}"
                                    photo_path = os.path.join(extracted_dir, photo_name)
                                    
                                    with open(photo_path, "wb") as f:
                                        f.write(img_bytes)
                                    
                                    legend_text = cell_leg.text.strip() if cell_leg else ""
                                    
                                    result.append({
                                        "id_nc": active_id,
                                        "constatacao": nc_details.get(active_id, {}).get("constatacao", current_nc_desc),
                                        "pista": current_pista,
                                        "trecho": current_trecho,
                                        "old_photo_path": photo_path,
                                        "old_legend": legend_text
                                    })
                                    break
                                    
    except Exception as e:
        print(f"Erro ao extrair NCs e fotos do documento anterior: {e}")
        
    return result

