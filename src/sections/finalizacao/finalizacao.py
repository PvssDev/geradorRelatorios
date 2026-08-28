from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import pandas as pd
from datetime import datetime
from utils import adicionar_titulo_secao, extrair_ano, formatar_data_extenso
from database.manager import carregar_responsaveis, carregar_coordenadores

def formatar_data_dd_mm_yyyy(data_val):
    if pd.isna(data_val) or not data_val:
        return ""
    try:
        if hasattr(data_val, "to_pydatetime"):
            dt = data_val.to_pydatetime()
        elif isinstance(data_val, datetime):
            dt = data_val
        else:
            data_str = str(data_val).strip()
            for fmt in ("%d/%m/%Y", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
                try:
                    dt = datetime.strptime(data_str, fmt)
                    break
                except ValueError:
                    continue
            else:
                return data_str
        return dt.strftime("%d/%m/%Y")
    except Exception:
        return str(data_val)

def _obter_caminho_foto_valido(foto_val, fotos_dir):
    """
    Resolve o caminho físico de um arquivo de imagem tanto relativo a fotos_dir
    quanto absoluto no sistema de arquivos.
    """
    if foto_val is None or pd.isna(foto_val):
        return None
    f_str = str(foto_val).strip()
    if not f_str or f_str.lower() in ("nan", "none", ""):
        return None
        
    # 1. Se já for um caminho absoluto existente
    if os.path.isabs(f_str) and os.path.exists(f_str):
        return f_str
        
    # 2. Se o caminho direto existir no sistema
    if os.path.exists(f_str):
        return f_str
        
    # 3. Se estiver dentro de fotos_dir
    if fotos_dir and os.path.isdir(fotos_dir):
        # Tenta com o nome completo
        c1 = os.path.join(fotos_dir, f_str)
        if os.path.exists(c1):
            return c1
        # Tenta com o basename (caso f_str tenha vindo com separadores de caminho)
        c2 = os.path.join(fotos_dir, os.path.basename(f_str))
        if os.path.exists(c2):
            return c2
            
        # Tenta correspondência insensível a maiúsculas/minúsculas
        f_base = os.path.basename(f_str).lower()
        try:
            for arq in os.listdir(fotos_dir):
                if arq.lower() == f_base:
                    return os.path.join(fotos_dir, arq)
        except Exception:
            pass

    return None

def criar_grade_fotos(doc, df_fotos, terminal_nc, fotos_dir, data_fisc, tipo_relatorio="CRA"):
    if df_fotos is None or df_fotos.empty:
        return
    
    # Compatibilidade com planilhas antigas
    df_fotos = df_fotos.copy()
    for col in ["Pista", "Trecho", "Foto", "Fotos", "Foto Anterior", "Legenda Anterior", "Observações", "Legenda da Foto", "Não Conformidade", "Identificação"]:
        if col not in df_fotos.columns:
            df_fotos[col] = ""
        
    records = df_fotos.to_dict('records')
    
    if "MONITORAMENTO" in str(tipo_relatorio).upper():
        # Agrupa por Pista mantendo a ordem de inserção
        pistas_unicas = []
        for p in df_fotos["Pista"].tolist():
            p_str = str(p).strip() if not pd.isna(p) else ""
            if not p_str:
                p_str = "Única"
            if p_str not in pistas_unicas:
                pistas_unicas.append(p_str)
                
        for idx_pista, pista_val in enumerate(pistas_unicas):
            # Filtra os itens desta pista
            mask = df_fotos["Pista"].apply(lambda x: (str(x).strip() if not pd.isna(x) else "") == (pista_val if pista_val != "Única" else ""))
            df_pista = df_fotos[mask].copy()
            if df_pista.empty:
                continue
                
            records_pista = df_pista.to_dict('records')
            for idx, rec in enumerate(records_pista):
                # Título da NC (ex: NC_01_SH04 - KM 2 + 300 - Afundamento...)
                p_nc_desc = doc.add_paragraph()
                p_nc_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_nc_desc.paragraph_format.space_before = Pt(12)
                p_nc_desc.paragraph_format.space_after = Pt(6)
                p_nc_desc.paragraph_format.line_spacing = 1.15
                
                ident = str(rec.get("Identificação", "")).strip()
                trecho_val = str(rec.get("Trecho", "")).strip()
                desc_nc = str(rec.get("Não Conformidade", "")).strip()
                
                parts = []
                if ident:
                    parts.append(ident)
                if trecho_val:
                    parts.append(trecho_val)
                if desc_nc:
                    parts.append(desc_nc)
                
                run_nc_desc = p_nc_desc.add_run(" – ".join(parts))
                run_nc_desc.bold = True
                run_nc_desc.font.name = 'Aptos'
                run_nc_desc.font.size = Pt(11)
                
                # Criar tabela de 2 colunas e 2 linhas
                table = doc.add_table(rows=2, cols=2)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                table.allow_autofit = False
                
                table.columns[0].width = Inches(3.635)
                table.columns[1].width = Inches(3.635)
                table.rows[0].cells[0].width = Inches(3.635)
                table.rows[0].cells[1].width = Inches(3.635)
                table.rows[1].cells[0].width = Inches(3.635)
                table.rows[1].cells[1].width = Inches(3.635)
                
                # Foto Esquerda (Anterior)
                p_img_left = table.rows[0].cells[0].paragraphs[0]
                p_img_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img_left.paragraph_format.left_indent = Inches(0)
                p_img_left.paragraph_format.first_line_indent = Inches(0)
                p_img_left.paragraph_format.space_before = Pt(4)
                p_img_left.paragraph_format.space_after = Pt(4)
                
                foto_ant = rec.get("Foto Anterior", "")
                foto_path_left = _obter_caminho_foto_valido(foto_ant, fotos_dir)
                if foto_path_left:
                    try:
                        run_img_left = p_img_left.add_run()
                        run_img_left.add_picture(foto_path_left, width=Inches(2.708), height=Inches(2.708))
                    except Exception as e:
                        print(f"Erro ao adicionar foto anterior: {e}")
                
                # Foto Direita (Nova)
                p_img_right = table.rows[0].cells[1].paragraphs[0]
                p_img_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img_right.paragraph_format.left_indent = Inches(0)
                p_img_right.paragraph_format.first_line_indent = Inches(0)
                p_img_right.paragraph_format.space_before = Pt(4)
                p_img_right.paragraph_format.space_after = Pt(4)
                
                foto_new = rec.get("Foto") or rec.get("Fotos") or rec.get("foto") or ""
                foto_path_right = _obter_caminho_foto_valido(foto_new, fotos_dir)
                if foto_path_right:
                    try:
                        run_img_right = p_img_right.add_run()
                        run_img_right.add_picture(foto_path_right, width=Inches(2.708), height=Inches(2.708))
                    except Exception as e:
                        print(f"Erro ao adicionar foto nova: {e}")
                    
                # Legenda Esquerda
                p_caption_left = table.rows[1].cells[0].paragraphs[0]
                p_caption_left.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_caption_left.paragraph_format.left_indent = Inches(0)
                p_caption_left.paragraph_format.first_line_indent = Inches(0)
                p_caption_left.paragraph_format.space_before = Pt(4)
                p_caption_left.paragraph_format.space_after = Pt(4)
                
                num_left = str(2 * idx + 1).zfill(2)
                legenda_ant = str(rec.get("Legenda Anterior", "")).strip()
                if legenda_ant:
                    if not legenda_ant.lower().startswith("foto"):
                        legenda_ant = f"Foto {num_left} – {legenda_ant}"
                    run_caption_left = p_caption_left.add_run(legenda_ant)
                    run_caption_left.font.name = 'Aptos'
                    run_caption_left.font.size = Pt(10)
                
                # Legenda Direita
                p_caption_right = table.rows[1].cells[1].paragraphs[0]
                p_caption_right.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_caption_right.paragraph_format.left_indent = Inches(0)
                p_caption_right.paragraph_format.first_line_indent = Inches(0)
                p_caption_right.paragraph_format.space_before = Pt(4)
                p_caption_right.paragraph_format.space_after = Pt(4)
                
                num_right = str(2 * idx + 2).zfill(2)
                obs_text = str(rec.get("Observações") or rec.get("Legenda da Foto") or "").strip()
                caption_right_text = f"Foto {num_right} – {obs_text}"
                if not caption_right_text.endswith("."):
                    caption_right_text += "."
                run_caption_right = p_caption_right.add_run(caption_right_text)
                run_caption_right.font.name = 'Aptos'
                run_caption_right.font.size = Pt(10)
                
                doc.add_paragraph() # Espaço entre tabelas
        return

    if str(tipo_relatorio).upper() in ["CRC", "SOCICAM"]:
        # No CRC/SOCICAM cada foto é uma tabela de 1 coluna
        for idx, rec in enumerate(records):
            obs_text = str(rec.get("Observações") or rec.get("Legenda da Foto") or "").strip()
            
            if str(tipo_relatorio).upper() == "CRC":
                p_nc_desc = doc.add_paragraph()
                p_nc_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_nc_desc.paragraph_format.space_before = Pt(12)
                p_nc_desc.paragraph_format.space_after = Pt(6)
                p_nc_desc.paragraph_format.line_spacing = 1.15
                
                ident = str(rec.get("Identificação", "")).strip()
                nc_txt = str(rec.get("Não Conformidade", "")).strip()
                label_item = f"{ident} – {obs_text}" if ident and obs_text else (ident or nc_txt or obs_text or f"Item {idx+1}")
                
                run_nc_desc = p_nc_desc.add_run(label_item)
                run_nc_desc.bold = True
                run_nc_desc.font.name = 'Aptos'
                run_nc_desc.font.size = Pt(11)
            
            # Tabela de 1 coluna com 2 linhas
            table = doc.add_table(rows=2, cols=1)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.allow_autofit = False
            
            # Largura de 3.12 polegadas para SOCICAM e 5.0 para CRC
            tbl_width = Inches(3.12) if str(tipo_relatorio).upper() == "SOCICAM" else Inches(5.0)
            table.rows[0].cells[0].width = tbl_width
            table.rows[1].cells[0].width = tbl_width
            
            # Row 0: Foto
            p_img = table.rows[0].cells[0].paragraphs[0]
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(4)
            p_img.paragraph_format.space_after = Pt(4)
            foto = rec.get("Foto") or rec.get("Fotos") or rec.get("foto") or ""
            foto_path = _obter_caminho_foto_valido(foto, fotos_dir)
            if foto_path:
                try:
                    run_img = p_img.add_run()
                    img_dim = Inches(2.96) if str(tipo_relatorio).upper() == "SOCICAM" else Inches(4.5)
                    run_img.add_picture(foto_path, width=img_dim, height=img_dim)
                except Exception as e:
                    print(f"Erro ao adicionar foto no CRC/SOCICAM: {e}")
                
            # Row 1: Legenda
            p_caption = table.rows[1].cells[0].paragraphs[0]
            p_caption.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_caption.paragraph_format.space_before = Pt(4)
            p_caption.paragraph_format.space_after = Pt(4)
            num = str(rec.get("Nº", idx+1)).zfill(2)
            data_str = f", em {data_fisc}" if data_fisc else ""
            caption_text = f"Foto {num} – {obs_text}{data_str}."
            run_caption = p_caption.add_run(caption_text)
            run_caption.font.name = 'Aptos'
            run_caption.font.size = Pt(10)
            
            doc.add_paragraph() # Espaço entre tabelas
    else:
        # CRA (Fiscalização) - Agrupa por Pista mantendo ordem de inserção
        pistas_unicas = []
        for p in df_fotos["Pista"].tolist():
            p_str = str(p).strip() if not pd.isna(p) else ""
            if not p_str:
                p_str = "Única"
            if p_str not in pistas_unicas:
                pistas_unicas.append(p_str)
                
        for idx_pista, pista_val in enumerate(pistas_unicas):
            # Filtra os itens desta pista
            mask = df_fotos["Pista"].apply(lambda x: (str(x).strip() if not pd.isna(x) else "") == (pista_val if pista_val != "Única" else ""))
            df_pista = df_fotos[mask].copy()
            if df_pista.empty:
                continue
                
            # Adiciona a tabela de grade (2 colunas)
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.columns[0].width = Inches(3.635)
            table.columns[1].width = Inches(3.635)
            
            # Cabeçalho da Pista
            if pista_val.lower().startswith("pista"):
                header_text = f"{terminal_nc}, {pista_val}"
            else:
                header_text = f"{terminal_nc}, Pista sentido {pista_val}"
                
            row_h = table.add_row()
            row_h.cells[0].width = Inches(3.635)
            row_h.cells[1].width = Inches(3.635)
            cell_merged = row_h.cells[0].merge(row_h.cells[1])
            cell_merged.width = Inches(7.27)
            p_h = cell_merged.paragraphs[0]
            p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_h.paragraph_format.space_before = Pt(6)
            p_h.paragraph_format.space_after = Pt(6)
            run_h = p_h.add_run(header_text)
            run_h.bold = True
            run_h.font.name = 'Aptos'
            run_h.font.size = Pt(11)
            
            records_pista = df_pista.to_dict('records')
            for i in range(0, len(records_pista), 2):
                rec_left = records_pista[i]
                rec_right = records_pista[i+1] if i+1 < len(records_pista) else None
                
                # Linha de Imagens
                row_img = table.add_row()
                row_img.cells[0].width = Inches(3.635)
                row_img.cells[1].width = Inches(3.635)
                
                # Imagem Esquerda
                p_img_left = row_img.cells[0].paragraphs[0]
                p_img_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img_left.paragraph_format.left_indent = Inches(0)
                p_img_left.paragraph_format.first_line_indent = Inches(0)
                p_img_left.paragraph_format.space_before = Pt(4)
                p_img_left.paragraph_format.space_after = Pt(4)
                foto_left = rec_left.get("Foto") or rec_left.get("Fotos") or rec_left.get("foto") or ""
                foto_path_left = _obter_caminho_foto_valido(foto_left, fotos_dir)
                if foto_path_left:
                    try:
                        run_img_left = p_img_left.add_run()
                        run_img_left.add_picture(foto_path_left, width=Inches(3.15), height=Inches(3.15))
                    except Exception as e:
                        print(f"Erro ao adicionar foto esquerda no CRA: {e}")
                    
                # Imagem Direita
                p_img_right = row_img.cells[1].paragraphs[0]
                p_img_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_img_right.paragraph_format.left_indent = Inches(0)
                p_img_right.paragraph_format.first_line_indent = Inches(0)
                p_img_right.paragraph_format.space_before = Pt(4)
                p_img_right.paragraph_format.space_after = Pt(4)
                if rec_right:
                    foto_right = rec_right.get("Foto") or rec_right.get("Fotos") or rec_right.get("foto") or ""
                    foto_path_right = _obter_caminho_foto_valido(foto_right, fotos_dir)
                    if foto_path_right:
                        try:
                            run_img_right = p_img_right.add_run()
                            run_img_right.add_picture(foto_path_right, width=Inches(3.15), height=Inches(3.15))
                        except Exception as e:
                            print(f"Erro ao adicionar foto direita no CRA: {e}")
                else:
                    row_img.cells[1].width = Inches(3.635)
                        
                # Linha de Descrições
                row_desc = table.add_row()
                row_desc.cells[0].width = Inches(3.635)
                row_desc.cells[1].width = Inches(3.635)
                
                # Descrição Esquerda
                p_desc_left = row_desc.cells[0].paragraphs[0]
                p_desc_left.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_desc_left.paragraph_format.left_indent = Inches(0)
                p_desc_left.paragraph_format.first_line_indent = Inches(0)
                p_desc_left.paragraph_format.space_before = Pt(4)
                p_desc_left.paragraph_format.space_after = Pt(4)
                
                num_left = str(rec_left.get("Nº", i+1)).zfill(2)
                trecho_left = str(rec_left.get("Trecho", "")).strip()
                obs_left = str(rec_left.get("Observações") or rec_left.get("Legenda da Foto") or "").strip()
                
                trecho_txt_left = f"Trecho {trecho_left} apresentando " if trecho_left else ""
                data_txt = f", ({data_fisc})" if data_fisc else ""
                desc_text_left = f"Foto {num_left} – {trecho_txt_left}{obs_left}{data_txt}."
                run_desc_left = p_desc_left.add_run(desc_text_left)
                run_desc_left.font.name = 'Aptos'
                run_desc_left.font.size = Pt(10)
                
                # Descrição Direita
                p_desc_right = row_desc.cells[1].paragraphs[0]
                p_desc_right.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_desc_right.paragraph_format.left_indent = Inches(0)
                p_desc_right.paragraph_format.first_line_indent = Inches(0)
                p_desc_right.paragraph_format.space_before = Pt(4)
                p_desc_right.paragraph_format.space_after = Pt(4)
                
                if rec_right:
                    num_right = str(rec_right.get("Nº", i+2)).zfill(2)
                    trecho_right = str(rec_right.get("Trecho", "")).strip()
                    obs_right = str(rec_right.get("Observações") or rec_right.get("Legenda da Foto") or "").strip()
                    
                    trecho_txt_right = f"Trecho {trecho_right} apresentando " if trecho_right else ""
                    desc_text_right = f"Foto {num_right} – {trecho_txt_right}{obs_right}{data_txt}."
                    run_desc_right = p_desc_right.add_run(desc_text_right)
                    run_desc_right.font.name = 'Aptos'
                    run_desc_right.font.size = Pt(10)
                
                # Forçar a largura de todas as células das duas linhas para 3.635 polegadas (metade exata)
                row_img.cells[0].width = Inches(3.635)
                row_img.cells[1].width = Inches(3.635)
                row_desc.cells[0].width = Inches(3.635)
                row_desc.cells[1].width = Inches(3.635)
                    
            if idx_pista < len(pistas_unicas) - 1:
                doc.add_paragraph()


def numero_por_extenso(n):
    extenso_map = {
        0: "zero", 1: "um", 2: "dois", 3: "três", 4: "quatro", 5: "cinco",
        6: "seis", 7: "sete", 8: "oito", 9: "nove", 10: "dez",
        11: "onze", 12: "doze", 13: "treze", 14: "quatorze", 15: "quinze",
        16: "dezesseis", 17: "dezessete", 18: "dezoito", 19: "dezenove", 20: "vinte",
        21: "vinte e um", 22: "vinte e dois", 23: "vinte e três", 24: "vinte e quatro",
        25: "vinte e cinco", 26: "vinte e seis", 27: "vinte e sete", 28: "vinte e oito",
        29: "vinte e nove", 30: "trinta", 31: "trinta e um", 32: "trinta e dois",
        33: "trinta e três", 34: "trinta e quatro", 35: "trinta e cinco",
        36: "trinta e seis", 37: "trinta e sete", 38: "trinta e oito", 39: "trinta e nove",
        40: "quarenta", 41: "quarenta e um", 42: "quarenta e dois", 43: "quarenta e três",
        44: "quarenta e quatro", 45: "quarenta e cinco", 46: "quarenta e seis",
    }
    return extenso_map.get(n, str(n))
def gerar_secao_finalizacao(doc: Document, row, total_ncs, nc_df=None, fotos_dir=None, report_config=None):
    """Gera o restante do relatório a partir da seção 5 (Determinações Gerais) até as assinaturas finais e apêndices com fotos."""
    
    ano = extrair_ano(row["Data"])
    local_val = str(row.get("Local", "Terminal Rodoviário de Passageiros do Recife (TIP)"))

    # Helper function to append paragraph runs
    def add_formatted_paragraph(paragraph_runs):
        if not paragraph_runs:
            doc.add_paragraph()
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        for text, bold, italic, color_rgb in paragraph_runs:
            run = p.add_run(text)
            run.font.name = 'Aptos'
            run.font.size = Pt(11)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if color_rgb:
                run.font.color.rgb = RGBColor(*color_rgb)

    # ----------------------------------------------------
    # DYNAMIC SECTIONS (Determinações, Recomendações, Conclusões)
    # ----------------------------------------------------
    sections_config = report_config.finalizacao_sections_config
    
    # Render sections in order
    for num_str, sec_type in sorted(sections_config.items(), key=lambda x: int(x[0])):
        if sec_type == "determinações":
            doc.add_paragraph()  # Espaço
            adicionar_titulo_secao(doc, f"{num_str}. DETERMINAÇÕES GERAIS")
            doc.add_paragraph()  # Pula linha abaixo do título
            for runs in report_config.get_determinations_paragraphs(total_ncs):
                add_formatted_paragraph(runs)
                
        elif sec_type == "recomendações":
            doc.add_paragraph()  # Espaço
            adicionar_titulo_secao(doc, f"{num_str}. RECOMENDAÇÕES")
            doc.add_paragraph()  # Pula linha abaixo do título
            
            # Special formatting for bullets
            recoms = report_config.get_recommendations_paragraphs()
            # The first paragraph is the introduction
            if recoms:
                add_formatted_paragraph(recoms[0])
                for runs in recoms[1:]:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.left_indent = Pt(18.0)
                    p.paragraph_format.first_line_indent = Pt(-18.0)
                    
                    run_bullet = p.add_run("•\t")
                    run_bullet.font.name = 'Aptos'
                    run_bullet.font.size = Pt(11)
                    
                    for text, bold, italic, color_rgb in runs:
                        run = p.add_run(text)
                        run.font.name = 'Aptos'
                        run.font.size = Pt(11)
                        if bold:
                            run.bold = True
                        if italic:
                            run.italic = True
                            
        elif sec_type == "conclusões":
            doc.add_paragraph()  # Espaço
            adicionar_titulo_secao(doc, f"{num_str}. CONCLUSÕES")
            doc.add_paragraph()  # Pula linha abaixo do título
            for runs in report_config.get_conclusions_paragraphs(total_ncs, local_val):
                # Replace placeholder template strings if present
                resolved_runs = []
                for text, bold, italic, color_rgb in runs:
                    resolved_text = text.replace("{ano}", ano).replace("{ano_anterior}", str(int(ano) - 1))
                    resolved_runs.append((resolved_text, bold, italic, color_rgb))
                add_formatted_paragraph(resolved_runs)

        elif sec_type in ("conclusões_monitoramento", "conclusoes_monitoramento"):
            data_extenso = formatar_data_extenso(row["Data"])
            doc.add_paragraph()  # Espaço
            adicionar_titulo_secao(doc, f"{num_str}. CONCLUSÕES E RECOMENDAÇÕES")
            doc.add_paragraph()  # Pula linha abaixo do título
            for runs in report_config.get_conclusions_monitoramento_paragraphs(total_ncs, data_extenso):
                resolved_runs = []
                for text, bold, italic, color_rgb in runs:
                    resolved_text = text.replace("{ano}", ano).replace("{ano_anterior}", str(int(ano) - 1))
                    resolved_runs.append((resolved_text, bold, italic, color_rgb))
                add_formatted_paragraph(resolved_runs)

    # ----------------------------------------------------
    # APÊNDICES E ASSINATURAS FINAIS
    # ----------------------------------------------------
    id_fisc = row["ID da Fiscalização"]
    data_fisc = formatar_data_dd_mm_yyyy(row["Data"])
    
    if nc_df is not None and not nc_df.empty and "ID da Fiscalização" in nc_df.columns:
        mask_id = nc_df["ID da Fiscalização"].astype(str).str.strip() == str(id_fisc).strip()
        current_ncs = nc_df[mask_id].copy()
        if current_ncs.empty:
            current_ncs = nc_df.copy()
    elif nc_df is not None and not nc_df.empty:
        current_ncs = nc_df.copy()
    else:
        current_ncs = pd.DataFrame()    
    ncs_reais = pd.DataFrame()
    pas_reais = pd.DataFrame()
    if not current_ncs.empty:
        is_monitoring = "MONITORAMENTO" in getattr(report_config, "key", "").upper() or getattr(report_config, "is_monitoramento", False)
        if is_monitoring:
            cols_check = [c for c in ["Não Conformidade", "Identificação", "Observações", "Legenda da Foto", "Determinação", "Foto", "Fotos"] if c in current_ncs.columns]
            if cols_check:
                mask_nc = current_ncs[cols_check].fillna("").astype(str).apply(lambda r_c: any(v.strip() != "" for v in r_c), axis=1)
                ncs_reais = current_ncs[mask_nc].copy()
            else:
                ncs_reais = current_ncs.copy()
        else:
            # Em Fiscalização:
            # 1. Pontos de Atenção (específico para CRA)
            if "Ponto de Atenção" in current_ncs.columns:
                mask_pa = current_ncs["Ponto de Atenção"].fillna("").astype(str).str.strip() != ""
                pas_reais = current_ncs[mask_pa].copy()
            
            # 2. Não Conformidades:
            # Se tiver coluna 'Não Conformidade', pega os preenchidos
            if "Não Conformidade" in current_ncs.columns:
                mask_nc = current_ncs["Não Conformidade"].fillna("").astype(str).str.strip() != ""
                ncs_reais = current_ncs[mask_nc].copy()
                
            # Se ncs_reais ficou vazio ou não tem a coluna, busca qualquer linha que tenha Foto/Observação/Identificação
            if ncs_reais.empty:
                cols_fisc = [c for c in ["Foto", "Fotos", "Observações", "Legenda da Foto", "Identificação", "Não Conformidade"] if c in current_ncs.columns]
                if cols_fisc:
                    mask_any = current_ncs[cols_fisc].fillna("").astype(str).apply(lambda r_c: any(v.strip() != "" for v in r_c), axis=1)
                    # Não duplicar com pas_reais se for CRA
                    if not pas_reais.empty and getattr(report_config, "key", "") == "CRA":
                        mask_any = mask_any & (~current_ncs.index.isin(pas_reais.index))
                    ncs_reais = current_ncs[mask_any].copy()
                else:
                    ncs_reais = current_ncs.copy()

    def render_apendices_fn():
        report_config.render_apendices(doc, row, ncs_reais, pas_reais, fotos_dir, data_fisc, ano, criar_grade_fotos)

    def render_assinaturas_fn():
        doc.add_paragraph()
        p_loc1 = doc.add_paragraph()
        p_loc1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_loc1.paragraph_format.space_before = Pt(12)
        p_loc1.paragraph_format.space_after = Pt(24)
        run_loc1 = p_loc1.add_run("Recife, data da assinatura eletrônica.")
        run_loc1.font.name = 'Aptos'
        run_loc1.font.size = Pt(11)
        
        doc.add_paragraph()
        
        responsaveis_list = [r.strip() for r in str(row["Pessoal Responsável"]).split(",") if r.strip()]
        db_resp = carregar_responsaveis()
        
        for nome in responsaveis_list:
            match = next((d for d in db_resp if d["nome"].strip().lower() == nome.lower()), None)
            if match:
                r_nome = match["nome"]
                r_funcao = report_config.analyst_title if report_config.key == "SOCICAM" else match["funcao"]
                r_matr = match["matricula"]
            else:
                r_nome = nome
                r_funcao = report_config.analyst_title
                r_matr = "xxxxxxx/xx"
                
            p_ass = doc.add_paragraph()
            p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_ass.paragraph_format.space_after = Pt(0)
            run_ass = p_ass.add_run(r_nome)
            run_ass.bold = True
            run_ass.font.name = 'Aptos'
            run_ass.font.size = Pt(11)
            
            p_carg = doc.add_paragraph()
            p_carg.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_carg.paragraph_format.space_after = Pt(0)
            run_carg = p_carg.add_run(f"{r_funcao}, Matrícula nº {r_matr}" if report_config.key == "SOCICAM" else r_funcao)
            run_carg.font.name = 'Aptos'
            run_carg.font.size = Pt(11)
            
            if report_config.key != "SOCICAM":
                p_mat = doc.add_paragraph()
                p_mat.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_mat.paragraph_format.space_after = Pt(12)
                run_mat = p_mat.add_run(f"Matrícula nº {r_matr}")
                run_mat.font.name = 'Aptos'
                run_mat.font.size = Pt(11)
            else:
                p_carg.paragraph_format.space_after = Pt(18)
            
            doc.add_paragraph()  # Espaço entre assinaturas
            
        # Ciente e de acordo (Coordenador)
        p_ciente = doc.add_paragraph()
        p_ciente.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ciente.paragraph_format.space_before = Pt(36)
        p_ciente.paragraph_format.space_after = Pt(36)
        run_ciente = p_ciente.add_run("Ciente e de acordo.")
        run_ciente.font.name = 'Aptos'
        run_ciente.font.size = Pt(11)
        
        db_coord = carregar_coordenadores()
        coord_name = str(row["Coordenador"]).strip()
        match_coord = next((c for c in db_coord if c["nome"].strip().lower() == coord_name.lower()), None)
        if match_coord:
            c_nome = match_coord["nome"]
            c_funcao = match_coord["funcao"]
            c_matr = match_coord["matricula"]
        else:
            c_nome = coord_name
            c_funcao = "Coordenador(a) de Transportes e Rodovias"
            c_matr = "xxxxxxx/xx"
            
        p_ass3 = doc.add_paragraph()
        p_ass3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ass3.paragraph_format.space_after = Pt(0)
        run_ass3 = p_ass3.add_run(c_nome)
        run_ass3.bold = True
        run_ass3.font.name = 'Aptos'
        run_ass3.font.size = Pt(11)
        
        p_carg3 = doc.add_paragraph()
        p_carg3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_carg3.paragraph_format.space_after = Pt(0)
        run_carg3 = p_carg3.add_run(c_funcao)
        run_carg3.font.name = 'Aptos'
        run_carg3.font.size = Pt(11)
        
        p_mat3 = doc.add_paragraph()
        p_mat3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_mat3.paragraph_format.space_after = Pt(12)
        run_mat3 = p_mat3.add_run(f"Matrícula nº {c_matr}")
        run_mat3.font.name = 'Aptos'
        run_mat3.font.size = Pt(11)

    if getattr(report_config, "signatures_before_apendices", False):
        render_assinaturas_fn()
        render_apendices_fn()
    else:
        render_apendices_fn()
        render_assinaturas_fn()
