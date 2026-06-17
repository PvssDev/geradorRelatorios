from docx import Document
from utils import adicionar_titulo_secao, adicionar_paragrafo_justificado
from docx.enum.text import WD_ALIGN_PARAGRAPH


def gerar_secao_fundamentacao_legal(doc: Document):
    """
    Gera a seção '2. FUNDAMENTAÇÃO LEGAL' no relatório.
    Esta seção contém texto fixo, baseado na legislação e regulamentação da ARPE.
    """

    adicionar_titulo_secao(doc, "2. FUNDAMENTAÇÃO LEGAL")

    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.add_run(
        "A presente fiscalização encontra fundamento nas seguintes normas legais e regulamentares:\n\n"
    )

    # Lei nº 12.524, de 30 de dezembro de 2003
    par.add_run("- ")
    par.add_run("Lei nº 12.524, de 30 de dezembro de 2003").bold = True
    par.add_run(" – Altera e consolida as disposições da ")
    par.add_run("Lei nº 12.126, de 12 de dezembro de 2001").bold = True
    par.add_run(
        ", que cria a Agência de Regulação dos Serviços Públicos Delegados do Estado de Pernambuco – ARPE, regulamentada pelo "
    )
    par.add_run("Decreto nº 30.200, de 09 de fevereiro de 2007").bold = True
    par.add_run(".\n\n")

    # Lei nº 13.254, de 21 de junho de 2007 e alterações, em especial a Lei Estadual nº 15.200, de 17 de dezembro de 2013
    par.add_run("- ")
    par.add_run("Lei nº 13.254, de 21 de junho de 2007").bold = True
    par.add_run(" e alterações, em especial a ")
    par.add_run("Lei Estadual nº 15.200, de 17 de dezembro de 2013").bold = True
    par.add_run(
        " – Estrutura o Sistema de Transporte Coletivo Intermunicipal de Passageiros do Estado de Pernambuco, regulamentada pelo "
    )
    par.add_run("Decreto nº 40.559, de 31 de março de 2014").bold = True
    par.add_run(".\n\n")

    # Resolução Arpe nº 46, de 07 de abril de 2008 (Antiga nº 06/2008)
    par.add_run("- ")
    par.add_run("Resolução Arpe nº 46, de 07 de abril de 2008").bold = True
    par.add_run(
        " (Antiga nº 06/2008) – Aprova o Regulamento dos Terminais Rodoviários do Estado de Pernambuco, alterada parcialmente pela "
    )
    par.add_run("Resolução ARPE nº 53, de 26 de janeiro de 2009").bold = True
    par.add_run(" (Antiga 003/2009).\n\n")

    # Resolução Arpe nº 083, de 30 de julho de 2013
    par.add_run("- ")
    par.add_run("Resolução Arpe nº 083, de 30 de julho de 2013").bold = True
    par.add_run(
        " – Dispõe sobre os procedimentos de fiscalização, autuação e aplicação de penalidades aos prestadores de serviços públicos delegados no Estado de Pernambuco fiscalizados pela ARPE mediante delegação.\n\n"
    )

    # Contrato de Concessão de Serviço Público nº 1.041.080/08, de 19 de setembro de 2008, e seus aditivos, especialmente o Segundo Termo Aditivo de 29 de setembro de 2017
    par.add_run("- ")
    par.add_run(
        "Contrato de Concessão de Serviço Público nº 1.041.080/08, de 19 de setembro de 2008"
    ).bold = True
    par.add_run(", e seus aditivos, especialmente o ")
    par.add_run("Segundo Termo Aditivo de 29 de setembro de 2017").bold = True
    par.add_run(
        " – contrato celebrado entre o Estado de Pernambuco, representado pela Secretaria de Transportes – SETRA, e a SOCICAM – Administração, Projetos e Representações Ltda."
    )
