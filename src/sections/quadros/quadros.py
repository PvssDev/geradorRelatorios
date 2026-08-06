from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

MAP_SIGLAS = {
    "FI": "Fissuras",
    "TTC": "Trincas isoladas transversais curtas",
    "TTL": "Trincas isoladas transversais longas",
    "TLC": "Trincas isoladas longitudinais curtas",
    "TLL": "Trincas isoladas longitudinais longas",
    "J": "Trincas interligadas sem erosão acentuada nas bordas das trincas",
    "JE": "Trincas interligadas com erosão acentuada nas bordas das trincas",
    "TRR": "Trincas isoladas no revestimento devido à retração térmica ou dissecação da base (solo-cimento) ou do revestimento",
    "TB": "Trincas interligadas sem erosão acentuada nas bordas das trincas",
    "TBE": "Trincas interligadas com erosão acentuada nas bordas das trincas",
    "ALP": "Afundamento local plástico devido à fluência plástica de uma ou mais camadas do pavimento ou do subleito",
    "ATP": "Afundamento da trilha plástico devido à fluência plástica de uma ou mais camadas do pavimento ou do subleito",
    "ALC": "Afundamento local de consolidação devido à consolidação diferencial ocorrente em camadas do pavimento ou do subleito",
    "ATC": "Afundamento da trilha de consolidação devido à consolidação diferencial ocorrente em camadas do pavimento ou do subleito",
    "O": "Ondulação/Corrugação, caracterizada por ondulações transversais causadas por instabilidade da mistura betuminosa constituinte do revestimento ou da base",
    "E": "Escorregamento do revestimento betuminoso",
    "EX": "Exsudação do ligante betuminoso no revestimento",
    "D": "Desgaste acentuado na superfície do revestimento",
    "P": "buracos decorrentes da desagregação do revestimento e às vezes de camadas inferiores",
    "R": "Remendo"
}

def expandir_siglas(siglas_str):
    if not siglas_str or pd.isna(siglas_str):
        return ""
    parts = [p.strip() for p in str(siglas_str).split(",") if p.strip()]
    
    # Load custom NCs and merge into local map
    local_map = MAP_SIGLAS.copy()
    try:
        import os
        import json
        quadros_dir = os.path.dirname(os.path.abspath(__file__))
        custom_ncs_path = os.path.abspath(os.path.join(quadros_dir, "..", "..", "database", "custom_ncs.json"))
        if os.path.exists(custom_ncs_path):
            with open(custom_ncs_path, "r", encoding="utf-8") as f:
                custom_list = json.load(f)
                for item in custom_list:
                    sigla = item.get("sigla", "")
                    desc = item.get("descricao", "")
                    if sigla:
                        local_map[sigla] = desc
    except Exception as e:
        print(f"Error loading custom NCs in expandir_siglas: {e}")
        
    expanded = []
    for p in parts:
        if p in local_map:
            expanded.append(f"{local_map[p]} ({p})")
        else:
            expanded.append(p)
    return ", ".join(expanded)

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    for tcMar in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(tcMar)
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    """Aplica bordas pretas e mais grossas (2.5 pt = sz 20) na tabela para maior legibilidade."""
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
        
    tblBorders = OxmlElement('w:tblBorders')
    
    # sz='20' significa 2.5 pt (cada unidade é 1/8 pt).
    borders = {
        'top': {'val': 'single', 'sz': '20', 'space': '0', 'color': '000000'},
        'left': {'val': 'single', 'sz': '20', 'space': '0', 'color': '000000'},
        'bottom': {'val': 'single', 'sz': '20', 'space': '0', 'color': '000000'},
        'right': {'val': 'single', 'sz': '20', 'space': '0', 'color': '000000'},
        'insideH': {'val': 'single', 'sz': '20', 'space': '0', 'color': '000000'},
        'insideV': {'val': 'single', 'sz': '20', 'space': '0', 'color': '000000'}
    }
    
    for border_name, border_attrs in borders.items():
        border = OxmlElement(f'w:{border_name}')
        for attr, val in border_attrs.items():
            border.set(qn(f'w:{attr}'), val)
        tblBorders.append(border)
        
    tblPr.append(tblBorders)

def agrupar_registros(df):
    """
    Agrupa registros que possuem as mesmas características de infração,
    combinando seus números de fotos na mesma linha.
    """
    if df.empty:
        return []
        
    records = df.to_dict('records')
    grouped = []
    
    for rec in records:
        found = False
        for g_rec in grouped:
            chaves = [
                "Identificação", "Não Conformidade", "Ponto de Atenção",
                "Direção (faixa)", "Pista", "Trecho", "Observações",
                "Fundamento da infração", "Determinação"
            ]
            
            match = True
            for key in chaves:
                val1 = str(rec.get(key, "")).strip()
                val2 = str(g_rec.get(key, "")).strip()
                if val1 != val2:
                    match = False
                    break
            
            if match:
                num_foto = rec.get("Nº")
                if num_foto not in g_rec["foto_numeros"]:
                    g_rec["foto_numeros"].append(num_foto)
                found = True
                break
                
        if not found:
            new_g_rec = rec.copy()
            new_g_rec["foto_numeros"] = [rec.get("Nº")]
            grouped.append(new_g_rec)
            
    # Formatar o texto de localização/foto de cada grupo
    for g_rec in grouped:
        faixa = str(g_rec.get("Direção (faixa)", "")).strip()
        nums = sorted(g_rec["foto_numeros"])
        
        nums_formatted = [str(n).zfill(2) for n in nums]
        if len(nums_formatted) == 1:
            fotos_str = f"Foto {nums_formatted[0]}"
        elif len(nums_formatted) == 2:
            fotos_str = f"Fotos {nums_formatted[0]} e {nums_formatted[1]}"
        else:
            fotos_str = f"Fotos {', '.join(nums_formatted[:-1])} e {nums_formatted[-1]}"
            
        g_rec["localizacao_formatada"] = f"{faixa}/ {fotos_str}" if faixa else fotos_str
        
    return grouped

def criar_tabela_quadros(doc, df_dados, is_pa, report_config):
    """Cria a tabela formatada de acordo com o padrão do documento de referência."""
    grouped_records = agrupar_registros(df_dados)
    num_rows = len(grouped_records)
    if num_rows == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Nenhum registro encontrado.")
        run.font.name = 'Aptos'
        run.font.size = Pt(10)
        return
        
    if report_config.key == "SOCICAM":
        table = doc.add_table(rows=1 + num_rows, cols=5)
    else:
        table = doc.add_table(rows=2 + num_rows, cols=5)
        
    table.style = 'Table Grid'
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    # 1. Mesclar células do cabeçalho
    if report_config.key != "SOCICAM":
        table.cell(0, 0).merge(table.cell(0, 1))
        table.cell(0, 2).merge(table.cell(1, 2))
        table.cell(0, 3).merge(table.cell(1, 3))
        table.cell(0, 4).merge(table.cell(1, 4))
    
    # 2. Definir os textos do cabeçalho
    if report_config.key == "SOCICAM":
        table.cell(0, 0).text = "IDENTIFICAÇÃO"
        table.cell(0, 1).text = "DESCRIÇÃO"
        table.cell(0, 2).text = "REGISTRO\nFOTOGRÁFICO"
        table.cell(0, 3).text = "FUNDAMENTO DA INFRAÇÃO\n(ANEXO V CONTRATO DE CONCESSÃO)"
        table.cell(0, 4).text = "DETERMINAÇÃO"
        header_color = "D9D9D9"
        col_widths = report_config.nc_table_col_widths
    elif not is_pa:
        if report_config.key == "CRA":
            table.cell(0, 0).text = "NÃO CONFORMIDADE (NC)"
            table.cell(1, 0).text = "IDENTIFICAÇÃO"
            table.cell(1, 1).text = "DESCRIÇÃO"
            table.cell(0, 2).text = "LOCALIZAÇÃO/\nREGISTRO FOTOGRÁFICO"
            table.cell(0, 3).text = "FUNDAMENTO DA INFRAÇÃO (PDCL)"
            table.cell(0, 4).text = "DETERMINAÇÃO"
            header_color = "D9D9D9"  # Cinza Quadro 2
            col_widths = [Inches(1.40), Inches(1.40), Inches(1.22), Inches(1.69), Inches(1.56)]
        else:
            table.cell(0, 0).text = "NÃO CONFORMIDADE (NC)"
            table.cell(1, 0).text = "IDENTIFICAÇÃO"
            table.cell(1, 1).text = "DESCRIÇÃO"
            table.cell(0, 2).text = "REGISTRO FOTOGRÁFICO"
            table.cell(0, 3).text = "FUNDAMENTO DA INFRAÇÃO (PER ANEXO IV - CONTRATO DE CONCESSÃO)"
            table.cell(0, 4).text = "DETERMINAÇÃO"
            header_color = "D9D9D9"
            col_widths = report_config.nc_table_col_widths
    else:
        table.cell(0, 0).text = "PONTOS DE ATENÇÃO"
        table.cell(1, 0).text = "IDENTIFICAÇÃO"
        table.cell(1, 1).text = "DESCRIÇÃO DA EVIDÊNCIA"
        table.cell(0, 2).text = "LOCALIZAÇÃO/\nREGISTRO FOTOGRÁFICO"
        table.cell(0, 3).text = "JUSTIFICATIVA PARA PONTOS DE ATENÇÃO"
        table.cell(0, 4).text = "SOLICITAÇÃO"
        header_color = "E7E6E6"  # Cinza Claro Quadro 3
        col_widths = [Inches(1.18), Inches(1.74), Inches(1.14), Inches(1.81), Inches(1.40)]
        
    # Formatar cabeçalhos
    header_rows = [0] if report_config.key == "SOCICAM" else [0, 1]
    for r_idx in header_rows:
        row = table.rows[r_idx]
        for c_idx, cell in enumerate(row.cells):
            set_cell_shading(cell, header_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            if cell.paragraphs:
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                
                # O python-docx ao atribuir text pode limpar runs, garantimos estilo nos runs gerados
                for run in p.runs:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(10)
                    run.bold = True
                    
    # Preencher dados
    start_r_idx = 1 if report_config.key == "SOCICAM" else 2
    for idx, rec in enumerate(grouped_records):
        r_idx = idx + start_r_idx
        row = table.rows[r_idx]
        
        ident = str(rec.get("Identificação", "")).strip()
        
        siglas_col = "Não Conformidade" if not is_pa else "Ponto de Atenção"
        siglas_str = rec.get(siglas_col, "")
        desc = expandir_siglas(siglas_str)
        
        localizacao = rec["localizacao_formatada"]
        
        fund = str(rec.get("Fundamento da infração", "")).strip()
        det = str(rec.get("Determinação", "")).strip()
        
        row.cells[0].text = ident
        row.cells[1].text = desc
        row.cells[2].text = localizacao
        row.cells[3].text = fund
        row.cells[4].text = det
        
        for c_idx, cell in enumerate(row.cells):
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            if cell.paragraphs:
                p = cell.paragraphs[0]
                if c_idx in [0, 2, 3, 4]:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                
                for run in p.runs:
                    run.font.name = 'Aptos'
                    run.font.size = Pt(10)
                    
    # Aplicar larguras fixas de forma robusta
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell.width = col_widths[c_idx]

    if report_config.key in ["CRC", "SOCICAM"] and not is_pa:
        # Add TOTAL row
        r_total = table.add_row()
        report_config.format_nc_table_total_row(table, len(table.rows) - 1, num_rows)
        
        # Format the cells
        for c_idx, cell in enumerate(r_total.cells):
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.font.name = 'Aptos'
                run.font.size = Pt(10)
                run.bold = True

def gerar_secao_quadros(doc: Document, row, nc_df, report_config):
    """Gera a seção com as descrições e títulos dos Quadros (Quadros 1 a 5)."""
    report_config.render_quadros(doc, row, nc_df, criar_tabela_quadros)
