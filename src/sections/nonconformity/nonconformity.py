from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import pandas as pd
from utils import (
    adicionar_paragrafo_justificado,
    adicionar_titulo_secao,
    aplicar_borda_paragrafo,
    adicionar_legenda_formatada,
    processar_imagem_para_relatorio,
)


def gerar_secao_nao_conformidades_constatadas(
    doc, row, nao_conformidades_df, fotos_dir, observacoes_df, recomendacoes_df
):
    """
    Gera a seção '3. NÃO CONFORMIDADES CONSTATADAS',
    incluindo Observações Importantes e Recomendações.
    """

    adicionar_titulo_secao(doc, "3. NÃO CONFORMIDADES CONSTATADAS")

    adicionar_paragrafo_justificado(
        doc,
        "A seguir, apresentam-se as não conformidades registradas nos diversos terminais fiscalizados:",
    )

    id_fisc = row["ID da Fiscalização"]
    nc_fisc = nao_conformidades_df[
        nao_conformidades_df["ID da Fiscalização"] == id_fisc
    ]

    if "Terminal" not in nc_fisc.columns:
        adicionar_paragrafo_justificado(
            doc, "⚠️ Coluna 'Terminal' não encontrada na planilha de não conformidades."
        )
        return

    num_terminal = 1
    for terminal, grupo_terminal in nc_fisc.groupby("Terminal"):
        if "(" in terminal and ")" in terminal:
            sigla_terminal = terminal.split("(")[-1].replace(")", "").strip()
        else:
            sigla_terminal = ""

        # Título do terminal
        par_terminal = doc.add_paragraph()
        par_terminal.add_run(f"3.{num_terminal} - {terminal.upper()}").bold = True

        num_nc = 1
        for nc_id, grupo_nc in grupo_terminal.groupby("Nº"):
            descricao = grupo_nc["Não Conformidade"].iloc[0]

            par_nc = doc.add_paragraph()
            run_nc_titulo = par_nc.add_run(
                f"Não Conformidade {sigla_terminal} {str(num_nc).zfill(2)}"
            )
            run_nc_titulo.bold = True
            run_nc_titulo.underline = True
            run_nc_titulo.font.size = Pt(10)
            run_nc_titulo.font.color.rgb = RGBColor(0, 0, 0)

            run_nc_desc = par_nc.add_run(f" – {descricao}")
            run_nc_desc.font.size = Pt(10)
            run_nc_desc.font.color.rgb = RGBColor(0, 0, 0)

            for _, linha in grupo_nc.iterrows():
                nomes_fotos = (
                    [f.strip() for f in str(linha["Foto"]).split(";") if f.strip()]
                    if pd.notna(linha["Foto"])
                    else []
                )
                legendas = (
                    [l.strip() for l in str(linha["Legenda da Foto"]).split(";")]
                    if pd.notna(linha["Legenda da Foto"])
                    else []
                )

                for idx, nome_foto in enumerate(nomes_fotos):
                    foto_path = os.path.join(fotos_dir, nome_foto)
                    legenda = legendas[idx] if idx < len(legendas) else ""
                    if os.path.exists(foto_path):
                        buffer = processar_imagem_para_relatorio(foto_path)
                        doc.add_picture(buffer, width=Inches(3))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        aplicar_borda_paragrafo(doc.paragraphs[-1])
                        adicionar_legenda_formatada(doc, legenda)
                    elif legenda:
                        adicionar_paragrafo_justificado(doc, legenda)

            num_nc += 1

        # OBSERVAÇÕES IMPORTANTES
        obs_terminais = observacoes_df[
            (observacoes_df["ID da Fiscalização"] == id_fisc)
            & (observacoes_df["Terminal"] == terminal)
        ]
        if not obs_terminais.empty:
            adicionar_titulo_secao(doc, "OBSERVAÇÕES IMPORTANTES:")
            run_titulo_obs = doc.paragraphs[-1].runs[0]
            run_titulo_obs.text = run_titulo_obs.text.upper()
            run_titulo_obs.underline = True
            run_titulo_obs.font.size = Pt(10)

            num_obs = 1
            for _, obs in obs_terminais.iterrows():
                textos_obs = (
                    [o.strip() for o in str(obs["Observações"]).split(";") if o.strip()]
                    if pd.notna(obs["Observações"])
                    else []
                )
                nomes_fotos_obs = (
                    [f.strip() for f in str(obs["Foto"]).split(";") if f.strip()]
                    if pd.notna(obs["Foto"])
                    else []
                )
                legendas_obs = (
                    [l.strip() for l in str(obs["Legenda da Foto"]).split(";")]
                    if pd.notna(obs["Legenda da Foto"])
                    else []
                )

                for i, texto_obs in enumerate(textos_obs):
                    adicionar_paragrafo_justificado(doc, f"{num_obs}. {texto_obs}")

                    if i < len(nomes_fotos_obs):
                        foto_path = os.path.join(fotos_dir, nomes_fotos_obs[i])
                        legenda_obs = legendas_obs[i] if i < len(legendas_obs) else ""
                        if os.path.exists(foto_path):
                            buffer = processar_imagem_para_relatorio(foto_path)
                            doc.add_picture(buffer, width=Inches(3))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            aplicar_borda_paragrafo(doc.paragraphs[-1])
                            if legenda_obs:
                                adicionar_legenda_formatada(doc, legenda_obs)
                        elif legenda_obs:
                            adicionar_legenda_formatada(doc, legenda_obs)

                    num_obs += 1

        # RECOMENDAÇÕES
        rec_terminais = recomendacoes_df[
            (recomendacoes_df["ID da Fiscalização"] == id_fisc)
            & (recomendacoes_df["Terminal"] == terminal)
        ]
        if not rec_terminais.empty:
            adicionar_titulo_secao(doc, "RECOMENDAÇÕES:")
            run_titulo_rec = doc.paragraphs[-1].runs[0]
            run_titulo_rec.text = run_titulo_rec.text.upper()
            run_titulo_rec.underline = True
            run_titulo_rec.font.size = Pt(10)

            num_rec = 1
            for _, rec in rec_terminais.iterrows():
                textos_rec = (
                    [r.strip() for r in str(rec["Recomendação"]).split(";") if r.strip()]
                    if pd.notna(rec["Recomendação"])
                    else []
                )
                nomes_fotos_rec = (
                    [f.strip() for f in str(rec["Foto"]).split(";") if f.strip()]
                    if pd.notna(rec["Foto"])
                    else []
                )
                legendas_rec = (
                    [l.strip() for l in str(rec["Legenda da Foto"]).split(";")]
                    if pd.notna(rec["Legenda da Foto"])
                    else []
                )

                for i, texto_rec in enumerate(textos_rec):
                    adicionar_paragrafo_justificado(doc, f"{num_rec}. {texto_rec}")

                    if i < len(nomes_fotos_rec):
                        foto_path = os.path.join(fotos_dir, nomes_fotos_rec[i])
                        legenda_rec = legendas_rec[i] if i < len(legendas_rec) else ""
                        if os.path.exists(foto_path):
                            buffer = processar_imagem_para_relatorio(foto_path)
                            doc.add_picture(buffer, width=Inches(3))
                            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            aplicar_borda_paragrafo(doc.paragraphs[-1])
                            if legenda_rec:
                                adicionar_legenda_formatada(doc, legenda_rec)
                        elif legenda_rec:
                            adicionar_legenda_formatada(doc, legenda_rec)

                    num_rec += 1

        num_terminal += 1
