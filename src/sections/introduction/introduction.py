from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils import adicionar_titulo_secao, formatar_data_extenso, extrair_ano
import os
import pandas as pd

def gerar_secao_introducao(doc: Document, row, total_achados, report_config, nc_df=None):
    """Gera as seções de Introdução, Objetivo, Informações Gerais, Metodologia e Fiscalização."""

    # Delega para implementação própria quando o relatório tem estrutura diferente
    if hasattr(report_config, "gerar_secoes_monitoramento"):
        report_config.gerar_secoes_monitoramento(doc, row, nc_df, total_achados)
        return


    ano = extrair_ano(row["Data"])
    data_extenso = formatar_data_extenso(row["Data"])

    # Helper function to append paragraph runs
    def add_formatted_paragraph(paragraph_runs):
        if not paragraph_runs:
            doc.add_paragraph()
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        for text, bold, italic, color_rgb in paragraph_runs:
            run = p.add_run(text)
            run.font.name = 'Aptos'
            run.font.size = Pt(11)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if color_rgb:
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor(*color_rgb)

    # ----------------------------------------------------
    # 1. SEÇÃO: INTRODUÇÃO (Seção 1)
    # ----------------------------------------------------
    adicionar_titulo_secao(doc, "1. INTRODUÇÃO")
    doc.add_paragraph()  # Pula uma linha abaixo do título
    
    intro_paragraphs = report_config.get_intro_paragraphs(row, ano, data_extenso)
    for paragraph_runs in intro_paragraphs:
        add_formatted_paragraph(paragraph_runs)

    # ----------------------------------------------------
    # 2. SEÇÃO: OBJETIVO (Seção 2)
    # ----------------------------------------------------
    adicionar_titulo_secao(doc, "2. OBJETIVO")
    doc.add_paragraph()  # Pula uma linha abaixo do título
    
    objective_paragraphs = report_config.get_objective_paragraphs(row)
    for paragraph_runs in objective_paragraphs:
        add_formatted_paragraph(paragraph_runs)

    # ----------------------------------------------------
    # 3. SEÇÃO: INFORMAÇÕES GERAIS (Tabela)
    # ----------------------------------------------------
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_info.paragraph_format.space_before = Pt(12)
    p_info.paragraph_format.space_after = Pt(12)
    
    info_title = "3. INFORMAÇÕES GERAIS" if report_config.key == "CRC" else "INFORMAÇÕES GERAIS"
    run_info = p_info.add_run(info_title)
    run_info.bold = True
    run_info.font.name = 'Aptos'
    run_info.font.size = Pt(12)
    
    # Construção da tabela de Informações Gerais
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    responsaveis_formatted = str(row["Pessoal Responsável"]).replace(",", " e" if "," not in str(row["Pessoal Responsável"]) else ";")
    periodo_val = str(row["Período"]).strip() if pd.notna(row["Período"]) and str(row["Período"]).strip() else f"{data_extenso}."
    
    rows_data = report_config.get_general_info_rows(row, responsaveis_formatted, periodo_val)
    num_rows = len(rows_data)
    
    table = doc.add_table(rows=num_rows, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    def format_header_row(row, text):
        cell = row.cells[0].merge(row.cells[1])
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Aptos'
            
        # Shading XML
        tcPr = cell._tc.get_or_add_tcPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'DDDDDD')
        tcPr.append(shd)
        
    def format_normal_row(row, label, val, val_bold=False):
        # Col 0 (Label)
        cell_lbl = row.cells[0]
        cell_lbl.text = label
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_lbl.paragraph_format.space_before = Pt(4)
        p_lbl.paragraph_format.space_after = Pt(4)
        for run in p_lbl.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Aptos'
            
        # Col 1 (Value)
        cell_val = row.cells[1]
        cell_val.text = val
        p_val = cell_val.paragraphs[0]
        p_val.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_val.paragraph_format.space_before = Pt(4)
        p_val.paragraph_format.space_after = Pt(4)
        for run in p_val.runs:
            if val_bold:
                run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Aptos'

    for r_idx, (field_label, field_val, is_header, val_bold) in enumerate(rows_data):
        if is_header:
            format_header_row(table.rows[r_idx], field_label)
        else:
            format_normal_row(table.rows[r_idx], field_label, field_val, val_bold)
            
    col_widths = report_config.general_info_col_widths
    headers_indices = report_config.general_info_headers_indices
    for r_idx, row_obj in enumerate(table.rows):
        if r_idx in headers_indices:
            row_obj.cells[0].width = Inches(7.57) if report_config.key == "SOCICAM" else Inches(7.5)
        else:
            row_obj.cells[0].width = col_widths[0]
            row_obj.cells[1].width = col_widths[1]

    # ----------------------------------------------------
    # 4. SEÇÃO: METODOLOGIA
    # ----------------------------------------------------
    doc.add_paragraph()  # Pula uma linha antes do título
    metodo_title = "3. METODOLOGIA" if report_config.key == "CRA" else "4. METODOLOGIA"
    adicionar_titulo_secao(doc, metodo_title)
    doc.add_paragraph()  # Pula uma linha abaixo do título
    
    # Parágrafos iniciais da Metodologia
    methodology_paragraphs = report_config.get_methodology_paragraphs(data_extenso)
    for paragraph_runs in methodology_paragraphs:
        add_formatted_paragraph(paragraph_runs)

    # Referências com recuo de parágrafo estruturado nativamente no Word (Hanging Indent)
    referencias = report_config.get_references_bullets()
    
    for text, recuado, is_bullet in referencias:
        if text == "":
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        if recuado:
            p.paragraph_format.left_indent = Pt(70.8)
            run = p.add_run(text)
            run.font.name = 'Aptos'
            run.font.size = Pt(9)
            run.font.italic = True
        elif is_bullet:
            p.paragraph_format.left_indent = Pt(report_config.references_left_indent_pt)
            p.paragraph_format.first_line_indent = Pt(-18.0)
            run_bullet = p.add_run("•\t")
            run_bullet.font.name = 'Aptos'
            run_bullet.font.size = Pt(11)
            run = p.add_run(text)
            run.font.name = 'Aptos'
            run.font.size = Pt(11)
        else:
            run = p.add_run(text)
            run.font.name = 'Aptos'
            run.font.size = Pt(11)

    doc.add_paragraph()
    
    post_methodology_paragraphs = report_config.get_post_methodology_paragraphs(total_achados)
    for paragraph_runs in post_methodology_paragraphs:
        add_formatted_paragraph(paragraph_runs)
        
    # Níveis de NC para relatórios rodoviários
    levels_data = report_config.get_levels_data()
    if levels_data:
        for niv in levels_data["niveis"]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(0)  # Uma linha abaixo da outra, sem espaço extra
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.left_indent = Pt(54.0)       # Recuo da margem para o texto (0.75 in)
            p.paragraph_format.first_line_indent = Pt(-18.0) # Hanging Indent para o marcador (bullet)
            
            run_bullet = p.add_run("•\t")
            run_bullet.font.name = 'Aptos'
            run_bullet.font.size = Pt(11)
            
            run = p.add_run(niv)
            run.font.name = 'Aptos'
            run.font.size = Pt(11)
            
        # Exemplo (Sem pular linha antes)
        p_ex_label = doc.add_paragraph()
        p_ex_label.paragraph_format.left_indent = Pt(72.0)
        p_ex_label.paragraph_format.space_after = Pt(6)
        run_ex_lbl = p_ex_label.add_run("Exemplo:")
        run_ex_lbl.font.name = 'Aptos'
        run_ex_lbl.font.size = Pt(11)
        
        p_ex_val = doc.add_paragraph()
        p_ex_val.paragraph_format.left_indent = Pt(120.5)
        p_ex_val.paragraph_format.space_after = Pt(6)
        run_ex_val = p_ex_val.add_run(levels_data["ex_str"])
        run_ex_val.bold = True
        run_ex_val.font.name = 'Aptos'
        run_ex_val.font.size = Pt(11)
        
        p_ex_desc = doc.add_paragraph()
        p_ex_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_ex_desc.paragraph_format.left_indent = Pt(120.5)
        p_ex_desc.paragraph_format.space_after = Pt(6)
        p_ex_desc.paragraph_format.line_spacing = 1.15
        run_ex_desc = p_ex_desc.add_run(levels_data["ex_desc"])
        run_ex_desc.font.name = 'Aptos'
        run_ex_desc.font.size = Pt(11)
        
        # Monitoramento e Avaliação (Sem pular linha antes)
        p_mon = doc.add_paragraph()
        p_mon.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_mon.paragraph_format.space_after = Pt(6)
        p_mon.paragraph_format.line_spacing = 1.15
        run_mon_bold = p_mon.add_run("Monitoramento e Avaliação")
        run_mon_bold.bold = True
        run_mon_bold.font.name = 'Aptos'
        run_mon_bold.font.size = Pt(11)
        
        mon_term = "Relatório de Monitoramento" if "MONITORAMENTO" in report_config.key else "Relatório de Fiscalização"
        run_mon_text = p_mon.add_run(f" - Esta etapa é fundamental para garantir a eficácia das ações corretivas a serem executadas pela Concessionária para a melhoria contínua dos serviços prestados. Os principais instrumentos do Monitoramento e Avaliação são: Termo de Notificação e respectivo {mon_term}, Plano de Ação da Concessionária e Relatórios de Monitoramento e Avaliação Final.")
        run_mon_text.font.name = 'Aptos'
        run_mon_text.font.size = Pt(11)

    # ----------------------------------------------------
    # Seções extras pós metodologia (ex: Seção 4 do CRA)
    # ----------------------------------------------------
    extra_title, extra_paragraphs = report_config.get_post_metodologia_extra_paragraphs(row, data_extenso, total_achados)
    if extra_title:
        adicionar_titulo_secao(doc, extra_title)
        doc.add_paragraph()
        for paragraph_runs in extra_paragraphs:
            add_formatted_paragraph(paragraph_runs)
