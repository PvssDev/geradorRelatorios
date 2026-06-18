from utils import adicionar_titulo_secao
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd


def gerar_secao_resumo_nao_conformidades(doc, row, nao_conformidades_df):
    """
    Gera a seção '4. RESUMO DAS NÃO CONFORMIDADES IDENTIFICADAS' no formato visual do relatório oficial,
    com agrupamento por terminal e número sequencial.
    """

    espaco1 = doc.add_paragraph()
    espaco1.paragraph_format.space_after = Pt(12)

    adicionar_titulo_secao(doc, "4. RESUMO DAS NÃO CONFORMIDADES IDENTIFICADAS")

    id_fisc = row["ID da Fiscalização"]

    # Verifica de forma segura se a coluna existe em nao_conformidades_df
    if "ID da Fiscalização" in nao_conformidades_df.columns:
        nc_fisc = nao_conformidades_df[
            nao_conformidades_df["ID da Fiscalização"] == id_fisc
        ]
    else:
        nc_fisc = pd.DataFrame()

    if nc_fisc.empty:
        doc.add_paragraph("Nenhuma não conformidade registrada.")
        return

    if "Terminal" not in nc_fisc.columns or "Nº" not in nc_fisc.columns:
        doc.add_paragraph("⚠️ Colunas obrigatórias não encontradas na planilha.")
        return

    # Criar tabela
    tabela = doc.add_table(rows=1, cols=2)
    tabela.style = "Table Grid"
    tabela.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Cabeçalhos
    cabecalho = tabela.rows[0].cells
    cabecalho[0].text = "TERMINAL"
    cabecalho[1].text = "NÃO CONFORMIDADE"

    # Estilo cabeçalhos
    for cell in cabecalho:
        for par in cell.paragraphs:
            run = par.runs[0]
            run.bold = True
            run.font.size = Pt(11)

    # Agrupar por Terminal
    for terminal, grupo in nc_fisc.groupby("Terminal"):
        grupo = grupo.sort_values(by="Nº")

        # Extrair sigla
        if "(" in terminal and ")" in terminal:
            sigla_terminal = terminal.split("(")[-1].replace(")", "").strip()
        else:
            sigla_terminal = terminal[:3].upper()

        # Remover "TERMINAL DE"/"TERMINAL DO" e deixar maiúsculo
        nome_terminal = terminal.upper()
        nome_terminal = nome_terminal.replace("TERMINAL DE ", "").replace("TERMINAL DO ", "").strip()

        num_nc = 1  # contador sequencial para NCs

        for idx, (_, linha) in enumerate(grupo.iterrows()):
            row_cells = tabela.add_row().cells

            # Primeira linha do grupo: título do terminal
            if idx == 0:
                run_terminal = row_cells[0].paragraphs[0].add_run(nome_terminal)
                run_terminal.bold = True
                run_terminal.font.size = Pt(11)
            else:
                row_cells[0].text = ""

            # Coluna da NC
            descricao = linha["Não Conformidade"].strip()

            paragrafo_nc = row_cells[1].paragraphs[0]
            run_titulo = paragrafo_nc.add_run(f"{sigla_terminal} {num_nc}")
            run_titulo.bold = True
            run_titulo.font.size = Pt(11)

            run_desc = paragrafo_nc.add_run(f" – {descricao}")
            run_desc.font.size = Pt(11)

            num_nc += 1

    espaco = doc.add_paragraph()
    espaco.paragraph_format.space_after = Pt(24)
