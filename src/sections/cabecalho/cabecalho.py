from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from utils import formatar_mes_ano, extrair_ano, extrair_mes_ano_numerico
from database.manager import carregar_responsaveis

def gerar_capa_primeira_pagina(doc, logo_path, row, report_config, nc_df=None, documento_anterior=None):
    """
    Gera a primeira página do relatório (Capa) baseada no modelo de referência com dados dinâmicos.
    """
    if hasattr(report_config, "gerar_capa_monitoramento"):
        report_config.gerar_capa_monitoramento(doc, logo_path, row, documento_anterior)
    else:
        # 0. Ajuste de Estilo e Margens (Explicitas de 0.5 polegadas conforme a referência)
        style = doc.styles['Normal']
        style.font.name = 'Aptos'
        
        section = doc.sections[0]
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        # 1. Texto Superior (Sem fundo cinza)
        ano = extrair_ano(row["Data"])
        mes_ano = extrair_mes_ano_numerico(row["Data"])
        id_fisc = str(row.get("ID da Fiscalização", "")).strip()
        ctr_text = report_config.capa_ctr_number_template.format(ano=ano, id_fisc=id_fisc, mes_ano=mes_ano)
        
        p_ctr = doc.add_paragraph()
        p_ctr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_ctr.paragraph_format.space_before = Pt(6)
        p_ctr.paragraph_format.space_after = Pt(12)
        run_ctr = p_ctr.add_run(ctr_text)
        run_ctr.bold = True
        run_ctr.font.name = 'Aptos'
        run_ctr.font.size = Pt(11)
        
        # 1. Imagem da Capa (Aumentada em ~50% mantendo proporção nativa)
        if os.path.exists(logo_path):
            img_w = Inches(6.0)
            doc.add_picture(logo_path, width=img_w)
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        doc.add_paragraph()  # Linha vazia antes de FISCALIZAÇÃO...
        
        # 2. Títulos Principais
        titulos = report_config.get_capa_titulos(row, ano)
            
        for idx, titulo in enumerate(titulos):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(titulo)
            run.bold = True
            # Espaçamento pós parágrafo nas duas primeiras linhas, e um espaçamento maior (24pt) na última
            if idx < len(titulos) - 1:
                p.paragraph_format.space_after = Pt(6)
            else:
                p.paragraph_format.space_after = Pt(12)

        # 3. Lista de Analistas Dinâmica
        responsaveis_list = [r.strip() for r in str(row["Pessoal Responsável"]).split(",") if r.strip()]
        db_resp = carregar_responsaveis()
        
        analistas = []
        for nome in responsaveis_list:
            match = next((d for d in db_resp if d["nome"].strip().lower() == nome.lower()), None)
            if match:
                funcao = report_config.analyst_title if report_config.key == "SOCICAM" else match["funcao"]
                analistas.append((match["nome"], f"{funcao}, matrícula nº {match['matricula']}"))
            else:
                funcao = report_config.analyst_title if report_config.key == "SOCICAM" else "Analista de Regulação"
                analistas.append((nome, f"{funcao}, matrícula nº xxxxxxx/xx"))
        
        for i, (nome, cargo) in enumerate(analistas):
            # Nome em negrito
            p_nome = doc.add_paragraph()
            p_nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_nome = p_nome.add_run(nome)
            run_nome.bold = True
            run_nome.font.size = Pt(12)
            
            p_nome.paragraph_format.space_after = Pt(0)
            
            # Cargo normal
            p_cargo = doc.add_paragraph()
            p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cargo = p_cargo.add_run(cargo)
            run_cargo.font.size = Pt(12)
            
            if i == len(analistas) - 1:
                p_cargo.paragraph_format.space_after = Pt(12)  # Espaço após o parágrafo do último
            else:
                p_cargo.paragraph_format.space_after = Pt(12)  # Espaço abaixo de cada analista
        
        # 4. Data Dinâmica (Apenas para relatórios que contêm data central na capa, como CRA/CRC)
        if report_config.key != "SOCICAM":
            p_data = doc.add_paragraph()
            p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
            data_extenso = formatar_mes_ano(row["Data"])
            run_data = p_data.add_run(data_extenso)
            run_data.bold = True
            run_data.font.size = Pt(12)
            # Espaçamento superior e inferior para isolar a data nativamente
            p_data.paragraph_format.space_before = Pt(12)
            p_data.paragraph_format.space_after = Pt(12)
        
        # 5. Processo e SEI Dinâmicos
        rodape_textos = report_config.get_process_sei_texts(row, ano)
            
        for idx, texto in enumerate(rodape_textos):
            p_rodape = doc.add_paragraph()
            p_rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_rodape = p_rodape.add_run(texto)
            run_rodape.bold = True
            run_rodape.font.size = Pt(12)
            p_rodape.paragraph_format.space_after = Pt(0)
    
    # 6. Geração de Sumário e Lista de Abreviaturas com ordem condicional
    has_abbr = bool(report_config.get_abbreviations())
    if not has_abbr:
        # Apenas Sumário na Página 2, sem Lista de Abreviaturas
        doc.add_page_break()
        gerar_sumario(doc, row, report_config, nc_df=nc_df)
    elif report_config.sumario_before_abreviaturas:
        # CRC/SOCICAM: Sumário na Página 2, Lista de Abreviaturas na Página 3
        doc.add_page_break()
        gerar_sumario(doc, row, report_config, nc_df=nc_df)
        
        doc.add_page_break()
        gerar_lista_abreviaturas(doc, report_config)
    else:
        # CRA: Lista de Abreviaturas na Página 2, Sumário na Página 3
        doc.add_page_break()
        gerar_lista_abreviaturas(doc, report_config)
        
        doc.add_page_break()
        gerar_sumario(doc, row, report_config, nc_df=nc_df)
        
    # 7. Quebra de seção para ir para a Página 4 (Introdução)
    doc.add_section(WD_SECTION.NEW_PAGE)


def gerar_lista_abreviaturas(doc, report_config):
    """Gera a segunda página do cabeçalho (Lista de Abreviaturas e Siglas)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LISTA DE ABREVIATURAS E SIGLAS")
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(24)
    
    abreviaturas = report_config.get_abbreviations()
        
    table = doc.add_table(rows=1 + len(abreviaturas), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    # Cabeçalho da tabela
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'SIGLA'
    hdr_cells[1].text = 'DEFINIÇÃO'
    
    # Formatação do cabeçalho
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(11)
            
    # Preencher dados
    for idx, (sigla, definicao) in enumerate(abreviaturas):
        row_cells = table.rows[idx + 1].cells
        row_cells[0].text = sigla
        row_cells[1].text = definicao
        
        # Alinhamento
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Garantir tamanho da fonte e normalizar estilo
        for cell in row_cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(11)
                
    # Definir larguras das colunas
    col_widths = [Inches(0.65), Inches(4.38)] if report_config.key == "SOCICAM" else [Inches(1.0), Inches(5.0)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width


def habilitar_atualizacao_campos(doc):
    """Garante que o Word atualize os campos nativos (como Sumário/TOC) automaticamente ao abrir o documento."""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    try:
        settings = doc.settings.element
        if settings.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}updateFields') is None:
            update_fields = parse_xml('<w:updateFields %s w:val="true"/>' % nsdecls('w'))
            settings.append(update_fields)
    except Exception as e:
        print(f"Aviso: Não foi possível configurar updateFields: {e}")


def gerar_sumario(doc, row, report_config, nc_df=None):
    """Gera o Sumário como ferramenta nativa do Word (Table of Contents / SDT / Campo TOC) com atualização automática."""
    import inspect
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # Habilita a atualização automática de campos pelo Word ao abrir o arquivo
    habilitar_atualizacao_campos(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SUMÁRIO")
    run.bold = True
    run.font.name = 'Aptos'
    run.font.size = Pt(14)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(24)
    
    sig = inspect.signature(report_config.get_sumario_linhas)
    if "nc_df" in sig.parameters:
        linhas = report_config.get_sumario_linhas(row, nc_df=nc_df)
    else:
        linhas = report_config.get_sumario_linhas(row)
        
    sdt_xml_parts = [
        f'<w:sdt {nsdecls("w")}>',
        '  <w:sdtPr>',
        '    <w:id w:val="914752099"/>',
        '    <w:docPartObj>',
        '      <w:docPartGallery w:val="Table of Contents"/>',
        '      <w:docPartUnique/>',
        '    </w:docPartObj>',
        '  </w:sdtPr>',
        '  <w:sdtEndPr>',
        '    <w:rPr>',
        '      <w:rFonts w:ascii="Aptos" w:hAnsi="Aptos"/>',
        '      <w:b/>',
        '      <w:bCs/>',
        '    </w:rPr>',
        '  </w:sdtEndPr>',
        '  <w:sdtContent>'
    ]

    for idx, linha in enumerate(linhas):
        parts = linha.split('\t')
        num_part = parts[0].strip() if len(parts) > 0 and parts[0].strip() else ""
        titulo_part = parts[1].strip() if len(parts) > 1 else parts[0].strip()
        page_part = parts[2].strip() if len(parts) > 2 else "4"
        
        if idx == 0:
            toc_begin_xml = (
                '      <w:r><w:rPr><w:rFonts w:ascii="Aptos" w:hAnsi="Aptos"/></w:rPr>'
                '<w:fldChar w:fldCharType="begin"/></w:r>'
                '      <w:r><w:rPr><w:rFonts w:ascii="Aptos" w:hAnsi="Aptos"/></w:rPr>'
                '<w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText></w:r>'
                '      <w:r><w:rPr><w:rFonts w:ascii="Aptos" w:hAnsi="Aptos"/></w:rPr>'
                '<w:fldChar w:fldCharType="separate"/></w:r>'
            )
        else:
            toc_begin_xml = ""

        p_xml = [
            '    <w:p>',
            '      <w:pPr>',
            '        <w:pStyle w:val="Sumrio3"/>',
            '        <w:tabs>',
            '          <w:tab w:val="left" w:pos="960"/>',
            '          <w:tab w:val="right" w:leader="dot" w:pos="10456"/>',
            '        </w:tabs>',
            '        <w:rPr><w:rFonts w:ascii="Aptos" w:hAnsi="Aptos"/><w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr>',
            '      </w:pPr>',
            toc_begin_xml,
            '      <w:hyperlink w:history="1">'
        ]
        
        if num_part:
            p_xml.append(
                f'        <w:r><w:rPr><w:rFonts w:ascii="Aptos" w:hAnsi="Aptos"/><w:b/><w:sz w:val="22"/></w:rPr>'
                f'<w:t>{num_part}</w:t></w:r>'
                f'        <w:r><w:tab/></w:r>'
            )
            
        p_xml.append(
            f'        <w:r><w:rPr><w:rFonts w:ascii="Aptos" w:hAnsi="Aptos"/><w:b/><w:sz w:val="22"/></w:rPr>'
            f'<w:t>{titulo_part}</w:t></w:r>'
            f'        <w:r><w:tab/></w:r>'
            f'        <w:r><w:rPr><w:rFonts w:ascii="Aptos" w:hAnsi="Aptos"/><w:b/><w:sz w:val="22"/></w:rPr>'
            f'<w:t>{page_part}</w:t></w:r>'
            '      </w:hyperlink>'
        )
        p_xml.append('    </w:p>')
        sdt_xml_parts.extend(p_xml)

    sdt_xml_parts.append(
        '    <w:p>'
        '      <w:r><w:rPr><w:rFonts w:ascii="Aptos" w:hAnsi="Aptos"/><w:b/></w:rPr>'
        '        <w:fldChar w:fldCharType="end"/>'
        '      </w:r>'
        '    </w:p>'
    )
    sdt_xml_parts.append('  </w:sdtContent>')
    sdt_xml_parts.append('</w:sdt>')

    full_sdt_xml = "\n".join(sdt_xml_parts)
    sdt_element = parse_xml(full_sdt_xml)
    p._p.addnext(sdt_element)
